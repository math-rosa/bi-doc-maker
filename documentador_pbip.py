"""
Documentador de Projetos Power BI (.pbip)

Script para gerar documentação automática em Markdown de projetos Power BI Desktop.
Extrai informações do modelo semântico e relatório.
"""

import os
import json
import re
import sys
import base64
import html
import unicodedata
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Tuple

import i18n
from i18n import t as _t


# ============================================================================
# DATACLASSES - Estruturas de Dados
# ============================================================================

@dataclass
class InfoColuna:
    """Informações de uma coluna regular de tabela"""
    nome: str
    tipo_dado: str = "string"
    formato: Optional[str] = None
    sumarizacao: str = "none"
    coluna_fonte: Optional[str] = None
    esta_oculta: bool = False
    categoria_dados: Optional[str] = None


@dataclass
class InfoColunaCalculada:
    """Informações de uma coluna calculada (com DAX)"""
    nome: str
    expressao_dax: str
    tipo_dado: str = "string"
    formato: Optional[str] = None
    sumarizacao: str = "none"


@dataclass
class InfoMedida:
    """Informações de uma medida DAX"""
    nome: str
    expressao_dax: str
    formato: Optional[str] = None
    formato_dinamico: Optional[str] = None
    descricao: Optional[str] = None


@dataclass
class InfoHierarquia:
    """Informações de uma hierarquia"""
    nome: str
    niveis: List[str] = field(default_factory=list)


@dataclass
class InfoParticao:
    """Informações da partição (fonte de dados)"""
    nome: str
    modo: str = "import"
    grupo_consulta: Optional[str] = None
    codigo_fonte: str = ""


@dataclass
class InfoTabela:
    """Informações completas de uma tabela"""
    nome: str
    esta_oculta: bool = False
    excluida_refresh: bool = False
    descricao: Optional[str] = None
    eh_fato: bool = False
    colunas: List[InfoColuna] = field(default_factory=list)
    medidas: List[InfoMedida] = field(default_factory=list)
    colunas_calculadas: List[InfoColunaCalculada] = field(default_factory=list)
    hierarquias: List[InfoHierarquia] = field(default_factory=list)
    particao: Optional[InfoParticao] = None


@dataclass
class InfoRelacionamento:
    """Informações de um relacionamento entre tabelas"""
    id: str
    tabela_origem: str
    coluna_origem: str
    tabela_destino: str
    coluna_destino: str
    filtro_bidirecional: bool = False
    esta_ativo: bool = True
    comportamento_data: Optional[str] = None


@dataclass
class InfoModelo:
    """Informações gerais do modelo semântico"""
    cultura: str = "pt-BR"
    versao_datasource: str = ""
    grupos_consulta: List[Dict] = field(default_factory=list)
    tabelas_referenciadas: List[str] = field(default_factory=list)
    annotations: Dict = field(default_factory=dict)


@dataclass
class BrandingConfig:
    """Configuracao visual usada na documentacao gerada."""
    document_title: str = "Documentação Power BI"
    logo_path: Optional[Path] = None
    primary_color: str = "#003D6B"
    secondary_color: str = "#006DAA"
    light_color: str = "#D6E8F5"


@dataclass
class OcorrenciaTermo:
    """Local onde um termo do dicionario foi encontrado."""
    origem: str
    tipo_origem: str
    nome_objeto: str
    contexto: str = ""


@dataclass
class TermoDicionario:
    """Termo de negocio inferido a partir dos metadados do PBIP."""
    termo: str
    frequencia: int = 0
    score: float = 0.0
    categoria: str = "Termo Geral"
    fontes: List[str] = field(default_factory=list)
    exemplos: List[str] = field(default_factory=list)
    ocorrencias: List[OcorrenciaTermo] = field(default_factory=list)


@dataclass
class EtapaPowerQuery:
    """Leitura estática de uma etapa Power Query M."""
    nome: str
    expressao: str
    categoria: str
    descricao: str


@dataclass
class LinhaRegraPowerQuery:
    """Linha executiva da regra inferida do Power Query M."""
    etapa: str
    regra: str
    descricao: str


@dataclass
class RegraPowerQuery:
    """Resumo offline inferido a partir do Power Query M."""
    observacoes: Dict[str, List[str]] = field(default_factory=dict)
    funcoes_relevantes: Dict[str, List[str]] = field(default_factory=dict)
    origem_dados: List[str] = field(default_factory=list)
    filtros_aplicados: List[str] = field(default_factory=list)
    transformacoes: List[str] = field(default_factory=list)
    integracoes: List[str] = field(default_factory=list)
    resultado_final: str = ""
    etapa_final: str = ""
    etapas: List[EtapaPowerQuery] = field(default_factory=list)
    linhas_regra: List[LinhaRegraPowerQuery] = field(default_factory=list)


@dataclass
class LeituraDax:
    """Leitura técnica simples de uma expressão DAX."""
    funcoes: List[str] = field(default_factory=list)
    itens: List[str] = field(default_factory=list)
    categorias: Dict[str, List[str]] = field(default_factory=dict)
    # Candidatos extraidos via regex (validar depois contra o modelo)
    referencias_medidas: List[str] = field(default_factory=list)  # [Nome] sem prefixo de tabela
    referencias_tabelas: List[str] = field(default_factory=list)  # Tabela[col] ou 'Tabela'[col]
    variaveis_dax: List[str] = field(default_factory=list)  # VAR nome (excluir de medidas)


@dataclass(frozen=True)
class CatalogoFuncao:
    """Entrada de catálogo offline para leitura técnica de M/DAX."""
    nome: str
    categoria: str
    descricao: str
    leitura_negocio: str
    tags: Tuple[str, ...] = ()


POWER_QUERY_CATEGORY_LABELS = {
    "origem": "Origens e conexões",
    "filtro": "Filtros e redução de linhas",
    "integracao": "Integrações entre consultas",
    "transformacao": "Transformações de estrutura",
    "agregacao": "Agregações e modelagem",
    "qualidade": "Tratamento e qualidade de dados",
    "logica": "Lógica e regras condicionais",
}


DAX_CATEGORY_LABELS = {
    "contexto_filtro": "Contexto e filtros",
    "agregacao": "Agregações e iteradores",
    "relacionamento": "Relacionamentos",
    "tempo": "Inteligência temporal",
    "logica": "Lógica e tratamento",
    "texto": "Texto e apresentação",
    "tabela": "Funções tabulares",
}


DAX_CATEGORY_ORDER = [
    "contexto_filtro",
    "agregacao",
    "relacionamento",
    "tempo",
    "logica",
    "texto",
    "tabela",
]


def _catalogo_funcao(
    categoria: str,
    descricao: str,
    leitura_negocio: str,
    *tags: str,
) -> CatalogoFuncao:
    return CatalogoFuncao("", categoria, descricao, leitura_negocio, tuple(tags))


def _catalogo_com_nomes(definicoes: Dict[str, CatalogoFuncao]) -> Dict[str, CatalogoFuncao]:
    return {
        nome: CatalogoFuncao(
            nome=nome,
            categoria=item.categoria,
            descricao=item.descricao,
            leitura_negocio=item.leitura_negocio,
            tags=item.tags,
        )
        for nome, item in definicoes.items()
    }


# Catálogos curados a partir das referências oficiais Microsoft Learn.
# As descrições abaixo são resumos próprios em PT-BR para uso offline.
POWER_QUERY_FUNCTION_CATALOG = _catalogo_com_nomes({
    "Sql.Database": _catalogo_funcao("origem", "Conecta a uma base SQL Server.", "conecta a consulta a uma base SQL Server.", "sql", "banco"),
    "Oracle.Database": _catalogo_funcao("origem", "Conecta a uma base Oracle.", "conecta a consulta a uma base Oracle.", "banco"),
    "PostgreSQL.Database": _catalogo_funcao("origem", "Conecta a uma base PostgreSQL.", "conecta a consulta a uma base PostgreSQL.", "banco"),
    "MySQL.Database": _catalogo_funcao("origem", "Conecta a uma base MySQL.", "conecta a consulta a uma base MySQL.", "banco"),
    "Snowflake.Databases": _catalogo_funcao("origem", "Lista bancos e objetos Snowflake.", "usa Snowflake como origem corporativa de dados.", "banco", "cloud"),
    "AnalysisServices.Database": _catalogo_funcao("origem", "Conecta ao Analysis Services ou modelo tabular.", "lê dados de um modelo analítico existente.", "modelo"),
    "Odbc.DataSource": _catalogo_funcao("origem", "Conecta a uma fonte via ODBC.", "usa um conector ODBC para acessar dados externos.", "conector"),
    "OleDb.DataSource": _catalogo_funcao("origem", "Conecta a uma fonte via OLE DB.", "usa um conector OLE DB para acessar dados externos.", "conector"),
    "Excel.Workbook": _catalogo_funcao("origem", "Interpreta conteúdo de arquivo Excel.", "lê planilhas, tabelas ou intervalos de um arquivo Excel.", "arquivo"),
    "Csv.Document": _catalogo_funcao("origem", "Interpreta conteúdo CSV.", "lê dados tabulares de um arquivo CSV.", "arquivo"),
    "Json.Document": _catalogo_funcao("origem", "Interpreta conteúdo JSON.", "lê dados estruturados em JSON.", "arquivo", "api"),
    "Web.Contents": _catalogo_funcao("origem", "Obtém conteúdo por HTTP/HTTPS.", "consulta uma URL, API ou arquivo publicado na web.", "api", "web"),
    "OData.Feed": _catalogo_funcao("origem", "Lê feed OData.", "consome dados expostos por serviço OData.", "api"),
    "SharePoint.Files": _catalogo_funcao("origem", "Lista arquivos de um site SharePoint.", "usa arquivos armazenados no SharePoint como origem.", "sharepoint"),
    "SharePoint.Contents": _catalogo_funcao("origem", "Navega conteúdo de um site SharePoint.", "navega pastas e bibliotecas do SharePoint.", "sharepoint"),
    "Folder.Files": _catalogo_funcao("origem", "Lista arquivos de uma pasta.", "combina ou seleciona arquivos de um diretório.", "arquivo"),
    "File.Contents": _catalogo_funcao("origem", "Lê bytes de um arquivo.", "carrega um arquivo local ou de rede.", "arquivo"),
    "PowerPlatform.Dataflows": _catalogo_funcao("origem", "Acessa dataflows da Power Platform.", "reaproveita dados preparados em dataflows.", "dataflow"),
    "Dataverse.Contents": _catalogo_funcao("origem", "Acessa tabelas do Dataverse.", "lê dados de entidades Dataverse.", "dataverse"),
    "Value.NativeQuery": _catalogo_funcao("origem", "Executa consulta nativa na fonte.", "usa SQL ou linguagem nativa da origem para trazer dados já tratados.", "sql", "native"),

    "Table.SelectRows": _catalogo_funcao("filtro", "Mantém linhas que atendem a uma condição.", "reduz a base para registros aderentes à regra de filtro.", "where"),
    "Table.Skip": _catalogo_funcao("filtro", "Ignora linhas iniciais.", "descarta linhas por posição, comum em arquivos com cabeçalho extra.", "linha"),
    "Table.FirstN": _catalogo_funcao("filtro", "Mantém as primeiras N linhas.", "limita a amostra ou recorte inicial da tabela.", "linha"),
    "Table.LastN": _catalogo_funcao("filtro", "Mantém as últimas N linhas.", "limita a amostra ou recorte final da tabela.", "linha"),
    "Table.RemoveRows": _catalogo_funcao("filtro", "Remove linhas por posição ou quantidade.", "descarta registros por posição na tabela.", "linha"),
    "Table.RemoveRowsWithErrors": _catalogo_funcao("qualidade", "Remove linhas com erro.", "elimina registros inválidos gerados por conversão ou cálculo.", "erro"),
    "Table.SelectRowsWithErrors": _catalogo_funcao("qualidade", "Seleciona linhas com erro.", "isola registros problemáticos para análise ou tratamento.", "erro"),

    "Table.SelectColumns": _catalogo_funcao("transformacao", "Mantém apenas colunas selecionadas.", "define quais campos seguem para o modelo.", "coluna"),
    "Table.RemoveColumns": _catalogo_funcao("transformacao", "Remove colunas da tabela.", "descarta campos que não entram na modelagem final.", "coluna"),
    "Table.RenameColumns": _catalogo_funcao("transformacao", "Renomeia colunas.", "padroniza nomes de campos para leitura do negócio.", "coluna"),
    "Table.TransformColumnTypes": _catalogo_funcao("transformacao", "Altera tipos de dados.", "define tipos corretos para datas, números e textos.", "tipo"),
    "Table.TransformColumns": _catalogo_funcao("transformacao", "Aplica transformação a colunas.", "normaliza ou calcula valores em campos específicos.", "coluna"),
    "Table.AddColumn": _catalogo_funcao("transformacao", "Cria coluna calculada no Power Query.", "deriva um novo atributo a partir de regras de preparação.", "coluna"),
    "Table.DuplicateColumn": _catalogo_funcao("transformacao", "Duplica uma coluna.", "preserva um campo original antes de nova transformação.", "coluna"),
    "Table.SplitColumn": _catalogo_funcao("transformacao", "Divide uma coluna em outras colunas.", "separa informações compostas em campos analíticos.", "coluna"),
    "Table.CombineColumns": _catalogo_funcao("transformacao", "Combina colunas em uma coluna.", "gera campo consolidado a partir de múltiplos atributos.", "coluna"),
    "Table.ReorderColumns": _catalogo_funcao("transformacao", "Reordena colunas.", "organiza a estrutura final das colunas.", "coluna"),
    "Table.Sort": _catalogo_funcao("transformacao", "Ordena linhas.", "aplica ordenação por campos relevantes.", "ordem"),
    "Table.PromoteHeaders": _catalogo_funcao("transformacao", "Promove primeira linha para cabeçalhos.", "usa a primeira linha do arquivo como nomes de colunas.", "cabecalho"),
    "Table.DemoteHeaders": _catalogo_funcao("transformacao", "Rebaixa cabeçalhos para linha de dados.", "transforma nomes de colunas em registros.", "cabecalho"),
    "Table.Buffer": _catalogo_funcao("transformacao", "Materializa tabela em memória.", "estabiliza a leitura de uma etapa antes das próximas transformações.", "performance"),

    "Table.NestedJoin": _catalogo_funcao("integracao", "Cria junção com outra tabela em coluna aninhada.", "relaciona consultas usando chaves de negócio.", "join"),
    "Table.Join": _catalogo_funcao("integracao", "Une duas tabelas por chaves.", "combina atributos de tabelas relacionadas.", "join"),
    "Table.Combine": _catalogo_funcao("integracao", "Anexa tabelas com estrutura compatível.", "empilha consultas para formar uma base consolidada.", "append"),
    "Table.ExpandTableColumn": _catalogo_funcao("integracao", "Expande coluna de tabela aninhada.", "traz campos de consulta relacionada após merge.", "expand"),
    "Table.ExpandRecordColumn": _catalogo_funcao("integracao", "Expande coluna de registro.", "abre atributos internos de um registro estruturado.", "expand"),
    "Table.ExpandListColumn": _catalogo_funcao("integracao", "Expande listas em múltiplas linhas.", "transforma itens de lista em registros analisáveis.", "expand"),

    "Table.Group": _catalogo_funcao("agregacao", "Agrupa linhas e calcula agregações.", "resume dados por chaves de análise.", "group"),
    "Table.Pivot": _catalogo_funcao("agregacao", "Transforma valores de linhas em colunas.", "reorganiza categorias como colunas analíticas.", "pivot"),
    "Table.Unpivot": _catalogo_funcao("agregacao", "Transforma colunas em linhas.", "normaliza colunas repetidas para formato longo.", "unpivot"),
    "Table.UnpivotOtherColumns": _catalogo_funcao("agregacao", "Mantém algumas colunas e despivota as demais.", "normaliza colunas de métricas mantendo chaves fixas.", "unpivot"),
    "Table.Distinct": _catalogo_funcao("qualidade", "Remove registros duplicados.", "mantém combinações únicas de registros.", "duplicidade"),
    "Table.RemoveDuplicates": _catalogo_funcao("qualidade", "Remove duplicidades considerando colunas.", "garante unicidade por uma ou mais chaves.", "duplicidade"),
    "Table.ReplaceValue": _catalogo_funcao("qualidade", "Substitui valores na tabela.", "padroniza ou corrige valores em campos selecionados.", "limpeza"),
    "Table.ReplaceErrorValues": _catalogo_funcao("qualidade", "Substitui erros por valores definidos.", "trata falhas de conversão sem descartar a linha.", "erro"),
    "Table.FillDown": _catalogo_funcao("qualidade", "Preenche valores nulos com valor acima.", "propaga valores de cabeçalho/grupo para linhas inferiores.", "nulo"),
    "Table.FillUp": _catalogo_funcao("qualidade", "Preenche valores nulos com valor abaixo.", "propaga valores posteriores para registros anteriores.", "nulo"),

    "Text.Contains": _catalogo_funcao("filtro", "Verifica se texto contém trecho.", "aplica regra textual de presença de palavra ou código.", "texto"),
    "Text.StartsWith": _catalogo_funcao("filtro", "Verifica prefixo de texto.", "filtra ou classifica registros por início do campo.", "texto"),
    "Text.EndsWith": _catalogo_funcao("filtro", "Verifica sufixo de texto.", "filtra ou classifica registros por final do campo.", "texto"),
    "Text.Upper": _catalogo_funcao("qualidade", "Converte texto para maiúsculas.", "padroniza caixa de texto.", "texto"),
    "Text.Lower": _catalogo_funcao("qualidade", "Converte texto para minúsculas.", "padroniza caixa de texto.", "texto"),
    "Text.Trim": _catalogo_funcao("qualidade", "Remove espaços nas extremidades.", "limpa espaços indevidos em campos de texto.", "texto"),
    "Text.Clean": _catalogo_funcao("qualidade", "Remove caracteres não imprimíveis.", "higieniza campos importados de arquivos ou sistemas.", "texto"),
    "Text.Replace": _catalogo_funcao("qualidade", "Substitui trechos de texto.", "corrige padrões textuais em campos.", "texto"),
    "Text.From": _catalogo_funcao("transformacao", "Converte valor para texto.", "prepara valores para concatenação, limpeza ou identificação.", "texto"),

    "Date.From": _catalogo_funcao("transformacao", "Converte valor para data.", "garante tipo de data para análise temporal.", "data"),
    "DateTime.From": _catalogo_funcao("transformacao", "Converte valor para data/hora.", "garante granularidade temporal com horário.", "data"),
    "DateTime.LocalNow": _catalogo_funcao("transformacao", "Retorna data/hora local atual.", "usa o momento da atualização como referência temporal.", "data"),
    "DateTime.Date": _catalogo_funcao("transformacao", "Extrai a data de um valor data/hora.", "remove a granularidade de horário para análise por data.", "data"),
    "Date.StartOfMonth": _catalogo_funcao("transformacao", "Retorna início do mês.", "cria referência de competência mensal.", "data"),
    "Date.EndOfMonth": _catalogo_funcao("transformacao", "Retorna fim do mês.", "cria referência de fechamento mensal.", "data"),
    "Date.Year": _catalogo_funcao("transformacao", "Extrai ano.", "cria atributo anual para segmentação.", "data"),
    "Date.Month": _catalogo_funcao("transformacao", "Extrai mês.", "cria atributo mensal para segmentação.", "data"),
    "Date.AddDays": _catalogo_funcao("transformacao", "Desloca data em dias.", "calcula datas relativas por quantidade de dias.", "data"),
    "Date.AddMonths": _catalogo_funcao("transformacao", "Desloca data em meses.", "calcula datas relativas por quantidade de meses.", "data"),

    "List.Sum": _catalogo_funcao("agregacao", "Soma valores de uma lista.", "calcula total de um conjunto de valores.", "lista"),
    "List.Count": _catalogo_funcao("agregacao", "Conta itens de uma lista.", "calcula quantidade de itens.", "lista"),
    "List.Distinct": _catalogo_funcao("qualidade", "Remove itens duplicados de lista.", "mantém valores únicos em uma lista.", "lista"),
    "Record.Field": _catalogo_funcao("transformacao", "Obtém campo de um registro.", "extrai atributo específico de estrutura record.", "record"),
    "Record.ToTable": _catalogo_funcao("transformacao", "Converte registro em tabela.", "transforma estrutura de atributos em linhas.", "record"),
    "try otherwise": _catalogo_funcao("logica", "Trata erro com valor alternativo.", "define regra de contingência quando uma transformação falha.", "erro"),
    "if then else": _catalogo_funcao("logica", "Aplica condição com resultado alternativo.", "implementa regra condicional de negócio.", "condicao"),
})


DAX_FUNCTION_CATALOG = _catalogo_com_nomes({
    "CALCULATE": _catalogo_funcao("contexto_filtro", "Avalia expressão em contexto de filtro modificado.", "recalcula a métrica aplicando regras específicas de filtro.", "contexto"),
    "CALCULATETABLE": _catalogo_funcao("contexto_filtro", "Avalia tabela em contexto de filtro modificado.", "gera uma tabela intermediária com filtros específicos.", "contexto"),
    "FILTER": _catalogo_funcao("contexto_filtro", "Retorna tabela filtrada por expressão.", "aplica uma regra linha a linha para limitar dados.", "filtro"),
    "ALL": _catalogo_funcao("contexto_filtro", "Remove filtros de tabela ou coluna.", "calcula indicador ignorando filtros selecionados.", "contexto"),
    "ALLEXCEPT": _catalogo_funcao("contexto_filtro", "Remove filtros exceto os informados.", "preserva apenas dimensões relevantes no cálculo.", "contexto"),
    "ALLSELECTED": _catalogo_funcao("contexto_filtro", "Considera filtros externos selecionados.", "calcula respeitando seleção visual mais ampla.", "contexto"),
    "REMOVEFILTERS": _catalogo_funcao("contexto_filtro", "Remove filtros de colunas ou tabelas.", "limpa filtros para obter base de comparação.", "contexto"),
    "KEEPFILTERS": _catalogo_funcao("contexto_filtro", "Preserva filtros existentes ao adicionar novos.", "refina o contexto sem sobrescrever seleções.", "contexto"),
    "VALUES": _catalogo_funcao("contexto_filtro", "Retorna valores distintos visíveis.", "obtém conjunto de valores disponíveis no contexto atual.", "contexto"),
    "SELECTEDVALUE": _catalogo_funcao("contexto_filtro", "Retorna valor único selecionado ou alternativa.", "lê uma seleção única feita pelo usuário ou contexto.", "contexto"),
    "HASONEVALUE": _catalogo_funcao("contexto_filtro", "Verifica se há apenas um valor no contexto.", "valida se o cálculo está em granularidade única.", "contexto"),
    "ISFILTERED": _catalogo_funcao("contexto_filtro", "Verifica filtro direto em coluna/tabela.", "identifica se o usuário ou visual aplicou filtro.", "contexto"),
    "TREATAS": _catalogo_funcao("contexto_filtro", "Aplica valores como filtro em colunas não relacionadas diretamente.", "simula relacionamento lógico entre conjuntos.", "contexto"),

    "SUM": _catalogo_funcao("agregacao", "Soma uma coluna numérica.", "calcula total direto de um campo.", "agregacao"),
    "SUMX": _catalogo_funcao("agregacao", "Itera tabela e soma expressão.", "calcula total linha a linha antes de somar.", "iterador"),
    "AVERAGE": _catalogo_funcao("agregacao", "Calcula média de uma coluna.", "obtém valor médio direto de um campo.", "agregacao"),
    "AVERAGEX": _catalogo_funcao("agregacao", "Itera tabela e calcula média de expressão.", "calcula média a partir de regra linha a linha.", "iterador"),
    "COUNT": _catalogo_funcao("agregacao", "Conta valores numéricos.", "mede quantidade de valores preenchidos.", "agregacao"),
    "COUNTROWS": _catalogo_funcao("agregacao", "Conta linhas de tabela.", "mede quantidade de registros no contexto.", "agregacao"),
    "DISTINCTCOUNT": _catalogo_funcao("agregacao", "Conta valores distintos.", "mede quantidade única de entidades ou chaves.", "agregacao"),
    "MIN": _catalogo_funcao("agregacao", "Retorna menor valor.", "identifica valor mínimo no contexto.", "agregacao"),
    "MAX": _catalogo_funcao("agregacao", "Retorna maior valor.", "identifica valor máximo no contexto.", "agregacao"),
    "MINX": _catalogo_funcao("agregacao", "Itera tabela e retorna menor expressão.", "calcula mínimo após regra linha a linha.", "iterador"),
    "MAXX": _catalogo_funcao("agregacao", "Itera tabela e retorna maior expressão.", "calcula máximo após regra linha a linha.", "iterador"),

    "RELATED": _catalogo_funcao("relacionamento", "Busca valor em tabela relacionada.", "usa relacionamento do modelo para trazer atributo de dimensão.", "relacao"),
    "RELATEDTABLE": _catalogo_funcao("relacionamento", "Retorna linhas relacionadas.", "navega do registro atual para registros relacionados.", "relacao"),
    "USERELATIONSHIP": _catalogo_funcao("relacionamento", "Ativa relacionamento específico no cálculo.", "usa relacionamento alternativo para calcular a métrica.", "relacao"),
    "CROSSFILTER": _catalogo_funcao("relacionamento", "Altera direção de filtro de relacionamento.", "controla propagação de filtro durante o cálculo.", "relacao"),
    "LOOKUPVALUE": _catalogo_funcao("relacionamento", "Busca valor por pares de chave/valor.", "recupera atributo quando a regra usa correspondência explícita.", "relacao"),

    "DATEADD": _catalogo_funcao("tempo", "Desloca datas por intervalo.", "compara períodos deslocados no tempo.", "tempo"),
    "DATESYTD": _catalogo_funcao("tempo", "Retorna datas do ano até a data atual.", "cria cálculo acumulado no ano.", "tempo"),
    "DATESMTD": _catalogo_funcao("tempo", "Retorna datas do mês até a data atual.", "cria cálculo acumulado no mês.", "tempo"),
    "DATESQTD": _catalogo_funcao("tempo", "Retorna datas do trimestre até a data atual.", "cria cálculo acumulado no trimestre.", "tempo"),
    "TOTALYTD": _catalogo_funcao("tempo", "Calcula expressão acumulada no ano.", "gera indicador year-to-date.", "tempo"),
    "SAMEPERIODLASTYEAR": _catalogo_funcao("tempo", "Retorna período equivalente no ano anterior.", "compara desempenho com o mesmo período do ano anterior.", "tempo"),
    "PREVIOUSMONTH": _catalogo_funcao("tempo", "Retorna mês anterior.", "compara ou calcula período mensal anterior.", "tempo"),
    "NEXTMONTH": _catalogo_funcao("tempo", "Retorna mês seguinte.", "projeta ou compara período mensal seguinte.", "tempo"),
    "STARTOFMONTH": _catalogo_funcao("tempo", "Retorna primeira data do mês.", "ancora cálculo no início da competência.", "tempo"),
    "ENDOFMONTH": _catalogo_funcao("tempo", "Retorna última data do mês.", "ancora cálculo no fechamento da competência.", "tempo"),
    "TODAY": _catalogo_funcao("tempo", "Retorna a data atual.", "usa a data de atualização como referência de regra.", "tempo"),
    "NOW": _catalogo_funcao("tempo", "Retorna data e hora atuais.", "usa o momento da atualização como referência de regra.", "tempo"),
    "YEAR": _catalogo_funcao("tempo", "Extrai o ano de uma data.", "segmenta ou compara resultados por ano.", "tempo"),
    "MONTH": _catalogo_funcao("tempo", "Extrai o mês de uma data.", "segmenta ou compara resultados por mês.", "tempo"),

    "IF": _catalogo_funcao("logica", "Aplica regra condicional simples.", "retorna resultado conforme condição de negócio.", "condicao"),
    "SWITCH": _catalogo_funcao("logica", "Seleciona resultado entre múltiplos casos.", "modela classificação ou regra com vários cenários.", "condicao"),
    "DIVIDE": _catalogo_funcao("logica", "Divide com tratamento para erro/divisão por zero.", "calcula razão ou percentual com fallback seguro.", "calculo"),
    "COALESCE": _catalogo_funcao("logica", "Retorna primeiro valor não vazio.", "define valor alternativo para lacunas.", "nulo"),
    "ISBLANK": _catalogo_funcao("logica", "Verifica valor vazio.", "trata ausência de valor no cálculo.", "nulo"),
    "BLANK": _catalogo_funcao("logica", "Retorna valor vazio.", "omite resultado quando a regra não deve exibir valor.", "nulo"),
    "AND": _catalogo_funcao("logica", "Combina condições obrigatórias.", "exige que múltiplas regras sejam verdadeiras.", "condicao"),
    "OR": _catalogo_funcao("logica", "Combina condições alternativas.", "aceita qualquer regra verdadeira.", "condicao"),

    "FORMAT": _catalogo_funcao("texto", "Formata valor como texto.", "controla apresentação textual de número, data ou medida.", "texto"),
    "CONCATENATEX": _catalogo_funcao("texto", "Concatena expressão sobre linhas.", "gera lista textual a partir de registros.", "texto"),
    "ISINSCOPE": _catalogo_funcao("tabela", "Verifica se coluna está no nível atual de hierarquia.", "ajusta cálculo conforme granularidade visível no relatório.", "hierarquia"),
    "RANKX": _catalogo_funcao("tabela", "Calcula ranking sobre tabela.", "classifica entidades por métrica.", "ranking"),
    "TOPN": _catalogo_funcao("tabela", "Retorna os N maiores/menores registros.", "seleciona principais itens conforme ordenação.", "ranking"),
    "SUMMARIZE": _catalogo_funcao("tabela", "Cria tabela resumida por grupos.", "gera tabela agregada para cálculo intermediário.", "tabela"),
    "ADDCOLUMNS": _catalogo_funcao("tabela", "Adiciona colunas calculadas a uma tabela.", "enriquece tabela virtual com métricas auxiliares.", "tabela"),
})


@dataclass
class InfoVisual:
    """Informações de um visual na página"""
    tipo: str
    titulo: Optional[str] = None
    nome: str = ""  # ID interno ou nome do container
    campos: List[str] = field(default_factory=list)


@dataclass
class InfoPagina:
    """Informações de uma página do relatório"""
    id: str
    nome_exibicao: str
    largura: int = 1280
    altura: int = 720
    opcao_exibicao: str = "FitToPage"
    tipo: str = "Normal"
    filtros: Optional[List[Dict]] = None
    visuais: List[InfoVisual] = field(default_factory=list)


# ============================================================================
# FUNÇÕES UTILITÁRIAS
# ============================================================================

def limpar_nome(nome: str) -> str:
    """Remove aspas simples do nome se existirem"""
    return nome.strip().strip("'\"")


def arquivo_existe(caminho: str) -> bool:
    """Verifica se arquivo existe"""
    return os.path.isfile(caminho)


def obter_com_fallback(dicionario: dict, chave: str, padrao=None):
    """Obtém valor do dicionário ou retorna padrão"""
    return dicionario.get(chave, padrao)


def _limitar_texto(texto: str, limite: int = 180) -> str:
    texto = re.sub(r"\s+", " ", str(texto or "")).strip()
    if len(texto) <= limite:
        return texto
    return texto[: limite - 3].rstrip() + "..."


def _normalizar_codigo_analise(codigo: str) -> str:
    return (
        str(codigo or "")
        .replace("#(lf)", "\n")
        .replace("#(cr)", "\r")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
        .strip()
    )


TERM_STOPWORDS = {
    "a", "o", "as", "os", "um", "uma", "uns", "umas", "de", "do", "da", "dos", "das",
    "e", "ou", "em", "no", "na", "nos", "nas", "por", "para", "com", "sem", "ao", "aos",
    "ate", "apos", "entre", "sobre", "sob", "que", "qual", "quais", "como", "mais",
    "menos", "se", "sua", "seu", "suas", "seus", "meu", "minha", "meus", "minhas",
    "the", "and", "or", "of", "to", "in", "on", "for", "from", "with", "without",
    "by", "at", "into", "is", "are", "was", "were", "be", "been", "this", "that",
    "these", "those", "not", "null", "blank", "true", "false",
    "novo", "nova", "novos", "novas",
}

TERM_TECHNICAL_WORDS = {
    "table", "tables", "column", "columns", "row", "rows", "record", "records",
    "value", "values", "literal", "source", "sourceRef", "expression", "query",
    "filter", "filters", "visual", "config", "object", "objects", "property",
    "properties", "select", "selected", "single", "localdatetable", "datetabletemplate",
    "power", "query", "dax", "m", "pbip", "tmdl", "json", "model", "report",
    "count", "sum", "avg", "min", "max", "format", "type", "name", "entity",
    "source", "kind", "item", "sheet", "database", "schema", "navigation",
    "sourceref", "queryref", "nativequeryref", "displayname", "singlevisual",
    "visualcontainer", "semantic", "expr",
    "date", "cd", "cod", "nr", "num", "vl", "dt", "nm", "ds", "tp", "fl", "st",
    "sk", "fk", "pk", "local", "etapa", "emp",
}

TERM_IDENTIFIER_WORDS = {"id", "ids", "codigo", "codigos", "code", "codes", "chave", "key", "keys", "guid"}

TERM_TIME_WORDS = {
    # PT
    "data", "ano", "mes", "mês", "dia", "semana", "periodo", "período",
    "trimestre", "competencia", "competência", "calendario", "calendário",
    # EN
    "date", "year", "month", "day", "week", "period", "quarter", "calendar",
    "hour", "today", "yesterday", "tomorrow",
    # Universal abbreviations
    "ytd", "mtd", "qtd", "ftd",
}

TERM_INDICATOR_WORDS = {
    # PT
    "receita", "despesa", "valor", "total", "saldo", "soma", "media", "média",
    "quantidade", "qtd", "percentual", "perc", "taxa", "indice", "índice",
    "meta", "previsto", "realizado", "orcamento", "orçamento", "custo", "margem",
    "lucro", "resultado", "venda", "vendas", "faturamento", "empenho", "empenhos",
    # EN
    "value", "revenue", "expense", "sales", "sale", "amount", "balance", "sum",
    "average", "avg", "count", "percent", "percentage", "rate", "ratio", "share",
    "index", "target", "planned", "actual", "budget", "cost", "margin", "profit",
    "income", "invoice", "expenditure", "paid", "payment", "credit", "debit",
    "price", "qty", "quantity", "fee", "spend",
}

TERM_ATTRIBUTE_WORDS = {
    # PT
    "status", "tipo", "categoria", "grupo", "classe", "segmento", "regiao", "região",
    "gerencia", "gerência", "diretoria", "departamento", "area", "área", "nome",
    "descricao", "descrição", "atributo", "natureza", "programa", "projeto", "processo",
    # EN
    "type", "category", "group", "class", "segment", "region", "name",
    "description", "attribute", "project", "process", "department",
}


def _sem_acentos(texto: str) -> str:
    return unicodedata.normalize("NFKD", str(texto or "")).encode("ascii", "ignore").decode("ascii")


def _normalizar_palavra_termo(palavra: str) -> str:
    palavra = _sem_acentos(palavra).lower().strip()
    palavra = re.sub(r"[^a-z0-9]+", "", palavra)
    if len(palavra) > 4 and palavra.endswith("s"):
        palavra = palavra[:-1]
    return palavra


def _texto_parece_tecnico_ou_valor(texto: str) -> bool:
    texto = str(texto or "").strip()
    if not texto:
        return True
    if len(texto) > 180:
        return True
    if re.search(r"https?://|www\.|[A-Za-z]:\\|\\\\|/[A-Za-z0-9_.-]+/", texto):
        return True
    if re.fullmatch(r"[{(]?[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}[)}]?", texto):
        return True
    if re.fullmatch(r"[\d\s.,:/\\_-]+", texto):
        return True
    return False


def _separar_palavras_termo(texto: str) -> List[Tuple[str, str]]:
    if _texto_parece_tecnico_ou_valor(texto):
        return []

    texto = str(texto or "")
    texto = re.sub(r"(?<=[a-zà-ÿ0-9])(?=[A-ZÀ-Ý])", " ", texto)
    texto = re.sub(r"[\[\]{}()'\"`]+", " ", texto)
    texto = re.sub(r"[_\-./\\|:;,+]+", " ", texto)

    palavras: List[Tuple[str, str]] = []
    for bruto in re.findall(r"[A-Za-zÀ-ÿ0-9]+", texto):
        norm = _normalizar_palavra_termo(bruto)
        if not norm:
            continue
        if norm in TERM_STOPWORDS:
            continue
        if norm in TERM_TECHNICAL_WORDS:
            continue
        if len(norm) <= 1:
            continue
        if re.fullmatch(r"\d+", norm):
            continue
        palavras.append((bruto.strip(), norm))

    return palavras


def _termo_display(palavras: List[str]) -> str:
    partes = []
    for palavra in palavras:
        palavra = str(palavra or "").strip()
        if not palavra:
            continue
        if palavra.isupper() and len(palavra) <= 6:
            partes.append(palavra)
        elif palavra.lower() in {"id", "ids"}:
            partes.append(palavra.upper())
        else:
            partes.append(palavra[:1].upper() + palavra[1:])
    return " ".join(partes)


# Traducoes pt-BR para tipos de visuais do PBIR (mapping nativo -> rotulo amigavel).
# Cobre os visuais padrao do Power BI + nomes mais comuns.
VISUAL_TYPE_TRANSLATIONS: Dict[str, str] = {
    "barChart": "Gráfico de Barras",
    "clusteredBarChart": "Gráfico de Barras Agrupadas",
    "stackedBarChart": "Gráfico de Barras Empilhadas",
    "hundredPercentStackedBarChart": "Gráfico de Barras Empilhadas 100%",
    "columnChart": "Gráfico de Colunas",
    "clusteredColumnChart": "Gráfico de Colunas Agrupadas",
    "stackedColumnChart": "Gráfico de Colunas Empilhadas",
    "hundredPercentStackedColumnChart": "Gráfico de Colunas Empilhadas 100%",
    "lineChart": "Gráfico de Linhas",
    "areaChart": "Gráfico de Áreas",
    "stackedAreaChart": "Gráfico de Áreas Empilhadas",
    "pieChart": "Gráfico de Pizza",
    "donutChart": "Gráfico de Rosca",
    "scatterChart": "Gráfico de Dispersão",
    "treemap": "Treemap",
    "card": "Cartão",
    "cardVisual": "Cartão",
    "multiRowCard": "Cartão Multilinhas",
    "kpi": "KPI",
    "gauge": "Medidor",
    "slicer": "Segmentação de Dados",
    "advancedSlicerVisual": "Segmentação Avançada",
    "tableEx": "Tabela",
    "pivotTable": "Matriz",
    "matrix": "Matriz",
    "map": "Mapa",
    "filledMap": "Mapa Preenchido",
    "shapeMap": "Mapa de Formas",
    "azureMap": "Mapa do Azure",
    "textbox": "Caixa de Texto",
    "image": "Imagem",
    "shape": "Forma",
    "basicShape": "Forma Básica",
    "actionButton": "Botão",
    "lineStackedColumnComboChart": "Combo Linha + Coluna Empilhada",
    "lineClusteredColumnComboChart": "Combo Linha + Coluna Agrupada",
    "waterfallChart": "Gráfico de Cascata",
    "funnel": "Funil",
    "ribbonChart": "Gráfico de Faixas",
    "decompositionTreeVisual": "Árvore de Decomposição",
    "qnaVisual": "Perguntas e Respostas",
    "keyDriversVisual": "Influenciadores-Chave",
    "smartNarrativeVisual": "Narrativa Inteligente",
    "Group": "Grupo",
    "Unknown": "Tipo desconhecido",
}


# Traducoes para valores de enum do PBIR (tipos de pagina, filtros, etc.).
PBIR_ENUM_TRANSLATIONS: Dict[str, str] = {
    "Drillthrough": "Detalhamento",
    "Normal": "Padrão",
    "Tooltip": "Dica de Ferramenta",
    "Categorical": "Categórico",
    "Exclusion": "Exclusão",
    "Numeric": "Numérico",
    "TopN": "Top N",
    "Advanced": "Avançado",
    "Tuple": "Tupla",
    "RelativeDate": "Data Relativa",
    "RelativeTime": "Tempo Relativo",
}


def traduzir_tipo_visual(tipo: str) -> str:
    """Converte o id de tipo de visual do PBIR para nome amigavel em pt-BR.

    Custom visuals frequentemente trazem hash GUID no fim, ex:
    `payPalKPIDonutChart55A431AB15A540ED924ACD72ED8D259F`. Esse hash e
    removido antes da traducao.
    """
    if not tipo or tipo == "-":
        return tipo
    # Remove sufixo hexadecimal de 20+ chars (GUID de custom visual).
    limpo = re.sub(r"[A-F0-9]{20,}$", "", tipo)
    if limpo in VISUAL_TYPE_TRANSLATIONS:
        return VISUAL_TYPE_TRANSLATIONS[limpo]
    # Tipo nao mapeado: tenta separar camelCase em palavras Title-Case.
    palavras = re.findall(r"[A-Z][a-z]+|[A-Z]+(?=[A-Z][a-z])|[a-z]+|[A-Z]+|\d+", limpo)
    if palavras:
        return " ".join(p.title() if p.islower() else p for p in palavras)
    return limpo


def traduzir_enum_pbir(valor: str) -> str:
    """Traduz valor de enum do PBIR (Drillthrough, Categorical, etc.) para pt-BR."""
    if not valor:
        return valor
    return PBIR_ENUM_TRANSLATIONS.get(valor, valor)


def _extrair_termos_texto(texto: str, incluir_frases: bool = True) -> List[Tuple[str, str]]:
    palavras = _separar_palavras_termo(texto)
    if not palavras:
        return []

    termos: List[Tuple[str, str]] = []
    vistos = set()

    def adicionar(slice_palavras: List[Tuple[str, str]], permitir_identificador_solo: bool = False):
        norm_words = [p[1] for p in slice_palavras]
        if not norm_words:
            return
        if not permitir_identificador_solo and len(norm_words) == 1 and norm_words[0] in TERM_IDENTIFIER_WORDS:
            return
        chave = " ".join(norm_words)
        if chave in vistos:
            return
        vistos.add(chave)
        termos.append((_termo_display([p[0] for p in slice_palavras]), chave))

    for palavra in palavras:
        adicionar([palavra])

    if incluir_frases:
        for tamanho in (2, 3):
            if len(palavras) < tamanho:
                continue
            for idx in range(0, len(palavras) - tamanho + 1):
                fatia = palavras[idx: idx + tamanho]
                if all(p[1] in TERM_IDENTIFIER_WORDS for p in fatia):
                    continue
                adicionar(fatia, permitir_identificador_solo=True)

    return termos[:24]


def _extrair_referencias_dax(expressao: str) -> List[str]:
    referencias: List[str] = []
    texto = str(expressao or "")
    for tabela in re.findall(r"'([^']+)'", texto):
        if tabela and not _texto_parece_tecnico_ou_valor(tabela):
            referencias.append(tabela)
    for coluna in re.findall(r"\[([^\]]+)\]", texto):
        if coluna and not _texto_parece_tecnico_ou_valor(coluna):
            referencias.append(coluna)
    return list(dict.fromkeys(referencias))


def _codigo_fonte_eh_dax(codigo: str) -> bool:
    codigo_norm = _normalizar_codigo_analise(codigo)
    if not codigo_norm:
        return False

    cabecalho = codigo_norm[:300].upper()
    if "DAX" in cabecalho[:80]:
        return True
    if re.match(r"^\s*(DEFINE|EVALUATE|VAR|RETURN)\b", codigo_norm, re.I):
        return True
    if not re.search(r"(?im)^\s*let\b", codigo_norm):
        funcoes_dax = _detectar_funcoes_catalogo(codigo_norm, DAX_FUNCTION_CATALOG)
        funcoes_m = _detectar_funcoes_catalogo(codigo_norm, POWER_QUERY_FUNCTION_CATALOG)
        return len(funcoes_dax) > 0 and len(funcoes_m) == 0
    return False


def _extrair_comentarios_bi_doc(codigo: str) -> Dict[str, List[str]]:
    mapa = {
        "BI_DOC": "geral",
        "BI_DOC_ORIGEM": "origem",
        "BI_DOC_REGRA": "regra",
        "BI_DOC_OBSERVACAO": "observacao",
    }
    comentarios = {valor: [] for valor in mapa.values()}
    padrao = re.compile(r"^\s*//\s*(BI_DOC(?:_ORIGEM|_REGRA|_OBSERVACAO)?)\s*:\s*(.+?)\s*$", re.I)

    for linha in _normalizar_codigo_analise(codigo).splitlines():
        match = padrao.match(linha)
        if match:
            chave = mapa[match.group(1).upper()]
            comentarios[chave].append(match.group(2).strip())

    return {chave: valores for chave, valores in comentarios.items() if valores}


def _remover_linhas_comentario(codigo: str) -> str:
    return "\n".join(
        linha for linha in codigo.splitlines()
        if not linha.strip().startswith("//")
    )


def _split_top_level(texto: str, separador: str = ",") -> List[str]:
    partes: List[str] = []
    atual: List[str] = []
    profundidade = 0
    em_string = False
    i = 0

    while i < len(texto):
        char = texto[i]
        prox = texto[i + 1] if i + 1 < len(texto) else ""

        if em_string:
            atual.append(char)
            if char == '"' and prox == '"':
                atual.append(prox)
                i += 2
                continue
            if char == '"':
                em_string = False
            i += 1
            continue

        if char == '"':
            em_string = True
            atual.append(char)
        elif char in "([{":
            profundidade += 1
            atual.append(char)
        elif char in ")]}":
            profundidade = max(0, profundidade - 1)
            atual.append(char)
        elif char == separador and profundidade == 0:
            parte = "".join(atual).strip()
            if parte:
                partes.append(parte)
            atual = []
        else:
            atual.append(char)

        i += 1

    parte = "".join(atual).strip()
    if parte:
        partes.append(parte)
    return partes


def _split_atribuicao_m(texto: str) -> Tuple[str, str]:
    profundidade = 0
    em_string = False
    i = 0
    while i < len(texto):
        char = texto[i]
        prox = texto[i + 1] if i + 1 < len(texto) else ""
        if em_string:
            if char == '"' and prox == '"':
                i += 2
                continue
            if char == '"':
                em_string = False
            i += 1
            continue
        if char == '"':
            em_string = True
        elif char in "([{":
            profundidade += 1
        elif char in ")]}":
            profundidade = max(0, profundidade - 1)
        elif char == "=" and profundidade == 0:
            return texto[:i].strip(), texto[i + 1:].strip()
        i += 1
    return texto.strip(), ""


def _limpar_identificador_m(nome: str) -> str:
    nome = re.sub(r"^\s*shared\s+", "", nome.strip(), flags=re.I)
    if nome.startswith('#"') and nome.endswith('"'):
        return nome[2:-1].replace('""', '"')
    return limpar_nome(nome)


def _extrair_strings_m(expressao: str) -> List[str]:
    valores = []
    i = 0
    while i < len(expressao):
        if expressao[i] == "#" and i + 1 < len(expressao) and expressao[i + 1] == '"':
            i += 2
            while i < len(expressao):
                if expressao[i] == '"' and i + 1 < len(expressao) and expressao[i + 1] == '"':
                    i += 2
                    continue
                if expressao[i] == '"':
                    i += 1
                    break
                i += 1
            continue

        if expressao[i] == '"':
            i += 1
            chars = []
            while i < len(expressao):
                if expressao[i] == '"' and i + 1 < len(expressao) and expressao[i + 1] == '"':
                    chars.append('"')
                    i += 2
                    continue
                if expressao[i] == '"':
                    i += 1
                    break
                chars.append(expressao[i])
                i += 1
            valor = "".join(chars).strip()
            if valor:
                valores.append(valor)
            continue

        i += 1
    return valores


def _resumir_lista(valores: List[str], limite: int = 6) -> str:
    unicos = list(dict.fromkeys(v for v in valores if v))
    if not unicos:
        return ""
    texto = ", ".join(unicos[:limite])
    if len(unicos) > limite:
        texto += f" e mais {len(unicos) - limite}"
    return texto


def _detectar_funcoes_catalogo(expressao: str, catalogo: Dict[str, CatalogoFuncao]) -> List[CatalogoFuncao]:
    texto = str(expressao or "")
    encontrados: List[CatalogoFuncao] = []

    for nome, item in catalogo.items():
        if nome == "try otherwise":
            if re.search(r"\btry\b", texto, re.I) and re.search(r"\botherwise\b", texto, re.I):
                encontrados.append(item)
            continue
        if nome == "if then else":
            if (
                re.search(r"\bif\b", texto, re.I)
                and re.search(r"\bthen\b", texto, re.I)
                and re.search(r"\belse\b", texto, re.I)
            ):
                encontrados.append(item)
            continue

        padrao = rf"(?<![A-Za-z0-9_]){re.escape(nome)}\s*\("
        if re.search(padrao, texto, re.I):
            encontrados.append(item)

    return encontrados


def _resumo_funcoes_catalogo(funcoes: List[CatalogoFuncao], principal: str = "", limite: int = 3) -> str:
    itens = []
    principal_lower = principal.lower()
    vistos = set()
    for funcao in funcoes:
        if funcao.nome.lower() == principal_lower or funcao.nome.lower() in vistos:
            continue
        vistos.add(funcao.nome.lower())
        itens.append(f"{funcao.nome}: {funcao.leitura_negocio.rstrip('.')}")

    if not itens:
        return ""

    texto = "; ".join(itens[:limite])
    if len(itens) > limite:
        texto += f"; e mais {len(itens) - limite} função(ões)"
    return f" Funções auxiliares detectadas: {texto}."


def _descricao_catalogo_generica(nome_etapa: str, funcao: CatalogoFuncao, funcoes: List[CatalogoFuncao]) -> str:
    descricao = f"Etapa `{nome_etapa}` {funcao.leitura_negocio}"
    if not descricao.endswith("."):
        descricao += "."
    return descricao + _resumo_funcoes_catalogo(funcoes, funcao.nome)


def _adicionar_unico(lista: List[str], valor: str) -> None:
    if valor and valor not in lista:
        lista.append(valor)


def _registrar_funcoes_power_query(regra: RegraPowerQuery, expressao: str) -> None:
    for funcao in _detectar_funcoes_catalogo(expressao, POWER_QUERY_FUNCTION_CATALOG):
        rotulo = POWER_QUERY_CATEGORY_LABELS.get(funcao.categoria, funcao.categoria.title())
        texto = f"`{funcao.nome}`: {funcao.descricao}"
        regra.funcoes_relevantes.setdefault(rotulo, [])
        _adicionar_unico(regra.funcoes_relevantes[rotulo], texto)


def _token_codigo(texto: str, limite: int = 90) -> str:
    texto = _limitar_texto(texto, limite).replace("\n", " ").strip()
    return f"`{texto}`" if texto else ""


def _nomes_funcoes_power_query(expressao: str) -> List[str]:
    return [funcao.nome for funcao in _detectar_funcoes_catalogo(expressao, POWER_QUERY_FUNCTION_CATALOG)]


def _principal_funcao_power_query(expressao: str, padrao: str = "") -> str:
    funcoes = _nomes_funcoes_power_query(expressao)
    return funcoes[0] if funcoes else padrao


def _descricao_executiva_power_query(texto: str) -> str:
    texto = re.sub(r"\s+Funções auxiliares detectadas:.*$", ".", str(texto or "")).strip()
    texto = re.sub(r"\.{2,}$", ".", texto)
    return texto


def _linha_regra_power_query(etapa: EtapaPowerQuery) -> LinhaRegraPowerQuery:
    expr_lower = etapa.expressao.lower()
    strings = _extrair_strings_m(etapa.expressao)
    colunas = _resumir_lista(strings)
    funcao_principal = _principal_funcao_power_query(etapa.expressao)
    funcao_token = _token_codigo(funcao_principal, 60)
    descricao_etapa = _descricao_executiva_power_query(etapa.descricao)

    if etapa.categoria == "origem":
        if re.search(r"\{\s*\[.+?\]\s*\}\s*\[", etapa.expressao, re.S):
            return LinhaRegraPowerQuery(
                "Navegação",
                "Objeto da origem",
                f"Etapa `{etapa.nome}` navega para um objeto específico dentro da origem de dados.",
            )
        regra = funcao_token or "Fonte de dados"
        return LinhaRegraPowerQuery(
            "Fonte",
            regra,
            descricao_etapa,
        )

    if "table.selectrows" in expr_lower:
        condicao = _extrair_condicao_each(etapa.expressao)
        regra = _token_codigo(condicao, 140) if condicao else funcao_token or "Filtro"
        descricao = (
            f"Etapa `{etapa.nome}` mantém apenas os registros que atendem à condição informada."
            if condicao else descricao_etapa
        )
        return LinhaRegraPowerQuery("Filtro", regra, descricao)

    if "table.selectcolumns" in expr_lower:
        descricao = f"Etapa `{etapa.nome}` mantém somente as colunas: {colunas}." if colunas else etapa.descricao
        return LinhaRegraPowerQuery("Seleção de Colunas", funcao_token or "Colunas selecionadas", descricao)

    if "table.removecolumns" in expr_lower:
        descricao = f"Etapa `{etapa.nome}` remove as colunas: {colunas}." if colunas else etapa.descricao
        return LinhaRegraPowerQuery("Remoção de Colunas", funcao_token or "Colunas removidas", descricao)

    if "table.renamecolumns" in expr_lower:
        pares = [f"{strings[i]} -> {strings[i + 1]}" for i in range(0, len(strings) - 1, 2)]
        regra = _token_codigo(_resumir_lista(pares, 4), 120) if pares else funcao_token
        descricao = f"Etapa `{etapa.nome}` padroniza nomes de colunas." if pares else etapa.descricao
        return LinhaRegraPowerQuery("Renomeação", regra or "Renomeação de colunas", descricao)

    if "table.transformcolumntypes" in expr_lower:
        descricao = f"Etapa `{etapa.nome}` ajusta os tipos de dados das colunas: {colunas}." if colunas else etapa.descricao
        return LinhaRegraPowerQuery("Tipagem", funcao_token or "Conversão de tipos", descricao)

    if "table.addcolumn" in expr_lower:
        nome_coluna = strings[0] if strings else ""
        regra = f"Coluna calculada: {_token_codigo(nome_coluna, 70)}" if nome_coluna else funcao_token
        descricao = f"Etapa `{etapa.nome}` cria uma coluna calculada a partir de regra definida em M."
        return LinhaRegraPowerQuery("Coluna Calculada", regra or "Nova coluna", descricao)

    if "table.nestedjoin" in expr_lower or "table.join" in expr_lower:
        descricao = f"Etapa `{etapa.nome}` combina dados com outra consulta/tabela."
        if colunas:
            descricao += f" Chaves ou campos citados: {colunas}."
        return LinhaRegraPowerQuery("Merge / Junção", funcao_token or "Junção entre consultas", descricao)

    if "table.combine" in expr_lower:
        return LinhaRegraPowerQuery("Append", funcao_token or "Anexação de consultas", descricao_etapa)

    if (
        "table.expandtablecolumn" in expr_lower
        or "table.expandrecordcolumn" in expr_lower
        or "table.expandlistcolumn" in expr_lower
    ):
        descricao = f"Etapa `{etapa.nome}` expande campos relacionados."
        if colunas:
            descricao += f" Campos expandidos: {colunas}."
        return LinhaRegraPowerQuery("Expansão", funcao_token or "Campos expandidos", descricao)

    if "table.group" in expr_lower:
        descricao = f"Etapa `{etapa.nome}` resume registros por chaves de análise."
        if colunas:
            descricao += f" Campos citados: {colunas}."
        return LinhaRegraPowerQuery("Agrupamento", funcao_token or "Agrupamento e agregação", descricao)

    if "table.pivot" in expr_lower:
        return LinhaRegraPowerQuery("Pivot", funcao_token or "Linhas para colunas", descricao_etapa)

    if "table.unpivot" in expr_lower:
        return LinhaRegraPowerQuery("Unpivot", funcao_token or "Colunas para linhas", descricao_etapa)

    if "table.sort" in expr_lower:
        descricao = f"Etapa `{etapa.nome}` ordena os registros."
        if colunas:
            descricao += f" Campos de ordenação: {colunas}."
        return LinhaRegraPowerQuery("Ordenação", funcao_token or "Ordenação de linhas", descricao)

    if "table.distinct" in expr_lower or "table.removeduplicates" in expr_lower:
        descricao = f"Etapa `{etapa.nome}` remove registros duplicados."
        if colunas:
            descricao += f" Campos considerados: {colunas}."
        return LinhaRegraPowerQuery("Deduplicação", funcao_token or "Remoção de duplicidades", descricao)

    if "table.replacevalue" in expr_lower:
        regra = _token_codigo(" -> ".join(strings[:2]), 90) if len(strings) >= 2 else funcao_token
        return LinhaRegraPowerQuery("Substituição de Valores", regra or "Substituição", descricao_etapa)

    if "table.promoteheaders" in expr_lower:
        return LinhaRegraPowerQuery("Cabeçalhos", funcao_token or "Primeira linha como cabeçalho", descricao_etapa)

    if (
        "table.removerowswitherrors" in expr_lower
        or "table.selectrowswitherrors" in expr_lower
        or "table.replaceerrorvalues" in expr_lower
    ):
        return LinhaRegraPowerQuery("Tratamento de Erros", funcao_token or "Tratamento de erros", descricao_etapa)

    if "table.skip" in expr_lower or "table.removerows" in expr_lower or "table.firstn" in expr_lower or "table.lastn" in expr_lower:
        return LinhaRegraPowerQuery("Recorte de Linhas", funcao_token or "Limite/remoção por posição", descricao_etapa)

    if "try" in expr_lower and "otherwise" in expr_lower:
        return LinhaRegraPowerQuery("Tratamento de Erro", _token_codigo("try ... otherwise"), descricao_etapa)

    if re.search(r"\bif\b", etapa.expressao, re.I) and re.search(r"\bthen\b", etapa.expressao, re.I):
        return LinhaRegraPowerQuery("Regra Condicional", _token_codigo("if ... then ... else"), descricao_etapa)

    if etapa.categoria == "personalizada":
        return LinhaRegraPowerQuery(
            "Transformação Personalizada",
            "Transformação personalizada",
            f"Etapa `{etapa.nome}` não foi classificada automaticamente. Consulte o código M original para detalhe completo.",
        )

    return LinhaRegraPowerQuery(
        POWER_QUERY_CATEGORY_LABELS.get(etapa.categoria, etapa.categoria.title()),
        funcao_token or "Transformação",
        descricao_etapa,
    )


def _extrair_condicao_each(expressao: str) -> str:
    match = re.search(r"\beach\b(.+)", expressao, re.I | re.S)
    if not match:
        return ""
    condicao = match.group(1).strip()
    condicao = re.sub(r"\)\s*$", "", condicao).strip()
    return _limitar_texto(condicao, 220)


def _detectar_origem_power_query(expressao: str) -> str:
    origens = [
        ("Sql.Database", "SQL Server"),
        ("Oracle.Database", "Oracle"),
        ("PostgreSQL.Database", "PostgreSQL"),
        ("MySQL.Database", "MySQL"),
        ("Snowflake.Databases", "Snowflake"),
        ("AnalysisServices.Database", "Analysis Services"),
        ("Odbc.DataSource", "ODBC"),
        ("OleDb.DataSource", "OLE DB"),
        ("Excel.Workbook", "arquivo Excel"),
        ("Csv.Document", "arquivo CSV"),
        ("Json.Document", "arquivo JSON"),
        ("Web.Contents", "Web/API"),
        ("OData.Feed", "OData"),
        ("SharePoint.Files", "SharePoint"),
        ("SharePoint.Contents", "SharePoint"),
        ("Folder.Files", "pasta de arquivos"),
        ("File.Contents", "arquivo local"),
        ("PowerPlatform.Dataflows", "Dataflow Power Platform"),
        ("Dataverse.Contents", "Dataverse"),
        ("Value.NativeQuery", "consulta nativa na fonte de dados"),
    ]
    expr_lower = expressao.lower()
    for funcao, descricao in origens:
        if funcao.lower() in expr_lower:
            return descricao
    return ""


def _classificar_etapa_power_query(nome: str, expressao: str) -> EtapaPowerQuery:
    expr_lower = expressao.lower()
    strings = _extrair_strings_m(expressao)
    colunas = _resumir_lista(strings)
    origem = _detectar_origem_power_query(expressao)
    funcoes_catalogo = _detectar_funcoes_catalogo(expressao, POWER_QUERY_FUNCTION_CATALOG)

    if origem:
        principal = funcoes_catalogo[0].nome if funcoes_catalogo else ""
        return EtapaPowerQuery(
            nome,
            expressao,
            "origem",
            f"Etapa `{nome}` conecta ou lê dados de {origem}."
            + _resumo_funcoes_catalogo(funcoes_catalogo, principal),
        )
    if re.search(r"\{\s*\[.+?\]\s*\}\s*\[", expressao, re.S):
        return EtapaPowerQuery(nome, expressao, "origem", f"Etapa `{nome}` navega para um objeto específico dentro da origem de dados.")
    if "table.selectrows" in expr_lower:
        condicao = _extrair_condicao_each(expressao)
        detalhe = f" usando a condição `{condicao}`" if condicao else ""
        return EtapaPowerQuery(nome, expressao, "filtro", f"Etapa `{nome}` filtra linhas{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.SelectRows"))
    if "table.nestedjoin" in expr_lower or "table.join" in expr_lower:
        detalhe = f" Colunas/chaves citadas: {colunas}." if colunas else ""
        principal = "Table.NestedJoin" if "table.nestedjoin" in expr_lower else "Table.Join"
        return EtapaPowerQuery(nome, expressao, "integracao", f"Etapa `{nome}` combina dados com outra consulta ou tabela.{detalhe}" + _resumo_funcoes_catalogo(funcoes_catalogo, principal))
    if "table.combine" in expr_lower:
        return EtapaPowerQuery(nome, expressao, "integracao", f"Etapa `{nome}` empilha/anexa dados de múltiplas consultas ou tabelas." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.Combine"))
    if "table.expandtablecolumn" in expr_lower or "table.expandrecordcolumn" in expr_lower:
        detalhe = f" Campos expandidos: {colunas}." if colunas else ""
        principal = "Table.ExpandTableColumn" if "table.expandtablecolumn" in expr_lower else "Table.ExpandRecordColumn"
        return EtapaPowerQuery(nome, expressao, "integracao", f"Etapa `{nome}` expande colunas vindas de consulta relacionada.{detalhe}" + _resumo_funcoes_catalogo(funcoes_catalogo, principal))
    if "table.removecolumns" in expr_lower:
        detalhe = f": {colunas}" if colunas else ""
        return EtapaPowerQuery(nome, expressao, "transformacao", f"Etapa `{nome}` remove colunas{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.RemoveColumns"))
    if "table.selectcolumns" in expr_lower:
        detalhe = f": {colunas}" if colunas else ""
        return EtapaPowerQuery(nome, expressao, "transformacao", f"Etapa `{nome}` mantém somente colunas selecionadas{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.SelectColumns"))
    if "table.renamecolumns" in expr_lower:
        pares = [f"{strings[i]} -> {strings[i + 1]}" for i in range(0, len(strings) - 1, 2)]
        detalhe = f": {_resumir_lista(pares, 5)}" if pares else ""
        return EtapaPowerQuery(nome, expressao, "transformacao", f"Etapa `{nome}` renomeia colunas{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.RenameColumns"))
    if "table.transformcolumntypes" in expr_lower:
        detalhe = f": {colunas}" if colunas else ""
        return EtapaPowerQuery(nome, expressao, "transformacao", f"Etapa `{nome}` ajusta tipos de dados das colunas{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.TransformColumnTypes"))
    if "table.addcolumn" in expr_lower:
        detalhe = f" `{strings[0]}`" if strings else ""
        return EtapaPowerQuery(nome, expressao, "transformacao", f"Etapa `{nome}` cria coluna calculada{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.AddColumn"))
    if "table.group" in expr_lower:
        detalhe = f" por {colunas}" if colunas else ""
        return EtapaPowerQuery(nome, expressao, "agregacao", f"Etapa `{nome}` agrupa registros{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.Group"))
    if "table.sort" in expr_lower:
        detalhe = f": {colunas}" if colunas else ""
        return EtapaPowerQuery(nome, expressao, "transformacao", f"Etapa `{nome}` ordena linhas{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.Sort"))
    if "table.removerowswitherrors" in expr_lower or "table.selectrowswitherrors" in expr_lower:
        principal = "Table.RemoveRowsWithErrors" if "table.removerowswitherrors" in expr_lower else "Table.SelectRowsWithErrors"
        return EtapaPowerQuery(nome, expressao, "qualidade", _descricao_catalogo_generica(nome, POWER_QUERY_FUNCTION_CATALOG[principal], funcoes_catalogo))
    if "table.distinct" in expr_lower or "table.removeduplicates" in expr_lower:
        detalhe = f" considerando {colunas}" if colunas else ""
        principal = "Table.Distinct" if "table.distinct" in expr_lower else "Table.RemoveDuplicates"
        return EtapaPowerQuery(nome, expressao, "qualidade", f"Etapa `{nome}` remove duplicidades{detalhe}." + _resumo_funcoes_catalogo(funcoes_catalogo, principal))
    if "table.replacevalue" in expr_lower:
        return EtapaPowerQuery(nome, expressao, "qualidade", f"Etapa `{nome}` substitui valores em colunas da tabela." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.ReplaceValue"))
    if "table.unpivot" in expr_lower:
        principal = "Table.UnpivotOtherColumns" if "table.unpivotothercolumns" in expr_lower else "Table.Unpivot"
        return EtapaPowerQuery(nome, expressao, "agregacao", f"Etapa `{nome}` transforma colunas em linhas (unpivot)." + _resumo_funcoes_catalogo(funcoes_catalogo, principal))
    if "table.pivot" in expr_lower:
        return EtapaPowerQuery(nome, expressao, "agregacao", f"Etapa `{nome}` transforma valores de linhas em colunas (pivot)." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.Pivot"))
    if "table.promoteheaders" in expr_lower:
        return EtapaPowerQuery(nome, expressao, "transformacao", f"Etapa `{nome}` promove a primeira linha para cabeçalhos." + _resumo_funcoes_catalogo(funcoes_catalogo, "Table.PromoteHeaders"))
    if "table.skip" in expr_lower or "table.removerows" in expr_lower or "table.firstn" in expr_lower:
        principal = "Table.Skip" if "table.skip" in expr_lower else ("Table.FirstN" if "table.firstn" in expr_lower else "Table.RemoveRows")
        return EtapaPowerQuery(nome, expressao, "filtro", f"Etapa `{nome}` limita ou remove linhas por posição/quantidade." + _resumo_funcoes_catalogo(funcoes_catalogo, principal))

    if funcoes_catalogo:
        funcao = funcoes_catalogo[0]
        return EtapaPowerQuery(nome, expressao, funcao.categoria, _descricao_catalogo_generica(nome, funcao, funcoes_catalogo))

    return EtapaPowerQuery(nome, expressao, "personalizada", f"Etapa `{nome}` executa transformação personalizada/não classificada.")


def analisar_power_query_m(codigo: str) -> RegraPowerQuery:
    codigo = _normalizar_codigo_analise(codigo)
    regra = RegraPowerQuery(observacoes=_extrair_comentarios_bi_doc(codigo))
    codigo_sem_comentarios = _remover_linhas_comentario(codigo)
    match_let = re.search(r"(?im)^\s*let\b", codigo_sem_comentarios)

    if not match_let:
        origem = _detectar_origem_power_query(codigo_sem_comentarios)
        _registrar_funcoes_power_query(regra, codigo_sem_comentarios)
        if codigo_sem_comentarios.strip():
            etapa_unica = _classificar_etapa_power_query("Consulta", codigo_sem_comentarios)
            regra.etapas.append(etapa_unica)
            regra.linhas_regra.append(_linha_regra_power_query(etapa_unica))
        if origem:
            regra.origem_dados.append(f"A consulta lê dados de {origem}.")
        elif codigo_sem_comentarios.strip():
            funcoes = _detectar_funcoes_catalogo(codigo_sem_comentarios, POWER_QUERY_FUNCTION_CATALOG)
            if funcoes:
                regra.transformacoes.append(_descricao_catalogo_generica("Consulta", funcoes[0], funcoes))
            else:
                regra.transformacoes.append("Consulta em formato personalizado/não classificado.")
        regra.resultado_final = "Resultado final inferido a partir de uma expressão Power Query sem bloco `let` explícito."
        return regra

    apos_let = codigo_sem_comentarios[match_let.end():]
    linhas = apos_let.splitlines()
    idx_in = None
    final_expr = ""
    for idx in range(len(linhas) - 1, -1, -1):
        linha = linhas[idx].strip()
        if re.match(r"^in(\s+.+)?$", linha, re.I):
            idx_in = idx
            final_expr = re.sub(r"^in\b", "", linha, flags=re.I).strip()
            if not final_expr and idx + 1 < len(linhas):
                final_expr = "\n".join(linhas[idx + 1:]).strip()
            break

    texto_etapas = "\n".join(linhas[:idx_in]) if idx_in is not None else apos_let
    regra.etapa_final = _limpar_identificador_m(final_expr) if final_expr else ""

    for trecho in _split_top_level(texto_etapas):
        nome, expressao = _split_atribuicao_m(trecho)
        if not nome or not expressao:
            continue
        etapa = _classificar_etapa_power_query(_limpar_identificador_m(nome), expressao)
        _registrar_funcoes_power_query(regra, expressao)
        regra.etapas.append(etapa)
        regra.linhas_regra.append(_linha_regra_power_query(etapa))
        if etapa.categoria == "origem":
            regra.origem_dados.append(etapa.descricao)
        elif etapa.categoria == "filtro":
            regra.filtros_aplicados.append(etapa.descricao)
        elif etapa.categoria == "integracao":
            regra.integracoes.append(etapa.descricao)
        else:
            regra.transformacoes.append(etapa.descricao)

    final = regra.etapa_final or (regra.etapas[-1].nome if regra.etapas else "")
    regra.resultado_final = (
        f"A consulta retorna a etapa `{final}` como resultado final."
        if final else
        "Resultado final não identificado no bloco `in`."
    )
    return regra


def analisar_dax(expressao: str) -> LeituraDax:
    expressao = str(expressao or "").strip()
    leitura = LeituraDax()
    if not expressao:
        return leitura

    for funcao in _detectar_funcoes_catalogo(expressao, DAX_FUNCTION_CATALOG):
        if funcao.nome in leitura.funcoes:
            continue
        leitura.funcoes.append(funcao.nome)
        item = f"`{funcao.nome}`: {funcao.descricao} Leitura de negócio: {funcao.leitura_negocio}"
        leitura.itens.append(item)
        leitura.categorias.setdefault(funcao.categoria, []).append(item)

    if re.search(r"\bSWITCH\s*\(\s*TRUE\s*\(", expressao, re.I):
        item = "`SWITCH(TRUE())`: padrão comum para regras condicionais em cascata, semelhante a múltiplos cenários de negócio."
        if item not in leitura.itens:
            leitura.itens.append(item)
            leitura.categorias.setdefault("logica", []).append(item)

    # Coleta nomes de VAR para nao tratar como medida referenciada
    for m in re.finditer(r"\bVAR\s+([A-Za-z_][A-Za-z0-9_]*)\s*=", expressao, re.I):
        nome_var = m.group(1)
        if nome_var not in leitura.variaveis_dax:
            leitura.variaveis_dax.append(nome_var)

    # Referencias a tabelas: Tabela[Coluna] ou 'Tabela com espacos'[Coluna]
    for m in re.finditer(r"(?:'([^']+)'|([A-Za-z_][A-Za-z0-9_]*))\s*\[([^\]\[]+)\]", expressao):
        tabela_nome = m.group(1) or m.group(2)
        if tabela_nome and tabela_nome not in leitura.referencias_tabelas:
            leitura.referencias_tabelas.append(tabela_nome)

    # Referencias a medidas: [Nome] SEM prefixo de tabela (sem ' ou identificador antes)
    for m in re.finditer(r"(?<![\w'\]])\[([^\]\[]+)\]", expressao):
        nome = m.group(1).strip()
        if not nome:
            continue
        if nome in leitura.variaveis_dax:
            continue
        if nome not in leitura.referencias_medidas:
            leitura.referencias_medidas.append(nome)

    return leitura


def _agrupar_linhas_regra_consecutivas(
    linhas: List[LinhaRegraPowerQuery], min_grupo: int = 3
) -> List[LinhaRegraPowerQuery]:
    """Agrupa N etapas consecutivas com mesma (etapa, regra) numa unica linha.

    Evita poluir a tabela com 25x a mesma linha 'Text.Replace' quando a query
    aplica varias substituicoes em sequencia. A descricao da linha agrupada
    cita o intervalo de nomes (`t1` ... `t25`) e indica o total.
    """
    if len(linhas) < min_grupo:
        return linhas

    resultado: List[LinhaRegraPowerQuery] = []
    i = 0
    while i < len(linhas):
        atual = linhas[i]
        j = i + 1
        while (
            j < len(linhas)
            and linhas[j].etapa == atual.etapa
            and linhas[j].regra == atual.regra
        ):
            j += 1

        count = j - i
        if count >= min_grupo:
            nomes_etapas: List[str] = []
            for k in range(i, j):
                m = re.search(r"`([^`]+)`", linhas[k].descricao)
                if m:
                    nomes_etapas.append(m.group(1))

            if not nomes_etapas:
                nova_desc = (
                    f"Aplicado {count} vezes consecutivas. "
                    "Detalhes no codigo fonte abaixo."
                )
            elif count <= 4:
                nomes_str = ", ".join(f"`{n}`" for n in nomes_etapas)
                nova_desc = (
                    f"Aplicado {count} vezes consecutivas em: {nomes_str}. "
                    "Detalhes no codigo fonte abaixo."
                )
            else:
                primeira = nomes_etapas[0]
                ultima = nomes_etapas[-1]
                nova_desc = (
                    f"Aplicado {count} vezes consecutivas "
                    f"(`{primeira}` ... `{ultima}`). "
                    "Detalhes no codigo fonte abaixo."
                )

            resultado.append(
                LinhaRegraPowerQuery(
                    etapa=atual.etapa,
                    regra=atual.regra,
                    descricao=nova_desc,
                )
            )
        else:
            for k in range(i, j):
                resultado.append(linhas[k])

        i = j

    return resultado


def regra_power_query_tem_conteudo(regra: RegraPowerQuery) -> bool:
    return any([
        regra.observacoes,
        regra.linhas_regra,
        regra.origem_dados,
        regra.filtros_aplicados,
        regra.transformacoes,
        regra.integracoes,
    ])


_MEDIDA_TIPO_MOEDA_RE = re.compile(
    r"\b(valor|receita|custo|faturamento|fatur|venda|saldo|pago|d[eé]bito|debito|credito|cr[eé]dito|"
    r"pre[cç]o|preco|despesa|gasto|or[cç]ado|amortiz|recupera|montante|capital|"
    r"reservad|empenh|nota[_ ]?fiscal|aporte|aliena|opera[cç][aã]o)",
    re.I,
)
_MEDIDA_TIPO_PCT_NOME_RE = re.compile(
    r"(^\s*%|\bpct\b|\bpercentual\b|\b[ií]ndice\b|\btaxa\b|\bratio\b)",
    re.I,
)
_MEDIDA_TIPO_CONTAGEM_NOME_RE = re.compile(
    r"\b(qtde|quantidade|n[uú]m\b|num[_ ]?|count|contagem|contad|total[_ ]?registros?)",
    re.I,
)

# Equivalentes EN -- aplicados quando locale = en_US (alem dos PT, ja que nomes
# de medidas EN aparecem mesmo em modelos PT-BR ocasionalmente).
_MEDIDA_TIPO_MOEDA_EN_RE = re.compile(
    # NB: NAO incluir "total" -- termo generico demais (total de qualquer coisa).
    r"\b(value|revenue|cost|sales|sale|amount|price|expense|spend|budget|"
    r"paid|payment|debit|credit|invoice|balance|principal|salary|wage|fee|"
    r"income|expenditure|cash|asset|liabilit|equity|gross|net)",
    re.I,
)
_MEDIDA_TIPO_PCT_EN_RE = re.compile(
    r"(^\s*%|\bpct\b|\bpercent(?:age)?\b|\bindex\b|\brate\b|\bratio\b|\bshare\b)",
    re.I,
)
_MEDIDA_TIPO_CONTAGEM_EN_RE = re.compile(
    r"\b(qty|quantity|count|number|num\b|num_|total[_ ]?records?|of[_ ]?records?)",
    re.I,
)


def inferir_tipo_medida(nome: str, dax: str, leitura: LeituraDax) -> str:
    """Heuristica conservadora para inferir o tipo de uma medida DAX.

    Retorna a string traduzida (no idioma ativo) para uma das categorias:
    Percentual / Contagem / Moeda / Media / Soma / Extremo / Numerico.
    Designed to FAIL CLOSED -- quando ambiguo cai em "Numerico" em vez de
    chutar errado. Aplica regex PT e EN simultaneamente porque medidas
    podem ter nome em qualquer idioma independente do locale.
    """
    nome = nome or ""
    dax = dax or ""
    # Normaliza underscore como separador para que \b casar em Total_Pago, etc.
    nome_norm = nome.replace("_", " ")
    funcoes = {f.upper() for f in (leitura.funcoes or [])}

    # 1) Percentual: nome com %, palavras-chave (PT ou EN), ou FORMAT com % no DAX
    if _MEDIDA_TIPO_PCT_NOME_RE.search(nome_norm) or _MEDIDA_TIPO_PCT_EN_RE.search(nome_norm):
        return _t("measure.type.percentage")
    if re.search(r'FORMAT\s*\([^)]+?["\']\s*[#0,.\s]*%', dax):
        return _t("measure.type.percentage")

    # 2) Contagem: funcoes de count ou nome explicito (PT ou EN)
    if "DISTINCTCOUNT" in funcoes or "COUNTROWS" in funcoes or "COUNTA" in funcoes:
        return _t("measure.type.count")
    if _MEDIDA_TIPO_CONTAGEM_NOME_RE.search(nome_norm) or _MEDIDA_TIPO_CONTAGEM_EN_RE.search(nome_norm):
        return _t("measure.type.count")

    # 3) Moeda: nome com palavras-chave fortes (PT ou EN)
    if _MEDIDA_TIPO_MOEDA_RE.search(nome_norm) or _MEDIDA_TIPO_MOEDA_EN_RE.search(nome_norm):
        return _t("measure.type.currency")

    # 4) Media: AVERAGE/AVERAGEX ou nome com "media" / "average"
    if "AVERAGE" in funcoes or "AVERAGEX" in funcoes:
        return _t("measure.type.average")
    if re.search(r"\bm[eé]dia\b|\baverage\b|\bavg\b|\bmean\b", nome_norm, re.I):
        return _t("measure.type.average")

    # 5) Soma generica
    if "SUM" in funcoes or "SUMX" in funcoes:
        return _t("measure.type.sum")

    # 6) Min/Max sem outra qualificacao
    if "MIN" in funcoes or "MAX" in funcoes:
        return _t("measure.type.minmax")

    return _t("measure.type.numeric")


# Mapeamento singular/plural para o sumario de etapas Power Query
_POWER_QUERY_ETAPA_PLURAIS = {
    "Fonte": ("fonte", "fontes"),
    "Tipagem": ("tipagem", "tipagens"),
    "Filtro": ("filtro", "filtros"),
    "Renomeação": ("renomeação", "renomeações"),
    "Seleção de Colunas": ("seleção de colunas", "seleções de colunas"),
    "Remoção de Colunas": ("remoção de colunas", "remoções de colunas"),
    "Coluna Calculada": ("coluna calculada", "colunas calculadas"),
    "Coluna Condicional": ("coluna condicional", "colunas condicionais"),
    "Merge / Junção": ("junção", "junções"),
    "Append": ("append", "appends"),
    "Expansão": ("expansão", "expansões"),
    "Agrupamento": ("agrupamento", "agrupamentos"),
    "Pivot": ("pivot", "pivots"),
    "Unpivot": ("unpivot", "unpivots"),
    "Ordenação": ("ordenação", "ordenações"),
    "Deduplicação": ("deduplicação", "deduplicações"),
    "Substituição de Valores": ("substituição de valores", "substituições de valores"),
    "Cabeçalhos": ("ajuste de cabeçalho", "ajustes de cabeçalho"),
    "Tratamento de Erros": ("tratamento de erros", "tratamentos de erros"),
    "Tratamento de Erro": ("tratamento de erro", "tratamentos de erro"),
    "Recorte de Linhas": ("recorte de linhas", "recortes de linhas"),
    "Regra Condicional": ("regra condicional", "regras condicionais"),
    "Transformação Personalizada": ("transformação personalizada", "transformações personalizadas"),
}

# Espelho EN -- valores nos analisadores ficam em PT, mas na renderizacao
# (sumario + coluna Etapa) traduzimos via lookup quando locale = en_US.
_POWER_QUERY_ETAPA_EN = {
    # PT name -> (EN label, EN singular, EN plural)
    "Fonte":                       ("Source",            "source",            "sources"),
    "Tipagem":                     ("Type cast",         "type cast",         "type casts"),
    "Filtro":                      ("Filter",            "filter",            "filters"),
    "Renomeação":                  ("Rename",            "rename",            "renames"),
    "Seleção de Colunas":          ("Select Columns",    "column selection",  "column selections"),
    "Remoção de Colunas":          ("Remove Columns",    "column removal",    "column removals"),
    "Coluna Calculada":            ("Calculated Column", "calculated column", "calculated columns"),
    "Coluna Condicional":          ("Conditional Column","conditional column","conditional columns"),
    "Merge / Junção":              ("Merge / Join",      "join",              "joins"),
    "Append":                      ("Append",            "append",            "appends"),
    "Expansão":                    ("Expand",            "expand",            "expands"),
    "Agrupamento":                 ("Group By",          "group by",          "group bys"),
    "Pivot":                       ("Pivot",             "pivot",             "pivots"),
    "Unpivot":                     ("Unpivot",           "unpivot",           "unpivots"),
    "Ordenação":                   ("Sort",              "sort",              "sorts"),
    "Deduplicação":                ("Deduplicate",       "deduplication",     "deduplications"),
    "Substituição de Valores":     ("Replace Values",    "value replacement", "value replacements"),
    "Cabeçalhos":                  ("Headers",           "header adjustment", "header adjustments"),
    "Tratamento de Erros":         ("Error Handling",    "error handling",    "error handlings"),
    "Tratamento de Erro":          ("Error Handling",    "error handling",    "error handlings"),
    "Recorte de Linhas":           ("Row Slice",         "row slice",         "row slices"),
    "Regra Condicional":           ("Conditional Rule",  "conditional rule",  "conditional rules"),
    "Transformação Personalizada": ("Custom Transform",  "custom transform",  "custom transforms"),
}


def _localize_etapa_pq(etapa_pt: str) -> str:
    """Traduz o nome da etapa PT para o locale ativo. Fallback: o proprio PT."""
    if i18n.get_locale() == "en_US":
        entry = _POWER_QUERY_ETAPA_EN.get(etapa_pt)
        if entry:
            return entry[0]
    return etapa_pt


def _plurais_etapa_pq(etapa_pt: str) -> Tuple[str, str]:
    """Retorna (singular, plural) para a etapa no locale ativo."""
    if i18n.get_locale() == "en_US":
        entry = _POWER_QUERY_ETAPA_EN.get(etapa_pt)
        if entry:
            return (entry[1], entry[2])
    return _POWER_QUERY_ETAPA_PLURAIS.get(
        etapa_pt, (etapa_pt.lower(), etapa_pt.lower())
    )


def _resumir_etapas_power_query(linhas: List["LinhaRegraPowerQuery"]) -> str:
    """Conta etapas por categoria e retorna texto resumido para cabecalho.

    Ex.: '1 fonte, 4 tipagens, 6 colunas calculadas, 1 filtro, 1 append'.
    """
    if not linhas:
        return ""

    contagem: Dict[str, int] = {}
    ordem: List[str] = []
    for linha in linhas:
        etapa = (linha.etapa or "").strip()
        if not etapa:
            continue
        if etapa not in contagem:
            ordem.append(etapa)
        contagem[etapa] = contagem.get(etapa, 0) + 1

    partes: List[str] = []
    for etapa in ordem:
        count = contagem[etapa]
        singular, plural = _plurais_etapa_pq(etapa)
        nome = plural if count > 1 else singular
        partes.append(f"{count} {nome}")
    return ", ".join(partes)


def adicionar_regra_power_query_markdown(md: List[str], regra: RegraPowerQuery) -> None:
    if not regra_power_query_tem_conteudo(regra):
        return

    def escape_table(valor: str) -> str:
        return (
            str(valor or "")
            .replace("\\", "\\\\")
            .replace("|", "\\|")
            .replace("\n", "<br>")
            .strip()
        )

    md.append(f"#### {_t('pq.business_rule_inferred')}")
    md.append("")
    md.append(f"> {_t('pq.intro')}")
    if regra.observacoes:
        md.append(f"> {_t('pq.intro_with_observations')}")
    md.append("")

    if regra.observacoes:
        md.append(_t("pq.documented_observations"))
        md.append("")
        rotulos = {
            "geral": _t("pq.obs.general"),
            "origem": _t("pq.obs.source"),
            "regra": _t("pq.obs.rule"),
            "observacao": _t("pq.obs.observation"),
        }
        for chave, valores in regra.observacoes.items():
            for valor in valores:
                md.append(f"- **{rotulos.get(chave, chave.title())}**: {valor}")
        md.append("")

    if regra.linhas_regra:
        resumo = _resumir_etapas_power_query(regra.linhas_regra)
        if resumo:
            md.append(f"{_t('pq.summary_label')}: {resumo}.")
            md.append("")
        md.append(_t("pq.rules_title"))
        md.append("")
        md.append(f"| {_t('pq.col.step')} | {_t('pq.col.rule_or_filter')} | {_t('pq.col.description')} |")
        md.append("|---|---|---|")
        linhas_agrupadas = _agrupar_linhas_regra_consecutivas(regra.linhas_regra)
        for linha in linhas_agrupadas:
            md.append(
                f"| {escape_table(_localize_etapa_pq(linha.etapa))} "
                f"| {escape_table(linha.regra)} "
                f"| {escape_table(linha.descricao)} |"
            )
        md.append("")


def adicionar_linhagem_dax_markdown(
    md: List[str],
    leitura: LeituraDax,
    medidas_conhecidas: Optional[Dict[str, str]] = None,
) -> None:
    """Renderiza a linhagem (medidas referenciadas) de uma expressao DAX.

    Lista nomes [X] sem prefixo de tabela que casam com medidas conhecidas.
    """
    medidas_conhecidas = medidas_conhecidas or {}
    medidas_refs = [r for r in leitura.referencias_medidas if r in medidas_conhecidas]
    if medidas_refs:
        nomes = ", ".join(f"`[{m}]`" for m in medidas_refs)
        md.append(f"{_t('linhagem.measures_referenced')}: {nomes}")
        md.append("")


def adicionar_leitura_dax_markdown(md: List[str], leitura: LeituraDax) -> None:
    """Versao compacta: lista apenas funcoes usadas, sem repetir o glossario.

    O glossario completo de cada funcao DAX e exibido uma unica vez na secao
    "Glossario DAX" no inicio do documento. Aqui mostramos apenas a lista
    de funcoes usadas nesta expressao, agrupadas por categoria.
    """
    if not leitura.funcoes:
        return
    md.append(_t("leitura_dax.functions_used"))
    md.append("")
    if leitura.categorias:
        # Para cada categoria, lista as funcoes (sem descricao - vai no glossario).
        for categoria in DAX_CATEGORY_ORDER:
            itens = leitura.categorias.get(categoria, [])
            if not itens:
                continue
            # Extrai apenas o nome da funcao do inicio de cada item (ex: "CALCULATE: ...")
            nomes_funcao = []
            for item in itens:
                if ":" in item:
                    nome = item.split(":", 1)[0].strip()
                    # Remove backticks se houver
                    nome = nome.replace("`", "")
                    nomes_funcao.append(f"`{nome}`")
            if nomes_funcao:
                label = i18n.t_or(
                    f"dax.cat.{categoria}",
                    DAX_CATEGORY_LABELS.get(categoria, categoria.title()),
                )
                md.append(f"- **{label}**: {', '.join(nomes_funcao)}")
    else:
        # Fallback: lista funcoes diretamente
        funcoes_str = ", ".join(f"`{f}`" for f in leitura.funcoes)
        md.append(f"- {funcoes_str}")
    md.append("")


# ============================================================================
# PARSING - JSON
# ============================================================================

def parse_json(caminho: str) -> dict:
    """
    Lê e parseia arquivo JSON.
    Retorna dict vazio se arquivo não existir ou houver erro.
    """
    if not arquivo_existe(caminho):
        return {}

    try:
        with open(caminho, 'r', encoding='utf-8-sig') as f:
            return json.load(f)
    except Exception as e:
        print(f"Erro ao ler JSON {caminho}: {e}")
        return {}


# ============================================================================
# PARSING - TMDL
# ============================================================================

def extrair_expressao_multilinhas(linhas: List[str], indice_inicio: int) -> Tuple[str, int]:
    """
    Extrai expressões DAX/M que podem ocupar múltiplas linhas.

    Returns:
        (expressao, indice_final)
    """
    expressao_partes = []
    i = indice_inicio

    # Verifica se começa com ```
    linha_atual = linhas[i].strip()

    if linha_atual.endswith('```'):
        # Expressão multilinha delimitada
        i += 1
        while i < len(linhas):
            linha = linhas[i].rstrip()
            if linha.strip() == '```':
                i += 1
                break
            expressao_partes.append(linha)
            i += 1
    else:
        # Expressão inline (na mesma linha)
        # Pega tudo após o =
        if '=' in linha_atual:
            expressao_partes.append(linha_atual.split('=', 1)[1].strip())
            i += 1

    return '\n'.join(expressao_partes), i


def parse_tmdl_relationships(caminho: str) -> List[InfoRelacionamento]:
    """
    Parseia arquivo relationships.tmdl
    """
    if not arquivo_existe(caminho):
        return []

    relacionamentos = []

    try:
        with open(caminho, 'r', encoding='utf-8-sig') as f:
            linhas = f.readlines()

        i = 0
        while i < len(linhas):
            linha = linhas[i].strip()

            # Detecta início de relacionamento
            match = re.match(r'^relationship\s+(.+)$', linha)
            if match:
                rel_id = match.group(1).strip()
                rel_info = {
                    'id': rel_id,
                    'fromTable': None,
                    'fromColumn': None,
                    'toTable': None,
                    'toColumn': None,
                    'crossFilteringBehavior': 'singleDirection',
                    'isActive': True,
                    'joinOnDateBehavior': None
                }

                # Lê propriedades do relacionamento
                i += 1
                while i < len(linhas):
                    linha_prop = linhas[i].strip()

                    # Fim do bloco de relacionamento
                    if linha_prop and not linha_prop.startswith(('fromColumn', 'toColumn', 'crossFilteringBehavior', 'isActive', 'joinOnDateBehavior')):
                        break

                    # fromColumn ou toColumn
                    match_col = re.match(r'^(fromColumn|toColumn):\s*(.+)\.(.+)$', linha_prop)
                    if match_col:
                        tipo = match_col.group(1)
                        tabela = limpar_nome(match_col.group(2))
                        coluna = limpar_nome(match_col.group(3))

                        if tipo == 'fromColumn':
                            rel_info['fromTable'] = tabela
                            rel_info['fromColumn'] = coluna
                        else:
                            rel_info['toTable'] = tabela
                            rel_info['toColumn'] = coluna

                    # Outras propriedades
                    match_prop = re.match(r'^(\w+):\s*(.+)$', linha_prop)
                    if match_prop:
                        prop_nome = match_prop.group(1)
                        prop_valor = match_prop.group(2).strip()

                        if prop_nome == 'crossFilteringBehavior':
                            rel_info['crossFilteringBehavior'] = prop_valor
                        elif prop_nome == 'isActive':
                            rel_info['isActive'] = prop_valor.lower() != 'false'
                        elif prop_nome == 'joinOnDateBehavior':
                            rel_info['joinOnDateBehavior'] = prop_valor

                    i += 1

                # Cria objeto InfoRelacionamento
                if rel_info['fromTable'] and rel_info['toTable']:
                    relacionamento = InfoRelacionamento(
                        id=rel_info['id'],
                        tabela_origem=rel_info['fromTable'],
                        coluna_origem=rel_info['fromColumn'],
                        tabela_destino=rel_info['toTable'],
                        coluna_destino=rel_info['toColumn'],
                        filtro_bidirecional=(rel_info['crossFilteringBehavior'] == 'bothDirections'),
                        esta_ativo=rel_info['isActive'],
                        comportamento_data=rel_info['joinOnDateBehavior']
                    )
                    relacionamentos.append(relacionamento)

                continue

            i += 1

    except Exception as e:
        print(f"Erro ao parsear relationships.tmdl: {e}")

    return relacionamentos


def parse_tmdl_model(caminho: str) -> InfoModelo:
    """
    Parseia arquivo model.tmdl
    """
    info = InfoModelo()

    if not arquivo_existe(caminho):
        return info

    try:
        with open(caminho, 'r', encoding='utf-8-sig') as f:
            linhas = f.readlines()

        grupo_atual = None

        for linha in linhas:
            linha_strip = linha.strip()

            # Cultura
            if linha_strip.startswith('culture:'):
                info.cultura = linha_strip.split(':', 1)[1].strip()

            # Versão datasource
            elif linha_strip.startswith('defaultPowerBIDataSourceVersion:'):
                info.versao_datasource = linha_strip.split(':', 1)[1].strip()

            # Query group
            elif 'queryGroup' in linha_strip:
                match = re.search(r"queryGroup\s+['\"](.+?)['\"]", linha_strip)
                if match:
                    grupo_atual = {'nome': match.group(1), 'ordem': 0}

            # Ordem do grupo
            elif grupo_atual and 'PBI_QueryGroupOrder' in linha_strip:
                match = re.search(r'=\s*(\d+)', linha_strip)
                if match:
                    grupo_atual['ordem'] = int(match.group(1))
                    info.grupos_consulta.append(grupo_atual)
                    grupo_atual = None

            # Referências de tabela
            elif linha_strip.startswith('ref table'):
                match = re.search(r"ref table\s+['\"](.+?)['\"]", linha_strip)
                if match:
                    info.tabelas_referenciadas.append(match.group(1))

            # Annotations
            elif linha_strip.startswith('annotation'):
                match = re.match(r'annotation\s+(\S+)\s*=\s*(.+)$', linha_strip)
                if match:
                    info.annotations[match.group(1)] = match.group(2)

    except Exception as e:
        print(f"Erro ao parsear model.tmdl: {e}")

    return info


def _is_fact_table(nome: str) -> bool:
    if not nome:
        return False

    nome_norm = nome.strip().lower()
    if re.search(r"(^|[^a-z0-9])(fato|fact|fat)([^a-z0-9]|$)", nome_norm):
        return True
    if re.search(r"(^|[^a-z0-9])f_", nome_norm):
        return True
    return False


def parse_tmdl_table(caminho: str) -> Optional[InfoTabela]:
    """
    Parseia arquivo .tmdl de uma tabela.
    DEFENSIVO: Retorna None em qualquer erro, nunca falha.
    """
    if not arquivo_existe(caminho):
        return None

    try:
        with open(caminho, 'r', encoding='utf-8-sig') as f:
            linhas = f.readlines()
    except Exception as e:
        print(f"  [AVISO] Erro ao ler {caminho}: {e}")
        return None

    # Validação básica
    if not linhas:
        return None

    # Extrai nome da tabela
    nome_tabela = None
    try:
        for linha in linhas[:min(10, len(linhas))]:  # Procura nas primeiras linhas
            match = re.match(r"^table\s+['\"]?(.+?)['\"]?\s*$", linha.strip())
            if match:
                nome_tabela = limpar_nome(match.group(1))
                break
    except Exception:
        pass

    if not nome_tabela:
        return None

    tabela = InfoTabela(nome=nome_tabela)
    tabela.eh_fato = _is_fact_table(nome_tabela)

    # Parse linha por linha
    i = 0
    while i < len(linhas):
        linha = linhas[i].strip()

        # Tabela oculta
        if linha == 'isHidden':
            tabela.esta_oculta = True

        # Excluída de refresh
        elif linha == 'excludeFromModelRefresh':
            tabela.excluida_refresh = True

        # Descrição da tabela
        elif linha.startswith('description:'):
            tabela.descricao = linha.split(':', 1)[1].strip().strip("'\"")

        # Coluna regular ou calculada
        elif linha.startswith('column '):
            # Tenta capturar nome entre aspas primeiro
            match_quoted = re.match(r"^column\s+['\"](.+?)['\"]\s*(?:=\s*(.+))?$", linha)
            # Se não tiver aspas, captura até o = (se houver) ou até o fim
            match_unquoted = re.match(r"^column\s+(\S+)\s*(?:=\s*(.+))?$", linha)

            match = match_quoted or match_unquoted
            if match:
                nome_col = limpar_nome(match.group(1))
                expressao = match.group(2)

                if expressao:
                    # Coluna calculada
                    if expressao.strip() == '```':
                        # Expressão multilinha
                        expressao_completa, i = extrair_expressao_multilinhas(linhas, i)
                    else:
                        expressao_completa = expressao.strip()
                        i += 1

                    col_calc = InfoColunaCalculada(nome=nome_col, expressao_dax=expressao_completa)

                    # Lê propriedades da coluna calculada
                    while i < len(linhas):
                        linha_prop = linhas[i].strip()
                        if not linha_prop.startswith(('dataType', 'formatString', 'summarizeBy', 'lineageTag', 'annotation',
                                                     'changedProperty')):
                            break

                        if linha_prop.startswith('dataType:'):
                            col_calc.tipo_dado = linha_prop.split(':', 1)[1].strip()
                        elif linha_prop.startswith('formatString:'):
                            col_calc.formato = linha_prop.split(':', 1)[1].strip()
                        elif linha_prop.startswith('summarizeBy:'):
                            col_calc.sumarizacao = linha_prop.split(':', 1)[1].strip()

                        i += 1

                    tabela.colunas_calculadas.append(col_calc)
                    continue

                else:
                    # Coluna regular
                    coluna = InfoColuna(nome=nome_col)
                    i += 1

                    # Lê propriedades da coluna
                    while i < len(linhas):
                        linha_prop = linhas[i].strip()

                        if not linha_prop.startswith(('dataType', 'formatString', 'summarizeBy', 'sourceColumn',
                                                     'isHidden', 'dataCategory', 'lineageTag', 'annotation',
                                                     'changedProperty', 'variation')):
                            break

                        # Bloco variation: pula todas as sub-linhas indentadas
                        if linha_prop.startswith('variation'):
                            i += 1
                            while i < len(linhas):
                                sub_linha = linhas[i]
                                # Sub-propriedades são indentadas com tabs/espaços extras
                                if sub_linha.strip() and not sub_linha.startswith('\t\t\t') and not sub_linha.startswith('            '):
                                    break
                                i += 1
                            continue


                        if linha_prop.startswith('dataType:'):
                            coluna.tipo_dado = linha_prop.split(':', 1)[1].strip()
                        elif linha_prop.startswith('formatString:'):
                            coluna.formato = linha_prop.split(':', 1)[1].strip()
                        elif linha_prop.startswith('summarizeBy:'):
                            coluna.sumarizacao = linha_prop.split(':', 1)[1].strip()
                        elif linha_prop.startswith('sourceColumn:'):
                            coluna.coluna_fonte = linha_prop.split(':', 1)[1].strip()
                        elif linha_prop == 'isHidden':
                            coluna.esta_oculta = True
                        elif linha_prop.startswith('dataCategory:'):
                            coluna.categoria_dados = linha_prop.split(':', 1)[1].strip()

                        i += 1

                    tabela.colunas.append(coluna)
                    continue

        # Medida
        elif linha.startswith('measure '):
            # Tenta capturar nome entre aspas primeiro (com ou sem =)
            match_quoted = re.match(r"^measure\s+['\"](.+?)['\"]\s*=?", linha)
            # Se não tiver aspas, captura até o = (se houver) ou fim da linha
            match_unquoted = re.match(r"^measure\s+(\S+)\s*=?", linha)

            match = match_quoted or match_unquoted
            if match:
                nome_medida = limpar_nome(match.group(1))

                # Extrai expressão - verifica se tem = na linha
                if '=' in linha:
                    # Tem = na mesma linha
                    partes_split = linha.split('=', 1)
                    if len(partes_split) > 1 and partes_split[1].strip():
                        # Tem conteúdo após o =
                        if linha.rstrip().endswith('```'):
                            expressao, i = extrair_expressao_multilinhas(linhas, i)
                        else:
                            # Expressão inline
                            expressao = partes_split[1].strip()
                            i += 1
                    else:
                        # Tem = mas nada depois - expressão na próxima linha
                        i += 1
                        expressao = ""
                        # Pula linhas vazias
                        while i < len(linhas) and not linhas[i].strip():
                            i += 1
                        if i < len(linhas):
                            proxima = linhas[i].strip()
                            if proxima and not proxima.startswith(('formatString', 'lineageTag', 'annotation', 'description:', 'displayFolder:')):
                                # A próxima linha é a expressão
                                if proxima == '```' or proxima.startswith('```'):
                                    expressao, i = extrair_expressao_multilinhas(linhas, i)
                                else:
                                    # Expressão DAX multilinha SEM delimitadores ```
                                    # Continua lendo até encontrar propriedade ou nova entidade
                                    linhas_dax = []
                                    while i < len(linhas):
                                        linha_atual = linhas[i]
                                        linha_strip = linha_atual.strip()

                                        # Para se encontrar uma propriedade TMDL ou nova entidade
                                        if linha_strip and (
                                            linha_strip.startswith(('formatString:', 'formatStringDefinition', 'lineageTag', 'annotation', 'changedProperty', 'description:', 'displayFolder:')) or
                                            linha_strip.startswith(('measure ', 'column ', 'hierarchy ', 'partition ', 'expression'))
                                        ):
                                            break

                                        # Adiciona a linha (mantém indentação original do DAX)
                                        linhas_dax.append(linha_atual.rstrip())
                                        i += 1

                                    # Junta as linhas e remove indentação comum
                                    expressao = '\n'.join(linhas_dax).strip()
                else:
                    # Não tem = na mesma linha - verifica próxima
                    i += 1
                    expressao = ""
                    if i < len(linhas):
                        proxima = linhas[i].strip()
                        # Medida sem expressão (apenas declarada)
                        if not proxima or proxima.startswith(('changedProperty', 'formatString', 'lineageTag', 'annotation', 'description:', 'displayFolder:',
                                                              'measure ', 'column ', 'hierarchy ', 'partition ')):
                            pass  # expressão fica vazia
                        elif proxima.startswith('='):
                            if '```' in proxima:
                                expressao, i = extrair_expressao_multilinhas(linhas, i)
                            else:
                                expressao = proxima.split('=', 1)[1].strip() if '=' in proxima else ""
                                i += 1
                        elif proxima == '```':
                            expressao, i = extrair_expressao_multilinhas(linhas, i)

                medida = InfoMedida(nome=nome_medida, expressao_dax=expressao)

                # Lê propriedades da medida
                while i < len(linhas):
                    linha_prop = linhas[i].strip()

                    if not linha_prop.startswith(('formatString:', 'formatStringDefinition', 'lineageTag', 'annotation',
                                                 'changedProperty', 'description:', 'displayFolder:')):
                        break

                    if linha_prop.startswith('formatString:'):
                        medida.formato = linha_prop.split(':', 1)[1].strip()
                    elif linha_prop.startswith('description:'):
                        medida.descricao = linha_prop.split(':', 1)[1].strip().strip("'\"")
                    elif linha_prop.startswith('formatStringDefinition'):
                        # Formato dinâmico
                        if linha_prop.endswith('```'):
                            formato_din, i = extrair_expressao_multilinhas(linhas, i)
                            medida.formato_dinamico = formato_din
                            continue

                    i += 1

                tabela.medidas.append(medida)
            else:
                # DEBUG: Se não conseguiu fazer match, mostra a linha
                print(f"  [DEBUG] Falha ao parsear medida: {linha[:80]}")
            continue

        # Hierarquia
        elif linha.startswith('hierarchy '):
            match = re.match(r"^hierarchy\s+['\"]?(.+?)['\"]?$", linha)
            if match:
                nome_hier = limpar_nome(match.group(1))
                hierarquia = InfoHierarquia(nome=nome_hier)
                i += 1

                # Lê níveis
                while i < len(linhas):
                    linha_prop = linhas[i].strip()

                    if linha_prop.startswith('level '):
                        match_nivel = re.match(r"^level\s+['\"]?(.+?)['\"]?$", linha_prop)
                        if match_nivel:
                            hierarquia.niveis.append(limpar_nome(match_nivel.group(1)))
                    elif not linha_prop.startswith(('lineageTag', 'column:')):
                        break

                    i += 1

                tabela.hierarquias.append(hierarquia)
                continue

        # Partição
        elif linha.startswith('partition '):
            match = re.match(r"^partition\s+['\"]?(.+?)['\"]?\s*=\s*(\w+)", linha)
            if match:
                nome_part = limpar_nome(match.group(1))
                tipo_part = match.group(2)

                particao = InfoParticao(nome=nome_part)
                i += 1

                # Lê propriedades
                while i < len(linhas):
                    linha_prop = linhas[i].strip()
                    linha_raw = linhas[i]  # Mantém indentação original

                    if linha_prop.startswith('mode:'):
                        particao.modo = linha_prop.split(':', 1)[1].strip()
                    elif linha_prop.startswith('queryGroup:'):
                        particao.grupo_consulta = limpar_nome(linha_prop.split(':', 1)[1].strip())
                    elif linha_prop.startswith('source') or linha_prop.startswith('expression'):
                        # Código fonte M ou DAX
                        if '```' in linha_prop:
                            # Com delimitadores ```
                            codigo, i = extrair_expressao_multilinhas(linhas, i)
                            particao.codigo_fonte = codigo
                            break
                        elif '=' in linha_prop:
                            # source = ou source = código
                            partes = linha_prop.split('=', 1)
                            resto = partes[1].strip() if len(partes) > 1 else ""

                            if resto:
                                # Código inline após =
                                particao.codigo_fonte = resto
                                i += 1
                            else:
                                # Código nas próximas linhas (indentado)
                                i += 1
                                linhas_codigo = []

                                # Pega todas as linhas indentadas até próxima propriedade
                                while i < len(linhas):
                                    linha_cod = linhas[i]
                                    linha_cod_strip = linha_cod.strip()

                                    # Para se encontrar linha não indentada (nova propriedade)
                                    # ou linha de annotation/outra seção
                                    if linha_cod_strip and not linha_cod.startswith('\t') and not linha_cod.startswith('    '):
                                        break
                                    if linha_cod_strip.startswith('annotation ') or linha_cod_strip.startswith('lineageTag'):
                                        break

                                    linhas_codigo.append(linha_cod.rstrip())
                                    i += 1

                                particao.codigo_fonte = '\n'.join(linhas_codigo).strip()
                            break
                        else:
                            i += 1
                            break
                    elif not linha_prop.startswith(('lineageTag', 'annotation')):
                        break

                    i += 1

                tabela.particao = particao
                continue

        i += 1

    return tabela


# ============================================================================
# CLASSE PRINCIPAL
# ============================================================================

class DocumentadorPBIP:
    """
    Classe principal para documentar projetos Power BI (.pbip)
    """

    def __init__(self, caminho_projeto: str):
        """
        Inicializa o documentador

        Args:
            caminho_projeto: Caminho para a pasta do projeto (que contém o .pbip)
        """
        self.caminho_projeto = Path(caminho_projeto)
        self.nome_projeto = None
        self.info_modelo = InfoModelo()
        self.tabelas: List[InfoTabela] = []
        self.relacionamentos: List[InfoRelacionamento] = []
        self.paginas: List[InfoPagina] = []
        self.info_report = {}
        self.total_visuais = 0
        self.visuais_personalizados = []
        self.recursos_imagem = []
        self._md_cache: Optional[str] = None
        self._dicionario_cache: Optional[List[TermoDicionario]] = None
        self.branding = BrandingConfig()

        # Layout do projeto: 'novo' (Model/, Report/) ou 'antigo' ({nome}.SemanticModel/definition/)
        self.layout = None

        # Encontra o projeto e detecta o layout
        self._encontrar_pbip()

    def _encontrar_pbip(self):
        """Encontra o projeto Power BI na pasta e detecta o layout"""

        # === Layout NOVO: pasta Model/ diretamente na raiz ===
        pasta_model = self.caminho_projeto / "Model"
        if pasta_model.is_dir():
            self.layout = 'novo'
            # Nome do projeto: arquivo .pbip, nome da pasta, ou stem da pasta
            arquivos_pbip = [p for p in self.caminho_projeto.glob("*.pbip") if p.is_file()]
            if arquivos_pbip:
                self.nome_projeto = arquivos_pbip[0].stem
            else:
                # Usa o nome da pasta (remove .pbip do nome se existir)
                nome_pasta = self.caminho_projeto.name
                if nome_pasta.lower().endswith('.pbip'):
                    self.nome_projeto = self.caminho_projeto.stem
                else:
                    self.nome_projeto = nome_pasta
            print(f"[OK] Projeto encontrado (layout novo): {self.nome_projeto}")
            return

        # === Layout ANTIGO: {nome}.SemanticModel/definition/ ===
        self.layout = 'antigo'

        # Padrão 1: Arquivo .pbip dentro da pasta
        arquivos_pbip = [p for p in self.caminho_projeto.glob("*.pbip") if p.is_file()]
        if arquivos_pbip:
            self.nome_projeto = arquivos_pbip[0].stem
            print(f"[OK] Projeto encontrado: {self.nome_projeto}")
            return

        # Padrão 2: Detecta pelo subdiretório *.SemanticModel
        semantic_dirs = list(self.caminho_projeto.glob("*.SemanticModel"))
        if semantic_dirs:
            self.nome_projeto = semantic_dirs[0].name.replace(".SemanticModel", "")
            print(f"[OK] Projeto encontrado (via SemanticModel): {self.nome_projeto}")
            return

        # Padrão 3: O nome da própria pasta termina com .pbip
        if self.caminho_projeto.suffix.lower() == '.pbip':
            nome_base = self.caminho_projeto.stem
            semantic_check = self.caminho_projeto / f"{nome_base}.SemanticModel"
            if semantic_check.exists():
                self.nome_projeto = nome_base
                print(f"[OK] Projeto encontrado (pasta .pbip): {self.nome_projeto}")
                return
            for item in self.caminho_projeto.iterdir():
                if item.is_dir() and item.name.endswith(".SemanticModel"):
                    self.nome_projeto = item.name.replace(".SemanticModel", "")
                    print(f"[OK] Projeto encontrado (pasta .pbip, SemanticModel): {self.nome_projeto}")
                    return

        raise FileNotFoundError(
            f"Nenhum projeto Power BI encontrado em {self.caminho_projeto}\n"
            f"Verifique se a pasta contém Model/ ou {'{nome}'}.SemanticModel/"
        )

    def _deve_filtrar_tabela(self, nome_tabela: str) -> bool:
        """Verifica se a tabela deve ser filtrada (tabelas técnicas)"""
        padroes_filtro = [
            r'^LocalDateTable_',
            r'^DateTableTemplate_',
        ]

        for padrao in padroes_filtro:
            if re.match(padrao, nome_tabela):
                return True

        return False

    def extrair_informacoes(self):
        """Extrai todas as informações do projeto"""
        self._md_cache = None  # Invalida cache ao re-extrair
        self._dicionario_cache = None
        print("\nExtraindo informações do modelo semântico...")
        self._extrair_modelo()

        print("Extraindo tabelas...")
        self._extrair_tabelas()

        print("Extraindo relacionamentos...")
        self._extrair_relacionamentos()

        print("Extraindo páginas do relatório...")
        self._extrair_paginas()

        print("Extraindo informações do relatório...")
        self._extrair_info_report()

        # Totaliza visuais a partir das páginas processadas
        self.total_visuais = sum(len(p.visuais) for p in self.paginas)

        # Indexa medidas para uso na linhagem (Medidas referenciadas)
        self._indexar_medidas()

        print("[OK] Extração concluída!\n")

    def _indexar_medidas(self):
        """Constroi medidas_index: nome_medida -> nome_tabela.

        Usado pela linhagem DAX (`adicionar_linhagem_dax_markdown`) para filtrar
        candidatos `[X]` extraidos via regex apenas para medidas realmente
        existentes no modelo.
        """
        self.medidas_index: Dict[str, str] = {}
        for tabela in self.tabelas:
            for medida in tabela.medidas:
                self.medidas_index[medida.nome] = tabela.nome

    def _obter_pasta_modelo(self) -> Path:
        """Retorna o caminho da pasta do modelo conforme o layout"""
        if self.layout == 'novo':
            return self.caminho_projeto / "Model"
        else:
            return self.caminho_projeto / f"{self.nome_projeto}.SemanticModel" / "definition"

    def _obter_pasta_report(self) -> Path:
        """Retorna o caminho da pasta do relatório conforme o layout"""
        if self.layout == 'novo':
            return self.caminho_projeto / "Report"
        else:
            return self.caminho_projeto / f"{self.nome_projeto}.Report" / "definition"

    def _extrair_modelo(self):
        """Extrai informações do modelo"""
        caminho_model = self._obter_pasta_modelo() / "model.tmdl"
        self.info_modelo = parse_tmdl_model(str(caminho_model))

    def _extrair_tabelas(self):
        """Extrai informações de todas as tabelas"""
        try:
            pasta_tables = self._obter_pasta_modelo() / "tables"

            if not pasta_tables.exists():
                print("  [AVISO] Pasta de tabelas não encontrada")
                return

            arquivos_tmdl = list(pasta_tables.glob("*.tmdl"))
            print(f"  Encontradas {len(arquivos_tmdl)} tabelas")

            for arquivo in arquivos_tmdl:
                try:
                    tabela = parse_tmdl_table(str(arquivo))

                    if tabela:
                        # Filtra tabelas técnicas
                        if self._deve_filtrar_tabela(tabela.nome):
                            print(f"  [X] Ignorando tabela técnica: {tabela.nome}")
                            continue

                        self.tabelas.append(tabela)
                        print(f"  [OK] {tabela.nome} ({len(tabela.colunas)} colunas, {len(tabela.medidas)} medidas)")
                except Exception as e:
                    print(f"  [AVISO] Erro ao processar {arquivo.name}: {e}")
                    continue
        except Exception as e:
            print(f"  [ERRO] Falha ao extrair tabelas: {e}")

    def _extrair_relacionamentos(self):
        """Extrai relacionamentos entre tabelas"""
        try:
            caminho_rel = self._obter_pasta_modelo() / "relationships.tmdl"
            self.relacionamentos = parse_tmdl_relationships(str(caminho_rel))
            print(f"  Encontrados {len(self.relacionamentos)} relacionamentos")
        except Exception as e:
            print(f"  [AVISO] Erro ao extrair relacionamentos: {e}")
            self.relacionamentos = []

    def _extrair_paginas(self):
        """Extrai informações das páginas do relatório"""
        try:
            pasta_report = self._obter_pasta_report()

            if self.layout == 'novo':
                # Layout novo: Report/sections/*/section.json
                pasta_sections = pasta_report / "sections"
                if not pasta_sections.exists():
                    print("  [AVISO] Pasta de seções não encontrada")
                    return

                for secao_dir in sorted(pasta_sections.iterdir()):
                    if not secao_dir.is_dir():
                        continue
                    try:
                        caminho_section = secao_dir / "section.json"
                        section_data = parse_json(str(caminho_section))

                        if section_data:
                            # displayOption pode ser int (0, 1) ou string
                            display_opt = section_data.get('displayOption', 'FitToPage')
                            if isinstance(display_opt, int):
                                display_opt = {0: 'FitToPage', 1: 'FitToPage', 2: 'FitToWidth'}.get(display_opt, str(display_opt))

                            pagina = InfoPagina(
                                id=section_data.get('name', secao_dir.name),
                                nome_exibicao=section_data.get('displayName', secao_dir.name),
                                largura=section_data.get('width', 1280),
                                altura=section_data.get('height', 720),
                                opcao_exibicao=display_opt,
                                tipo=traduzir_enum_pbir('Drillthrough' if section_data.get('pageBinding') else 'Normal'),
                                filtros=section_data.get('filterConfig', {}).get('filters') if section_data.get('filterConfig') else None
                            )

                            # Extrair visuais desta página
                            self._extrair_visuais_pagina(pagina, secao_dir)

                            self.paginas.append(pagina)
                            print(f"  [OK] {pagina.nome_exibicao} ({pagina.tipo}) - {len(pagina.visuais)} visuais")
                    except Exception as e:
                        print(f"  [AVISO] Erro ao processar seção {secao_dir.name}: {e}")
                        continue
            else:
                # Layout antigo: {nome}.Report/definition/pages/pages.json
                pasta_pages = pasta_report / "pages"
                if not pasta_pages.exists():
                    print("  [AVISO] Pasta de páginas não encontrada")
                    return

                caminho_pages_json = pasta_pages / "pages.json"
                pages_data = parse_json(str(caminho_pages_json))
                page_order = pages_data.get('pageOrder', [])

                for page_id in page_order:
                    try:
                        pasta_pagina = pasta_pages / page_id
                        caminho_page_json = pasta_pagina / "page.json"
                        page_data = parse_json(str(caminho_page_json))

                        if page_data:
                            pagina = InfoPagina(
                                id=page_id,
                                nome_exibicao=page_data.get('displayName', page_id),
                                largura=page_data.get('width', 1280),
                                altura=page_data.get('height', 720),
                                opcao_exibicao=page_data.get('displayOption', 'FitToPage'),
                                tipo=traduzir_enum_pbir('Drillthrough' if page_data.get('pageBinding') else 'Normal'),
                                filtros=page_data.get('filterConfig', {}).get('filters') if page_data.get('filterConfig') else None
                            )

                            # Extrair visuais desta página
                            self._extrair_visuais_pagina(pagina, pasta_pagina)

                            self.paginas.append(pagina)
                            print(f"  [OK] {pagina.nome_exibicao} ({pagina.tipo}) - {len(pagina.visuais)} visuais")
                    except Exception as e:
                        print(f"  [AVISO] Erro ao processar página {page_id}: {e}")
                        continue
        except Exception as e:
            print(f"  [AVISO] Erro ao extrair páginas: {e}")

    def _extrair_campos_visual_config(self, config: Dict) -> List[str]:
        """Extrai nomes de campos usados em visuais de forma best-effort."""
        campos: List[str] = []

        def adicionar(valor):
            if not isinstance(valor, str):
                return
            valor = valor.strip().strip("'\"")
            if not valor or _texto_parece_tecnico_ou_valor(valor):
                return

            referencias = _extrair_referencias_dax(valor)
            candidatos = referencias if referencias else [valor]
            for candidato in candidatos:
                candidato = str(candidato or "").strip().strip("'\"")
                if not candidato or _texto_parece_tecnico_ou_valor(candidato):
                    continue
                if candidato not in campos:
                    campos.append(candidato)

        def visitar(obj, chave=""):
            if len(campos) >= 40:
                return
            if isinstance(obj, dict):
                if "Column" in obj and isinstance(obj.get("Column"), dict):
                    col = obj.get("Column", {})
                    adicionar(col.get("Property"))
                    src = col.get("Expression", {}).get("SourceRef", {})
                    adicionar(src.get("Entity"))

                if "Measure" in obj and isinstance(obj.get("Measure"), dict):
                    medida = obj.get("Measure", {})
                    adicionar(medida.get("Property"))
                    src = medida.get("Expression", {}).get("SourceRef", {})
                    adicionar(src.get("Entity"))

                for k, v in obj.items():
                    if str(k).lower() in {"queryref", "nativequeryref", "displayname", "field", "property"}:
                        adicionar(v)
                    visitar(v, k)
            elif isinstance(obj, list):
                for item in obj:
                    visitar(item, chave)
            elif isinstance(obj, str):
                lower = str(chave or "").lower()
                if lower in {"queryref", "nativequeryref", "displayname", "field", "property"} or "[" in obj:
                    adicionar(obj)

        try:
            visitar(config)
        except Exception:
            return campos

        return campos[:40]

    def _extrair_visuais_pagina(self, pagina: InfoPagina, diretorio_pagina: Path):
        """Extrai detalhes dos visuais de uma página"""
        pasta_visuais = diretorio_pagina / "visualContainers"
        if not pasta_visuais.exists():
            return

        for visual_dir in pasta_visuais.iterdir():
            if not visual_dir.is_dir():
                continue

            caminho_config = visual_dir / "config.json"
            if not caminho_config.exists():
                continue

            try:
                config = parse_json(str(caminho_config))
                if not config:
                    continue

                # Tipo do visual
                tipo_visual = "Unknown"
                if 'singleVisual' in config:
                    tipo_visual = config['singleVisual'].get('visualType', 'Unknown')
                elif 'visualGroup' in config:
                    tipo_visual = 'Group'

                # Título do visual
                titulo = None

                # Tenta extrair título de vcObjects (comum)
                vc_objects = config.get('singleVisual', {}).get('vcObjects', {})
                if not vc_objects:
                    # Tenta no nível layout/config (ás vezes acontece)
                    vc_objects = config.get('vcObjects', {})

                if vc_objects:
                    title_obj = vc_objects.get('title', [])
                    if title_obj and isinstance(title_obj, list):
                        props = title_obj[0].get('properties', {})
                        text_prop = props.get('text', {})
                        # Estrutura complexa: text -> expr -> Literal -> Value
                        try:
                            titulo = text_prop.get('expr', {}).get('Literal', {}).get('Value', None)
                        except Exception as e:
                            print(f"  [AVISO] Falha ao extrair título (expr.Literal.Value) do visual {visual_dir.name}: {e}")

                # Se não achou, tenta em objects.general
                if not titulo:
                    objects = config.get('singleVisual', {}).get('objects', {})
                    general = objects.get('general', [])
                    if general and isinstance(general, list):
                        try:
                            # A estrutura pode variar muito aqui, é uma tentativa genérica
                            props = general[0].get('properties', {})
                            titulo = props.get('title', None)
                        except Exception as e:
                            print(f"  [AVISO] Falha ao extrair título (objects.general) do visual {visual_dir.name}: {e}")

                # Limpa o título (remove aspas e sufixo D/L se houver)
                if titulo:
                    # Se tiver aspas no começo e letras no final (ex: 'Título'D), remove a letra primeiro
                    if (titulo.endswith('D') or titulo.endswith('L')) and len(titulo) > 1 and titulo[-2] in ["'", '"']:
                         titulo = titulo[:-1]
                    titulo = titulo.strip("'\"")

                campos_visual = self._extrair_campos_visual_config(config)

                pagina.visuais.append(InfoVisual(
                    tipo=traduzir_tipo_visual(tipo_visual),
                    titulo=titulo if titulo else "-",
                    nome=visual_dir.name,
                    campos=campos_visual,
                ))
            except Exception as e:
                # Ignora visual se der erro no parse, mas loga
                print(f"  [AVISO] Erro total no parse do visual {visual_dir.name}: {e}")

    def _extrair_info_report(self):
        """Extrai informações do relatório"""
        try:
            caminho_report = self._obter_pasta_report() / "report.json"
            self.info_report = parse_json(str(caminho_report))

            # Extrai visuais personalizados (com segurança)
            if self.info_report:
                self.visuais_personalizados = self.info_report.get('publicCustomVisuals', [])

                # Extrai recursos de imagem
                resource_packages = self.info_report.get('resourcePackages', [])
                if isinstance(resource_packages, list):
                    for package in resource_packages:
                        # Layout novo: items dentro de resourcePackage
                        # Layout antigo: items direto no package
                        inner = package.get('resourcePackage', package)
                        items = inner.get('items', [])
                        if isinstance(items, list):
                            for item in items:
                                item_type = item.get('type', '')
                                # type pode ser string 'Image' ou int 100
                                if item_type == 'Image' or item_type == 100:
                                    self.recursos_imagem.append({
                                        'nome': item.get('name', 'sem_nome'),
                                        'tipo': 'Image'
                                    })
        except Exception as e:
            print(f"  [AVISO] Erro ao extrair informações do relatório: {e}")

    def _formatar_codigo_fonte(self, codigo: str) -> str:
        """Formata código fonte para melhor legibilidade"""
        if not codigo:
            return codigo
        # Substitui #(lf) por quebras de linha reais
        codigo = codigo.replace('#(lf)', '\n')
        # Remove indentação excessiva inicial
        linhas = codigo.split('\n')
        if linhas:
            # Encontra indentação mínima (ignorando linhas vazias)
            min_indent = float('inf')
            for linha in linhas:
                if linha.strip():
                    espacos = len(linha) - len(linha.lstrip())
                    min_indent = min(min_indent, espacos)

            if min_indent < float('inf') and min_indent > 0:
                linhas = [linha[min_indent:] if len(linha) >= min_indent else linha for linha in linhas]

            codigo = '\n'.join(linhas)
        return codigo.strip()

    def _interpretar_filtros_pagina(self, filtros: List[Dict]) -> List[Dict]:
        """
        Interpreta a estrutura JSON de filtros de página do Power BI
        e retorna uma lista simplificada: [{tabela, coluna, tipo, valores}]
        """
        resultado = []
        if not filtros or not isinstance(filtros, list):
            return resultado

        for filtro_raw in filtros:
            try:
                # filtro_raw pode ser string JSON ou dict
                if isinstance(filtro_raw, str):
                    import json
                    filtro_raw = json.loads(filtro_raw)

                if not isinstance(filtro_raw, dict):
                    continue

                tipo_filtro = filtro_raw.get('type', 'Unknown')
                filtro = filtro_raw.get('filter', filtro_raw)

                tabela = ''
                coluna = ''
                valores = []

                # Tenta extrair tabela/coluna de múltiplas estruturas possíveis
                # Estrutura "From" (filtros básicos)
                from_list = filtro.get('From', [])
                if from_list and isinstance(from_list, list):
                    tabela = from_list[0].get('Entity', from_list[0].get('Name', ''))

                # Estrutura "Column" (filtros avançados)
                col_obj = filtro.get('Column', {})
                if col_obj:
                    coluna = col_obj.get('Property', '')
                    src_ref = col_obj.get('Expression', {}).get('SourceRef', {})
                    if src_ref and not tabela:
                        tabela = src_ref.get('Entity', src_ref.get('Source', ''))

                # Estrutura "Where" para extrair valores
                where = filtro.get('Where', [])
                if where and isinstance(where, list):
                    for cond_wrapper in where:
                        cond = cond_wrapper.get('Condition', {})
                        # Comparação simples
                        comp = cond.get('Comparison', {})
                        if comp:
                            right = comp.get('Right', {}).get('Literal', {}).get('Value', '')
                            if right:
                                valores.append(str(right).strip("'"))
                        # In list
                        in_list = cond.get('In', {}).get('Values', [])
                        if in_list:
                            for v_list in in_list:
                                if isinstance(v_list, list):
                                    for v in v_list:
                                        lit = v.get('Literal', {}).get('Value', '')
                                        if lit:
                                            valores.append(str(lit).strip("'"))
                                elif isinstance(v_list, dict):
                                    lit = v_list.get('Literal', {}).get('Value', '')
                                    if lit:
                                        valores.append(str(lit).strip("'"))
                        # Not condition
                        not_cond = cond.get('Not', {})
                        if not_cond:
                            tipo_filtro = 'Exclusion'

                # Se conseguiu extrair algo útil, adiciona
                if tabela or coluna:
                    resultado.append({
                        'tabela': tabela,
                        'coluna': coluna,
                        'tipo': traduzir_enum_pbir(tipo_filtro),
                        'valores': valores
                    })
            except Exception:
                continue

        return resultado

    # Mapa de categorias internas (chave i18n) para PT canonico (usado tambem
    # como fallback de exibicao quando locale = pt-BR).
    _DICT_CAT_KEYS = {
        "Tempo":         "dict.cat.time",
        "Identificador": "dict.cat.identifier",
        "Indicador":     "dict.cat.indicator",
        "Filtro":        "dict.cat.filter",
        "Página/Visual": "dict.cat.page_visual",
        "Regra Técnica": "dict.cat.technical_rule",
        "Entidade":      "dict.cat.entity",
        "Atributo":      "dict.cat.attribute",
        "Termo Geral":   "dict.cat.general",
    }

    def _localizar_categoria_termo(self, categoria_pt: str) -> str:
        """Traduz categoria do dicionario para o locale ativo. Fallback: PT."""
        key = self._DICT_CAT_KEYS.get(categoria_pt)
        if not key:
            return categoria_pt
        return i18n.t_or(key, categoria_pt)

    def _categoria_termo_dicionario(self, chave: str, fontes: List[str]) -> str:
        palavras = set(str(chave or "").split())
        fontes_set = set(fontes or [])

        if palavras & {_normalizar_palavra_termo(p) for p in TERM_TIME_WORDS}:
            return "Tempo"
        if palavras & TERM_IDENTIFIER_WORDS:
            return "Identificador"
        if palavras & {_normalizar_palavra_termo(p) for p in TERM_INDICATOR_WORDS}:
            return "Indicador"
        if "Medida" in fontes_set or "Coluna calculada" in fontes_set:
            return "Indicador"
        if "Filtro" in fontes_set:
            return "Filtro"
        if fontes_set and fontes_set <= {"Página", "Visual"}:
            return "Página/Visual"
        if "Power Query M" in fontes_set or "DAX" in fontes_set:
            if fontes_set <= {"Power Query M", "DAX"}:
                return "Regra Técnica"
        if "Tabela" in fontes_set:
            return "Entidade"
        if fontes_set & {"Coluna", "Hierarquia", "Relacionamento"}:
            return "Atributo"
        if palavras & {_normalizar_palavra_termo(p) for p in TERM_ATTRIBUTE_WORDS}:
            return "Atributo"
        return "Termo Geral"

    def gerar_dicionario_dados(self, limite: int = 50) -> List[TermoDicionario]:
        """Gera dicionario offline a partir dos metadados do PBIP."""
        if self._dicionario_cache is not None:
            return self._dicionario_cache[:limite]

        termos: Dict[str, TermoDicionario] = {}
        funcoes_tecnicas = {
            _normalizar_palavra_termo(nome)
            for nome in list(POWER_QUERY_FUNCTION_CATALOG.keys()) + list(DAX_FUNCTION_CATALOG.keys())
        }

        def peso_tipo(tipo_origem: str) -> float:
            return {
                "Medida": 6.0,
                "Coluna calculada": 5.0,
                "Tabela": 4.5,
                "Página": 4.0,
                "Visual": 4.0,
                "Coluna": 2.8,
                "Filtro": 2.6,
                "Relacionamento": 2.4,
                "Hierarquia": 2.2,
                "Grupo de consulta": 2.0,
                "Power Query M": 0.8,
                "DAX": 0.7,
            }.get(tipo_origem, 1.5)

        def registrar(
            texto: str,
            tipo_origem: str,
            nome_objeto: str,
            contexto: str = "",
            origem: str = "",
            incluir_frases: bool = True,
        ) -> None:
            if not texto:
                return
            lista_termos = _extrair_termos_texto(texto, incluir_frases=incluir_frases)
            if not lista_termos:
                return

            peso = peso_tipo(tipo_origem)
            origem_final = origem or nome_objeto or tipo_origem
            contexto_limpo = _limitar_texto(contexto or str(texto), 120)
            exemplo = f"{tipo_origem}: {nome_objeto}" if nome_objeto else tipo_origem
            if contexto_limpo and contexto_limpo != str(texto).strip():
                exemplo = f"{exemplo} - {contexto_limpo}"

            vistos_ocorrencia = set()
            for termo_display, chave in lista_termos:
                palavras = chave.split()
                if not palavras or chave in vistos_ocorrencia:
                    continue
                if (
                    len(palavras) == 1
                    and len(palavras[0]) <= 3
                    and termo_display.isupper()
                    and tipo_origem in {"Coluna", "Relacionamento", "Power Query M", "DAX"}
                ):
                    continue
                if (
                    len(palavras) > 1
                    and all(len(p) <= 3 for p in palavras)
                    and tipo_origem in {"Coluna", "Relacionamento", "Power Query M", "DAX"}
                ):
                    continue
                if chave.replace(" ", "") in funcoes_tecnicas:
                    continue
                if all(p in TERM_IDENTIFIER_WORDS for p in palavras):
                    continue
                vistos_ocorrencia.add(chave)

                item = termos.get(chave)
                if not item:
                    item = TermoDicionario(termo=termo_display)
                    termos[chave] = item

                multiplicador = 1.25 if len(palavras) > 1 else 1.0
                item.frequencia += 1
                item.score += peso * multiplicador
                if tipo_origem not in item.fontes:
                    item.fontes.append(tipo_origem)
                if exemplo and exemplo not in item.exemplos and len(item.exemplos) < 5:
                    item.exemplos.append(exemplo)
                item.ocorrencias.append(OcorrenciaTermo(
                    origem=origem_final,
                    tipo_origem=tipo_origem,
                    nome_objeto=nome_objeto,
                    contexto=contexto_limpo,
                ))

        registrar(self.nome_projeto or self.caminho_projeto.name, "Projeto", self.nome_projeto or self.caminho_projeto.name)

        for grupo in self.info_modelo.grupos_consulta or []:
            registrar(str(grupo.get("nome", "")), "Grupo de consulta", str(grupo.get("nome", "")))

        for tabela in self.tabelas:
            registrar(tabela.nome, "Tabela", tabela.nome, "Nome da tabela")
            if tabela.descricao:
                registrar(tabela.descricao, "Tabela", tabela.nome, "Descrição da tabela")

            for coluna in tabela.colunas:
                registrar(coluna.nome, "Coluna", f"{tabela.nome}.{coluna.nome}", "Nome da coluna")
                if coluna.coluna_fonte:
                    registrar(coluna.coluna_fonte, "Coluna", f"{tabela.nome}.{coluna.nome}", "Coluna de origem")
                if coluna.categoria_dados:
                    registrar(coluna.categoria_dados, "Coluna", f"{tabela.nome}.{coluna.nome}", "Categoria de dados")

            for coluna_calc in tabela.colunas_calculadas:
                registrar(coluna_calc.nome, "Coluna calculada", f"{tabela.nome}.{coluna_calc.nome}", "Nome da coluna calculada")
                for ref in _extrair_referencias_dax(coluna_calc.expressao_dax):
                    registrar(ref, "DAX", f"{tabela.nome}.{coluna_calc.nome}", "Referência usada na expressão DAX")

            for medida in tabela.medidas:
                registrar(medida.nome, "Medida", f"{tabela.nome}.{medida.nome}", "Nome da medida")
                if medida.descricao:
                    registrar(medida.descricao, "Medida", f"{tabela.nome}.{medida.nome}", "Descrição da medida")
                for ref in _extrair_referencias_dax(medida.expressao_dax):
                    registrar(ref, "DAX", f"{tabela.nome}.{medida.nome}", "Referência usada na medida DAX")

            for hierarquia in tabela.hierarquias:
                registrar(hierarquia.nome, "Hierarquia", f"{tabela.nome}.{hierarquia.nome}", "Nome da hierarquia")
                for nivel in hierarquia.niveis:
                    registrar(nivel, "Hierarquia", f"{tabela.nome}.{hierarquia.nome}", "Nível da hierarquia")

            if tabela.particao:
                if tabela.particao.grupo_consulta:
                    registrar(tabela.particao.grupo_consulta, "Grupo de consulta", tabela.nome, "Grupo de consulta da tabela")
                codigo = self._formatar_codigo_fonte(tabela.particao.codigo_fonte or "")
                if codigo and not _codigo_fonte_eh_dax(codigo):
                    regra = analisar_power_query_m(codigo)
                    for valores in regra.observacoes.values():
                        for valor in valores:
                            registrar(valor, "Power Query M", tabela.nome, "Comentário BI_DOC")
                    for linha in regra.linhas_regra:
                        registrar(linha.etapa, "Power Query M", tabela.nome, "Etapa Power Query")
                        registrar(linha.regra, "Power Query M", tabela.nome, "Regra Power Query")
                    for texto in _extrair_strings_m(codigo):
                        registrar(texto, "Power Query M", tabela.nome, "Texto encontrado no código M")
                elif codigo:
                    for ref in _extrair_referencias_dax(codigo):
                        registrar(ref, "DAX", tabela.nome, "Referência usada em partição DAX")

        for rel in self.relacionamentos:
            if (
                rel.tabela_origem.startswith(("LocalDateTable_", "DateTableTemplate_"))
                or rel.tabela_destino.startswith(("LocalDateTable_", "DateTableTemplate_"))
            ):
                continue
            contexto_rel = f"{rel.tabela_origem}.{rel.coluna_origem} -> {rel.tabela_destino}.{rel.coluna_destino}"
            registrar(rel.tabela_origem, "Relacionamento", contexto_rel, "Tabela de origem")
            registrar(rel.coluna_origem, "Relacionamento", contexto_rel, "Coluna de origem")
            registrar(rel.tabela_destino, "Relacionamento", contexto_rel, "Tabela de destino")
            registrar(rel.coluna_destino, "Relacionamento", contexto_rel, "Coluna de destino")

        for pagina in self.paginas:
            registrar(pagina.nome_exibicao, "Página", pagina.nome_exibicao, "Nome da página")

            filtros = self._interpretar_filtros_pagina(pagina.filtros) if pagina.filtros else []
            for filtro in filtros:
                contexto = f"Filtro da página {pagina.nome_exibicao}"
                registrar(filtro.get("tabela", ""), "Filtro", pagina.nome_exibicao, contexto)
                registrar(filtro.get("coluna", ""), "Filtro", pagina.nome_exibicao, contexto)
                for valor in filtro.get("valores", [])[:4]:
                    registrar(valor, "Filtro", pagina.nome_exibicao, contexto)

            for visual in pagina.visuais:
                if visual.titulo and visual.titulo != "-":
                    registrar(visual.titulo, "Visual", f"{pagina.nome_exibicao} / {visual.tipo}", "Título do visual")
                for campo in visual.campos:
                    registrar(campo, "Visual", f"{pagina.nome_exibicao} / {visual.tipo}", "Campo usado no visual")

        dicionario = list(termos.values())
        for chave, termo in termos.items():
            termo.categoria = self._categoria_termo_dicionario(chave, termo.fontes)
            termo.fontes = termo.fontes[:6]
            termo.exemplos = termo.exemplos[:4]

        dicionario.sort(key=lambda t: (-t.score, -t.frequencia, t.termo.lower()))
        dicionario = self._deduplicar_termos_dicionario(dicionario)
        self._dicionario_cache = dicionario
        return dicionario[:limite]

    def _coletar_funcoes_dax_usadas(self) -> List[str]:
        """Varre todas as expressoes DAX do projeto e retorna a lista de funcoes
        unicas usadas (medidas + colunas calculadas + tabelas DAX)."""
        funcoes_set: set = set()
        for tabela in self.tabelas:
            for medida in tabela.medidas:
                if medida.expressao_dax:
                    leitura = analisar_dax(medida.expressao_dax)
                    funcoes_set.update(leitura.funcoes)
            for col_calc in tabela.colunas_calculadas:
                if col_calc.expressao_dax:
                    leitura = analisar_dax(col_calc.expressao_dax)
                    funcoes_set.update(leitura.funcoes)
            # Tabelas geradas via DAX (fonte_codigo DAX)
            if tabela.particao and tabela.particao.codigo_fonte:
                if _codigo_fonte_eh_dax(tabela.particao.codigo_fonte):
                    leitura = analisar_dax(tabela.particao.codigo_fonte)
                    funcoes_set.update(leitura.funcoes)
        return sorted(funcoes_set)

    def _gerar_glossario_dax_markdown(self, md: List[str]) -> None:
        """Insere uma secao 'Glossario DAX' com todas as funcoes usadas
        agrupadas por categoria. Substitui a repeticao de descricoes em
        cada medida individual."""
        funcoes_usadas = self._coletar_funcoes_dax_usadas()
        if not funcoes_usadas:
            return

        # Agrupa por categoria conforme catalogo
        por_categoria: Dict[str, List[CatalogoFuncao]] = {}
        for nome in funcoes_usadas:
            entrada = DAX_FUNCTION_CATALOG.get(nome)
            if not entrada:
                continue
            por_categoria.setdefault(entrada.categoria, []).append(entrada)

        if not por_categoria:
            return

        def _dax_cat_label(categoria: str) -> str:
            # i18n override (en) com fallback ao DAX_CATEGORY_LABELS (pt)
            return i18n.t_or(
                f"dax.cat.{categoria}",
                DAX_CATEGORY_LABELS.get(categoria, categoria.title()),
            )

        def _dax_desc(entrada: "CatalogoFuncao") -> str:
            key_root = f"dax.{entrada.nome.lower()}"
            return i18n.t_or(f"{key_root}.desc", entrada.descricao)

        def _dax_business(entrada: "CatalogoFuncao") -> str:
            key_root = f"dax.{entrada.nome.lower()}"
            return i18n.t_or(f"{key_root}.business", entrada.leitura_negocio)

        md.append(f"## {_t('doc.section.glossary_dax')}")
        md.append("")
        md.append(f"> {_t('glossary_dax.intro')}")
        md.append("")
        md.append(
            f"| {_t('glossary_dax.col.function')} | {_t('glossary_dax.col.category')} | "
            f"{_t('glossary_dax.col.description')} | {_t('glossary_dax.col.business_reading')} |"
        )
        md.append("|---|---|---|---|")
        for categoria in DAX_CATEGORY_ORDER:
            entradas = por_categoria.get(categoria, [])
            if not entradas:
                continue
            label = _dax_cat_label(categoria)
            for entrada in sorted(entradas, key=lambda e: e.nome):
                md.append(
                    f"| `{entrada.nome}` "
                    f"| {label} "
                    f"| {_dax_desc(entrada)} "
                    f"| {_dax_business(entrada)} |"
                )
        # Categorias fora da ordem padrao (caso existam)
        for categoria, entradas in por_categoria.items():
            if categoria in DAX_CATEGORY_ORDER:
                continue
            label = _dax_cat_label(categoria)
            for entrada in sorted(entradas, key=lambda e: e.nome):
                md.append(
                    f"| `{entrada.nome}` "
                    f"| {label} "
                    f"| {_dax_desc(entrada)} "
                    f"| {_dax_business(entrada)} |"
                )
        md.append("")
        md.append("---")
        md.append("")

    def _deduplicar_termos_dicionario(
        self, termos: List["TermoDicionario"]
    ) -> List["TermoDicionario"]:
        """Remove termos redundantes (substrings/superstrings com freq similar).

        Ex: 'GLA EXIGENCIAS' (14), 'TS GLA EXIGENCIAS' (14), 'TS GLA' (12),
        'GLA' (12) - mantem so o mais longo, que preserva mais contexto.
        Termos com frequencias muito diferentes (ratio < 0.75) sao preservados
        ambos (ex: 'Receita' 40 e 'Receita Realizada' 8 sao termos distintos).
        """
        if len(termos) <= 1:
            return termos

        # Ordena por comprimento DESC (preferir nome mais longo), depois score DESC.
        # Isso garante que processamos primeiro a forma mais informativa.
        candidatos = sorted(
            termos,
            key=lambda t: (-len(t.termo), -t.score, -t.frequencia),
        )

        mantidos: List["TermoDicionario"] = []
        for termo in candidatos:
            chave_norm = termo.termo.lower().strip()
            if not chave_norm:
                continue

            eh_redundante = False
            for ja in mantidos:
                ja_norm = ja.termo.lower().strip()
                # Considera "mesma familia" se um termo eh substring do outro
                # (limitado a palavras, nao caracteres dentro de palavra).
                eh_substring = (
                    f" {chave_norm} " in f" {ja_norm} "
                    or chave_norm == ja_norm
                    or ja_norm.startswith(chave_norm + " ")
                    or ja_norm.endswith(" " + chave_norm)
                )
                if not eh_substring:
                    continue
                # Frequencia similar (>=75% do maximo) -> redundante.
                max_freq = max(ja.frequencia, termo.frequencia)
                min_freq = min(ja.frequencia, termo.frequencia)
                if max_freq > 0 and (min_freq / max_freq) >= 0.75:
                    eh_redundante = True
                    break

            if not eh_redundante:
                mantidos.append(termo)

        # Re-ordena pelo criterio original (score/frequencia).
        mantidos.sort(key=lambda t: (-t.score, -t.frequencia, t.termo.lower()))
        return mantidos

    def _limpar_nome_mermaid(self, nome: str) -> str:
        import re
        resultado = re.sub(r'[^a-zA-Z0-9_]', '_', str(nome))
        if resultado and resultado[0].isdigit():
            resultado = "T_" + resultado
        return resultado

    def _tabelas_fato_mermaid(self) -> List[str]:
        tabelas_fato = []
        for tabela in self.tabelas:
            if 'LocalDateTable' in tabela.nome or 'DateTableTemplate' in tabela.nome:
                continue
            if tabela.eh_fato:
                tabelas_fato.append(self._limpar_nome_mermaid(tabela.nome))
        return sorted(set(tabelas_fato))

    def _gerar_codigo_mermaid(self) -> str:
        """Gera o código do diagrama ER em formato Mermaid"""
        LIMITE_COLUNAS_ER = 12
        mermaid = ["erDiagram"]

        # Adiciona definição das tabelas e colunas no diagrama
        for tabela in self.tabelas:
            # Omitir tabelas de calendário automáticas
            if 'LocalDateTable' in tabela.nome or 'DateTableTemplate' in tabela.nome:
                continue
            nome_tab = self._limpar_nome_mermaid(tabela.nome)
            mermaid.append(f"    {nome_tab} {{")
            # Limite de colunas no diagrama (top N para nao ficar gigantesco).
            # Se houver mais, sinaliza com pseudo-linha; lista completa fica
            # no detalhamento de cada tabela mais abaixo no documento.
            for col in tabela.colunas[:LIMITE_COLUNAS_ER]:
                tipo = self._limpar_nome_mermaid(col.tipo_dado or "string")
                nome_col = self._limpar_nome_mermaid(col.nome)
                mermaid.append(f"        {tipo} {nome_col}")
            restantes = len(tabela.colunas) - LIMITE_COLUNAS_ER
            if restantes > 0:
                sufixo = "coluna" if restantes == 1 else "colunas"
                mermaid.append(f"        string mais_{restantes}_{sufixo}_omitidas")
            mermaid.append("    }")

        # Adiciona as ligações (relacionamentos)
        for rel in self.relacionamentos:
            if 'LocalDateTable' in rel.tabela_destino or 'DateTableTemplate' in rel.tabela_destino:
                continue
            tabela_origem = self._limpar_nome_mermaid(rel.tabela_origem)
            tabela_destino = self._limpar_nome_mermaid(rel.tabela_destino)
            coluna_origem = rel.coluna_origem.replace("'", "")

            # }|--|| = muitos-para-um (bidirecional)  |  }o--|| = muitos-para-um (unidirecional)
            tipo_rel = "}|--||" if rel.filtro_bidirecional else "}o--||"
            mermaid.append("    " + tabela_origem + " " + tipo_rel + " " + tabela_destino + ' : "' + coluna_origem + '"')

        return "\n".join(mermaid)

    def gerar_resumo_estruturado(
        self,
        outputs: Optional[Dict[str, str]] = None,
        warnings: Optional[List[str]] = None
    ) -> Dict:
        """Retorna um resumo serializavel para a interface Tauri."""
        total_colunas = sum(len(t.colunas) for t in self.tabelas)
        total_medidas = sum(len(t.medidas) for t in self.tabelas)
        total_calc = sum(len(t.colunas_calculadas) for t in self.tabelas)

        def contar_filtros(pagina: InfoPagina) -> int:
            if not pagina.filtros:
                return 0
            return len(self._interpretar_filtros_pagina(pagina.filtros))

        def resumo_power_query(tabela: InfoTabela) -> Dict[str, int]:
            if not tabela.particao or not tabela.particao.codigo_fonte:
                return {}
            codigo = tabela.particao.codigo_fonte
            if _codigo_fonte_eh_dax(codigo):
                return {}
            regra = analisar_power_query_m(self._formatar_codigo_fonte(codigo))
            return {
                "steps": len(regra.etapas),
                "filters": len(regra.filtros_aplicados),
                "joins": len(regra.integracoes),
                "documentedNotes": sum(len(v) for v in regra.observacoes.values()),
            }

        dicionario = self.gerar_dicionario_dados()
        fontes_dicionario = sorted({
            fonte
            for termo in dicionario
            for fonte in termo.fontes
        })

        return {
            "ok": True,
            "projectName": self.nome_projeto or self.caminho_projeto.name,
            "layout": self.layout or "",
            "counts": {
                "tables": len(self.tabelas),
                "relationships": len(self.relacionamentos),
                "pages": len(self.paginas),
                "columns": total_colunas,
                "measures": total_medidas,
                "calculatedColumns": total_calc,
                "visuals": self.total_visuais,
            },
            "tables": [
                {
                    "name": tabela.nome,
                    "hidden": tabela.esta_oculta,
                    "columns": len(tabela.colunas),
                    "measures": len(tabela.medidas),
                    "calculatedColumns": len(tabela.colunas_calculadas),
                    "sourceGroup": tabela.particao.grupo_consulta if tabela.particao else None,
                    "powerQuery": resumo_power_query(tabela),
                }
                for tabela in self.tabelas
            ],
            "relationships": [
                {
                    "id": rel.id,
                    "fromTable": rel.tabela_origem,
                    "fromColumn": rel.coluna_origem,
                    "toTable": rel.tabela_destino,
                    "toColumn": rel.coluna_destino,
                    "bidirectional": rel.filtro_bidirecional,
                    "active": rel.esta_ativo,
                }
                for rel in self.relacionamentos
            ],
            "pages": [
                {
                    "id": pagina.id,
                    "name": pagina.nome_exibicao,
                    "type": pagina.tipo,
                    "width": pagina.largura,
                    "height": pagina.altura,
                    "visuals": len(pagina.visuais),
                    "filters": contar_filtros(pagina),
                }
                for pagina in self.paginas
            ],
            "dataDictionary": {
                "termsCount": len(dicionario),
                "topTerms": [
                    {
                        "term": termo.termo,
                        "frequency": termo.frequencia,
                        "category": termo.categoria,
                        "sources": termo.fontes[:5],
                    }
                    for termo in dicionario[:10]
                ],
                "sourcesUsed": fontes_dicionario,
            },
            "warnings": list(dict.fromkeys(warnings or [])),
            "outputs": outputs or {},
        }

    def aplicar_branding(self, branding: Optional[Dict] = None) -> None:
        """Aplica titulo, logo e cores customizadas para as exportacoes."""
        branding = branding or {}

        def valor_texto(*chaves: str) -> str:
            for chave in chaves:
                valor = branding.get(chave)
                if valor is not None:
                    return str(valor).strip()
            return ""

        def cor_hex(chave: str, padrao: str) -> str:
            valor = valor_texto(chave) or padrao
            if not re.fullmatch(r"#[0-9a-fA-F]{6}", valor):
                raise ValueError(f"Cor invalida em {chave}: use o formato #RRGGBB.")
            return valor.upper()

        titulo = valor_texto("documentTitle", "document_title") or _t("cover.document_title_default")
        logo_raw = valor_texto("logoPath", "logo_path")
        logo_path = None
        if logo_raw:
            logo_path = Path(logo_raw).expanduser()
            if not logo_path.is_file():
                raise ValueError(f"Logo nao encontrado: {logo_raw}")
            if logo_path.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
                raise ValueError("Logo deve ser uma imagem .png, .jpg ou .jpeg.")

        self.branding = BrandingConfig(
            document_title=titulo,
            logo_path=logo_path,
            primary_color=cor_hex("primaryColor", "#003D6B"),
            secondary_color=cor_hex("secondaryColor", "#006DAA"),
            light_color=cor_hex("lightColor", "#D6E8F5"),
        )
        self._md_cache = None

    def _hex_sem_hash(self, cor: str) -> str:
        return cor.lstrip("#").upper()

    def _obter_logo_path(self) -> Optional[Path]:
        """Retorna apenas a logo customizada informada para a documentacao."""
        if self.branding.logo_path and self.branding.logo_path.is_file():
            return self.branding.logo_path
        return None

    def _obter_logo_data_uri(self) -> str:
        """Retorna a logo como data URI para HTML offline."""
        logo_path = self._obter_logo_path()
        if not logo_path:
            return ""
        encoded = base64.b64encode(logo_path.read_bytes()).decode("ascii")
        mime = "image/jpeg" if logo_path.suffix.lower() in {".jpg", ".jpeg"} else "image/png"
        return f"data:{mime};base64,{encoded}"

    def _relacionamentos_usuario(self) -> List:
        """Relacionamentos relevantes para o usuario (exclui hierarquias automaticas de data: LocalDateTable_/DateTableTemplate_)."""
        TECNICAS = ('LocalDateTable_', 'DateTableTemplate_')
        return [
            r for r in self.relacionamentos
            if not r.tabela_destino.startswith(TECNICAS)
            and not r.tabela_origem.startswith(TECNICAS)
        ]

    @staticmethod
    def _classificar_tabela(nome: str) -> str:
        """Classifica tabela como 'fato', 'dimensao' ou 'outra' a partir do prefixo do nome."""
        n = (nome or "").lower()
        if n.startswith(("fat_", "fct_", "fact_")):
            return "fato"
        if n.startswith(("dim_", "d_")):
            return "dimensao"
        return "outra"

    def _resumo_modelo_dimensional(self) -> str:
        """Texto resumido do modelo: 'N fatos × M dimensoes + K auxiliares'.

        Retorna string vazia quando nao detecta convencao fat_/dim_ (modelo nao
        e star schema explicito; sumario seria enganoso).
        """
        fatos = sum(1 for tab in self.tabelas if self._classificar_tabela(tab.nome) == "fato")
        dims = sum(1 for tab in self.tabelas if self._classificar_tabela(tab.nome) == "dimensao")
        outras = len(self.tabelas) - fatos - dims

        if fatos == 0 and dims == 0:
            return ""

        def plural(n: int, key_one: str, key_many: str) -> str:
            return f"**{n}** {_t(key_many) if n != 1 else _t(key_one)}"

        partes: List[str] = []
        if fatos:
            partes.append(plural(fatos, "overview.dim.fact_one", "overview.dim.fact_many"))
        if dims:
            partes.append(plural(dims, "overview.dim.dim_one", "overview.dim.dim_many"))
        texto = " × ".join(partes)
        if outras:
            texto += f" + {plural(outras, 'overview.dim.aux_one', 'overview.dim.aux_many')}"
        return texto

    def gerar_documentacao(self) -> str:
        """
        Gera a documentação em Markdown.
        O resultado é cacheado internamente; chamadas subsequentes
        retornam o cache até que extrair_informacoes() seja invocado novamente.

        Returns:
            String com o conteúdo Markdown
        """
        if self._md_cache is not None:
            return self._md_cache

        md = []

        # Contadores para estatísticas
        total_colunas = sum(len(t.colunas) for t in self.tabelas)
        total_medidas = sum(len(t.medidas) for t in self.tabelas)
        total_calc = sum(len(t.colunas_calculadas) for t in self.tabelas)
        dicionario_dados = self.gerar_dicionario_dados()

        def escape_md_table(valor: str) -> str:
            return (
                str(valor or "")
                .replace("\\", "\\\\")
                .replace("|", "\\|")
                .replace("\n", "<br>")
                .strip()
            )

        # ========================================================================
        # CAPA (Página 1)
        # ========================================================================
        titulo_doc = html.escape(self.branding.document_title)
        nome_projeto_html = html.escape(str(self.nome_projeto))
        cor_primaria = self.branding.primary_color
        cor_secundaria = self.branding.secondary_color
        cor_clara = self.branding.light_color
        fact_tables = self._tabelas_fato_mermaid()
        fact_tables_js = json.dumps(fact_tables, ensure_ascii=True)
        logo_data_uri = self._obter_logo_data_uri()
        logo_img = (
            f"<img src='{logo_data_uri}' width='80' height='80' "
            "style='border-radius: 16px; object-fit: contain;'/>"
            if logo_data_uri else ""
        )
        data_criacao = datetime.now().strftime('%d/%m/%Y')
        md.append(
            '<div class="cover-page" '
            f'style="border-top: 6px solid {cor_primaria}; background: linear-gradient(180deg, {cor_clara}, #ffffff 42%);">'
        )
        if logo_img:
            md.append(f'<div class="cover-logo">{logo_img}</div>')
        md.append(
            f'<h1 class="cover-title" style="color: {cor_primaria}; '
            f'border-bottom: 3px solid {cor_secundaria};">{titulo_doc}</h1>'
        )
        md.append(f'<p class="cover-project" style="color: {cor_primaria};">{nome_projeto_html}</p>')
        md.append(f'<p class="cover-date">{_t("cover.created_on")}: {data_criacao}</p>')
        md.append('</div>')
        md.append('')
        md.append('<div class="page-break"></div>')
        md.append('')

        # ========================================================================
        # ÍNDICE (Página 2)
        # ========================================================================
        md.append(f"## {_t('doc.toc.title')}")
        md.append(f"")
        import unicodedata
        import re
        def gerar_slug(texto):
            # Função para emular o slugify nativo do Python-Markdown
            texto = unicodedata.normalize('NFKD', texto).encode('ascii', 'ignore').decode('ascii')
            texto = re.sub(r'[^\w\s-]', '', texto).strip().lower()
            return re.sub(r'[-\s]+', '-', texto)

        # Indice completo em HTML para controle total da numeracao.
        # Gera slugs dinamicamente a partir do label traduzido \u2014 funciona em
        # qualquer idioma sem catalogo separado de anchors.
        def _toc_li(key: str) -> str:
            label = _t(key)
            return f'<li><a href="#{gerar_slug(label)}">{label}</a></li>'

        toc_items = []
        toc_items.append('<ol class="toc-list">')
        toc_items.append(_toc_li("doc.toc.overview"))
        if dicionario_dados:
            toc_items.append(_toc_li("doc.toc.dictionary"))
        toc_items.append(_toc_li("doc.toc.pages"))
        toc_items.append(_toc_li("doc.toc.model"))
        toc_items.append(_toc_li("doc.toc.tables_summary"))
        det_label = _t("doc.toc.tables_detail")
        toc_items.append(f'<li><a href="#{gerar_slug(det_label)}">{det_label}</a>')
        toc_items.append('<div style="margin-left: 1em; margin-top: 0.3em;">')
        for idx, tabela in enumerate(self.tabelas, 1):
            slug = gerar_slug(f"{idx}. {tabela.nome}")
            letra = chr(ord('a') + (idx - 1) % 26)
            toc_items.append(f'{letra}) <a href="#{slug}">{tabela.nome}</a><br/>')
        toc_items.append('</div>')
        toc_items.append('</li>')
        toc_items.append(_toc_li("doc.toc.relationships"))
        if self.visuais_personalizados:
            toc_items.append(_toc_li("doc.toc.custom_visuals"))
        toc_items.append(_toc_li("doc.toc.image_resources"))
        toc_items.append('</ol>')
        md.append('\n'.join(toc_items))
        md.append(f"")

        md.append('<div class="page-break"></div>')
        md.append(f"")

        # ========================================================================
        # VISÃO GERAL (Página 3+)
        # ========================================================================
        rel_validos_visao = self._relacionamentos_usuario()
        md.append(f"## {_t('doc.section.overview')}")
        md.append(f"")
        md.append(
            f"| {_t('overview.col.tables')} | {_t('overview.col.measures')} | "
            f"{_t('overview.col.columns')} | {_t('overview.col.calc_columns')} | "
            f"{_t('overview.col.relationships')} | {_t('overview.col.pages')} |"
        )
        md.append(f"|:----------:|:----------:|:----------:|:-------------:|:-----------------:|:----------:|")
        md.append(f"| **{len(self.tabelas)}** | **{total_medidas}** | **{total_colunas}** | **{total_calc}** | **{len(rel_validos_visao)}** | **{len(self.paginas)}** |")
        md.append(f"")
        resumo_dim = self._resumo_modelo_dimensional()
        if resumo_dim:
            md.append(f"{_t('overview.dimensional_label')}: {resumo_dim}.")
            md.append(f"")
        md.append(f"---")
        md.append(f"")

        # ========================================================================
        # PÁGINAS DO RELATÓRIO
        # ========================================================================
        if dicionario_dados:
            md.append(f"## {_t('doc.section.dictionary')}")
            md.append("")
            md.append(f"> {_t('dict.intro')}")
            md.append("")
            md.append(
                f"| {_t('dict.col.term')} | {_t('dict.col.frequency')} | "
                f"{_t('dict.col.category')} | {_t('dict.col.sources')} | {_t('dict.col.examples')} |"
            )
            md.append("|---|---:|---|---|---|")
            for termo in dicionario_dados:
                onde_aparece = ", ".join(termo.fontes[:5]) or "-"
                exemplos = "; ".join(termo.exemplos[:3]) or "-"
                md.append(
                    f"| {escape_md_table(termo.termo)} "
                    f"| {termo.frequencia} "
                    f"| {escape_md_table(self._localizar_categoria_termo(termo.categoria))} "
                    f"| {escape_md_table(onde_aparece)} "
                    f"| {escape_md_table(exemplos)} |"
                )
            md.append("")
            md.append("---")
            md.append("")

        # Glossario DAX consolidado (uma vez para todo o documento).
        self._gerar_glossario_dax_markdown(md)

        md.append(f"## {_t('doc.section.pages')}")
        md.append(f"")
        md.append(_t("pages.total", n=len(self.paginas)))
        md.append(f"")
        if self.paginas:
            md.append(f"| # | Nome | Tipo | Dimensões | Filtros |")
            md.append(f"|---|------|------|-----------|---------|")

            for i, pagina in enumerate(self.paginas, 1):
                dimensoes = f"{pagina.largura} x {pagina.altura}"
                filtros_interpretados = self._interpretar_filtros_pagina(pagina.filtros) if pagina.filtros else []
                qtd_filtros = len(filtros_interpretados)
                filtro_badge = f"{qtd_filtros}" if qtd_filtros > 0 else "-"
                md.append(f"| {i} | {pagina.nome_exibicao} | {pagina.tipo} | {dimensoes} | {filtro_badge} |")

            md.append(f"")

            # Detalhamento dos filtros de página
            paginas_com_filtros = [
                (p, self._interpretar_filtros_pagina(p.filtros))
                for p in self.paginas if p.filtros
            ]
            paginas_com_filtros = [(p, f) for p, f in paginas_com_filtros if f]

            if paginas_com_filtros:
                md.append(f"### 🔍 Filtros de Página")
                md.append(f"")

                for pagina, filtros in paginas_com_filtros:
                    md.append(f"**{pagina.nome_exibicao}**")
                    md.append(f"")
                    md.append(f"| Tabela | Coluna | Tipo | Valores |")
                    md.append(f"|--------|--------|------|---------|")

                    for f in filtros:
                        valores_str = ', '.join(f['valores']) if f['valores'] else '-'
                        md.append(f"| {f['tabela']} | {f['coluna']} | {f['tipo']} | {valores_str} |")

                    md.append(f"")
        else:
            md.append(f"> {_t('pages.empty_thin_report')}")
            md.append(f"")

        md.append(f"---")
        md.append(f"")

        # ========================================================================
        # MODELO DE DADOS
        # ========================================================================
        md.append(f"## {_t('doc.section.model')}")
        md.append(f"")


        # Diagrama de Relacionamentos (Mermaid)
        if self.tabelas or self.relacionamentos:
            md.append(f"### {_t('doc.section.er_diagram')}")
            md.append(f"")
            md.append(f"```mermaid")
            md.append(self._gerar_codigo_mermaid())

            md.append("```")
            md.append("")
            md.append(_t("rel.legend"))
            md.append("")
        # Relacionamentos (Tabela)
        if self.relacionamentos:
            rel_validos = self._relacionamentos_usuario()

            if rel_validos:
                md.append(f"")
                md.append(f"### {_t('doc.section.relationships_list')}")
                md.append(f"")
                md.append(_t("rel.total_label", n=len(rel_validos)))
                md.append(f"")
                md.append(
                    f"| {_t('rel.col.from')} | {_t('rel.col.arrow')} | {_t('rel.col.to')} | "
                    f"{_t('rel.col.bidirectional')} | {_t('rel.col.active')} |"
                )
                md.append(f"|--------|---|---------|--------------|-------|")

                for rel in rel_validos:
                    origem = f"{rel.tabela_origem}.{rel.coluna_origem}"
                    destino = f"{rel.tabela_destino}.{rel.coluna_destino}"
                    bidirecional = _t("table.yes") if rel.filtro_bidirecional else _t("table.no")
                    ativo = _t("table.yes") if rel.esta_ativo else _t("table.no")
                    md.append(f"| {origem} | {_t('rel.col.arrow')} | {destino} | {bidirecional} | {ativo} |")

                md.append(f"")

        # Grupos de Consulta
        if self.info_modelo.grupos_consulta:
            md.append(f"### {_t('doc.section.query_groups')}")
            md.append(f"")
            md.append(f"| Nome | Ordem |")
            md.append(f"|------|-------|")
            for g in sorted(self.info_modelo.grupos_consulta, key=lambda x: x['ordem']):
                md.append(f"| {g['nome']} | {g['ordem']} |")
            md.append(f"")

        md.append(f"---")
        md.append(f"")

        # ========================================================================
        # TABELAS DETALHADAS
        # ========================================================================
        # RESUMO DAS TABELAS
        # ========================================================================
        md.append(f"---")
        md.append(f"")
        md.append(f"## {_t('doc.section.tables_summary')}")
        md.append(f"")
        md.append(_t("tables_summary.intro", n=len(self.tabelas)))
        md.append(f"")
        md.append(
            f"| {_t('tables_summary.col.idx')} | {_t('tables_summary.col.table')} | "
            f"{_t('tables_summary.col.columns')} | {_t('tables_summary.col.measures')} | "
            f"{_t('tables_summary.col.calc_columns')} | {_t('tables_summary.col.source')} |"
        )
        md.append(f"|---|--------|---------|---------|---------------|-------|")

        for idx, tabela in enumerate(self.tabelas, 1):
            num_colunas = len(tabela.colunas)
            num_medidas = len(tabela.medidas)
            num_calc = len(tabela.colunas_calculadas)
            fonte = "DAX" if tabela.particao and _codigo_fonte_eh_dax(tabela.particao.codigo_fonte) else "Importação"
            if tabela.particao and tabela.particao.grupo_consulta:
                fonte = tabela.particao.grupo_consulta
            md.append(f"| {idx} | **{tabela.nome}** | {num_colunas} | {num_medidas} | {num_calc} | {fonte} |")

        md.append(f"")

        # Totais
        total_colunas = sum(len(tab.colunas) for tab in self.tabelas)
        total_medidas = sum(len(tab.medidas) for tab in self.tabelas)
        total_calc = sum(len(tab.colunas_calculadas) for tab in self.tabelas)
        md.append(_t("tables_summary.totals", cols=total_colunas, meds=total_medidas, calc=total_calc))
        md.append(f"")

        # Tabelas
        if self.tabelas:
            md.append(f"---")
            md.append(f"")
            md.append(f"## {_t('doc.section.tables_in_model')}")
            md.append(f"")
            md.append(_t("tables_summary.doc_count", n=len(self.tabelas)))
            md.append(f"")
        else:
            md.append(f"---")
            md.append(f"")
            md.append(f"## {_t('doc.section.tables')}")
            md.append(f"")
            md.append(_t("tables_summary.empty"))
            md.append(f"")

        for idx, tabela in enumerate(self.tabelas, 1):
            # Separador entre tabelas (exceto antes da primeira)
            if idx > 1:
                md.append(f"---")
                md.append(f"")

            md.append(f"### {idx}. {tabela.nome}")
            md.append(f"")

            if hasattr(tabela, 'descricao') and tabela.descricao:
                md.append(f"*{tabela.descricao}*")
                md.append(f"")

            # Card de metadados com badges
            status_badge = _t("table.status.hidden") if tabela.esta_oculta else _t("table.status.visible")
            refresh_badge = _t("table.refresh.no") if tabela.excluida_refresh else _t("table.refresh.yes")

            # Determina tipo de fonte
            fonte_tipo = _t("table.source.import")
            if tabela.particao:
                if tabela.particao.grupo_consulta:
                    fonte_tipo = f"🗃️ {tabela.particao.grupo_consulta}"
                elif _codigo_fonte_eh_dax(tabela.particao.codigo_fonte) if tabela.particao.codigo_fonte else False:
                    fonte_tipo = _t("table.source.dax")
                elif "Oracle" in tabela.particao.codigo_fonte if tabela.particao.codigo_fonte else False:
                    fonte_tipo = _t("table.source.oracle")

            md.append(
                f"| {_t('table.col.status')} | {_t('table.col.refresh')} | "
                f"{_t('table.col.columns')} | {_t('table.col.measures')} | {_t('table.col.source')} |"
            )
            md.append(f"|:------:|:-----------:|:-------:|:-------:|:-----:|")
            md.append(f"| {status_badge} | {refresh_badge} | {len(tabela.colunas)} | {len(tabela.medidas)} | {fonte_tipo} |")
            md.append(f"")

            # Colunas
            if tabela.colunas:
                md.append(f"#### {_t('table.section.columns')}")
                md.append(f"")
                md.append(
                    f"| {_t('table.col.name')} | {_t('table.col.type')} | "
                    f"{_t('table.col.summarization')} | {_t('table.col.hidden')} |"
                )
                md.append(f"|:-----|:----:|:-----------:|:------:|")

                for coluna in tabela.colunas:
                    oculta = "🔴" if coluna.esta_oculta else "⚪"
                    md.append(f"| `{coluna.nome}` | `{coluna.tipo_dado}` | {coluna.sumarizacao} | {oculta} |")

                md.append(f"")

            # Colunas Calculadas - Resumo
            if tabela.colunas_calculadas:
                md.append(f"#### {_t('table.section.calc_columns_summary')}")
                md.append(f"")
                md.append(
                    f"| {_t('table.calc.col.name')} | {_t('table.calc.col.type')} | {_t('table.calc.col.format')} |"
                )
                md.append(f"|------|------|---------|")

                for coluna in tabela.colunas_calculadas:
                    tipo = coluna.tipo_dado or "string"
                    formato = coluna.formato or "-"
                    md.append(f"| {coluna.nome} | {tipo} | {formato} |")

                md.append(f"")

                # Colunas Calculadas - Código DAX Completo
                md.append(f"#### {_t('table.section.calc_columns_dax')}")
                md.append(f"")

                for coluna in tabela.colunas_calculadas:
                    md.append(f"##### {coluna.nome}")
                    md.append(f"")
                    info_parts = []
                    if coluna.tipo_dado:
                        info_parts.append(f"**Tipo**: `{coluna.tipo_dado}`")
                    if coluna.formato:
                        info_parts.append(f"**Formato**: `{coluna.formato}`")
                    if info_parts:
                        md.append(" | ".join(info_parts))
                        md.append(f"")
                    leitura_col = analisar_dax(coluna.expressao_dax)
                    adicionar_leitura_dax_markdown(md, leitura_col)
                    adicionar_linhagem_dax_markdown(
                        md,
                        leitura_col,
                        medidas_conhecidas=self.medidas_index,
                    )
                    md.append(f"```dax")
                    # Verifica se a expressão DAX está vazia ou None
                    if coluna.expressao_dax and str(coluna.expressao_dax).strip():
                        md.append(str(coluna.expressao_dax).strip())
                    else:
                        md.append("// [AVISO] Expressão DAX não capturada durante o parsing")
                    md.append(f"```")
                    md.append(f"")

                md.append(f"")

            # Medidas - Resumo
            if tabela.medidas:
                md.append(f"#### {_t('table.section.measures_summary')}")
                md.append(f"")
                md.append(f"| {_t('table.measure.col.name')} | {_t('table.measure.col.type')} |")
                md.append(f"|------|------|")

                for medida in tabela.medidas:
                    tipo_medida = inferir_tipo_medida(
                        medida.nome,
                        medida.expressao_dax or "",
                        analisar_dax(medida.expressao_dax or ""),
                    )
                    md.append(f"| {medida.nome} | {tipo_medida} |")

                md.append(f"")

                # Medidas - Código DAX Completo
                md.append(f"#### {_t('table.section.measures_dax')}")
                md.append(f"")

                for medida in tabela.medidas:
                    md.append(f"##### {medida.nome}")
                    md.append(f"")

                    if hasattr(medida, 'descricao') and medida.descricao:
                        md.append(f"*{medida.descricao}*")
                        md.append(f"")

                    leitura_medida = analisar_dax(medida.expressao_dax)
                    adicionar_leitura_dax_markdown(md, leitura_medida)
                    adicionar_linhagem_dax_markdown(
                        md,
                        leitura_medida,
                        medidas_conhecidas=self.medidas_index,
                    )
                    md.append(f"```dax")
                    # Verifica se a expressão DAX está vazia ou None
                    if medida.expressao_dax and str(medida.expressao_dax).strip():
                        md.append(str(medida.expressao_dax).strip())
                    else:
                        md.append("// [AVISO] Expressão DAX não capturada durante o parsing")
                    md.append(f"```")
                    md.append(f"")

                    # Formato dinâmico (se existir)
                    if medida.formato_dinamico and str(medida.formato_dinamico).strip():
                        md.append(f"**Formato Dinâmico** (expressão DAX que controla a máscara de exibição):")
                        md.append(f"")
                        md.append(f"```dax")
                        md.append(str(medida.formato_dinamico).strip())
                        md.append(f"```")
                        md.append(f"")

                md.append(f"")

            # Partição (regra inferida antes do código bruto)
            if tabela.particao and tabela.particao.codigo_fonte:
                codigo_formatado = self._formatar_codigo_fonte(tabela.particao.codigo_fonte)
                is_dax = _codigo_fonte_eh_dax(tabela.particao.codigo_fonte)
                if not is_dax:
                    adicionar_regra_power_query_markdown(md, analisar_power_query_m(codigo_formatado))

            # Hierarquias
            if tabela.hierarquias:
                md.append(f"#### 🔀 Hierarquias")
                md.append(f"")
                for hier in tabela.hierarquias:
                    niveis_str = " → ".join(hier.niveis)
                    md.append(f"- **{hier.nome}**: {niveis_str}")
                md.append(f"")

            # Fonte de Dados
            if tabela.particao and tabela.particao.codigo_fonte:
                md.append(f"#### {_t('table.section.source_data')}")
                md.append(f"")

                # Badges para modo e grupo
                modo_badge = f"`{tabela.particao.modo}`"
                grupo_badge = f"`{tabela.particao.grupo_consulta}`" if tabela.particao.grupo_consulta else ""

                if grupo_badge:
                    md.append(f"{_t('table.source.mode_label')}: {modo_badge} | {_t('table.source.group_label')}: {grupo_badge}")
                else:
                    md.append(f"{_t('table.source.mode_label')}: {modo_badge}")
                md.append(f"")

                # Código fonte formatado
                codigo_formatado = self._formatar_codigo_fonte(tabela.particao.codigo_fonte)
                is_dax = _codigo_fonte_eh_dax(tabela.particao.codigo_fonte)
                lang_fonte = "dax" if is_dax else "powerquery"
                src_label = _t("table.source.code_dax") if is_dax else _t("table.source.code_m")
                md.append(src_label)
                md.append(f"")
                md.append(f"```{lang_fonte}")
                md.append(codigo_formatado)
                md.append(f"```")
                md.append(f"")



        # Visuais Personalizados
        if self.visuais_personalizados:
            md.append(f"---")
            md.append(f"")
            md.append(f"## {_t('doc.toc.custom_visuals')}")
            md.append(f"")
            md.append(f"| ID do Visual |")
            md.append(f"|--------------|")

            for visual in self.visuais_personalizados:
                md.append(f"| {visual} |")

            md.append(f"")

        # Recursos de Imagem
        if self.recursos_imagem:
            md.append(f"---")
            md.append(f"")
            md.append(f"## {_t('doc.toc.image_resources')}")
            md.append(f"")
            md.append(f"| Nome | Tipo |")
            md.append(f"|------|------|")

            for recurso in self.recursos_imagem:
                md.append(f"| {recurso['nome']} | {recurso['tipo']} |")

            md.append(f"")

        self._md_cache = '\n'.join(md)
        return self._md_cache

    def salvar_documentacao(self, caminho_saida: Optional[str] = None):
        """
        Salva a documentação em arquivo

        Args:
            caminho_saida: Caminho para salvar. Se None, salva na pasta do projeto
        """
        if caminho_saida is None:
            caminho_saida = self.caminho_projeto / f"{self.nome_projeto}_documentacao.md"

        markdown = self.gerar_documentacao()

        with open(caminho_saida, 'w', encoding='utf-8-sig') as f:
            f.write(markdown)

        print(f"[OK] Documentação MD salva em: {caminho_saida}")
        return caminho_saida

    def _gerar_html_documentacao(self, auto_print: bool = False) -> str:
        """Gera HTML imprimivel a partir da documentacao Markdown."""
        try:
            import markdown
        except ImportError:
            raise ImportError("O pacote 'markdown' nao esta instalado. Execute: pip install markdown")

        markdown_text = self.gerar_documentacao()
        html_body = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code', 'toc'])
        cor_primaria = self.branding.primary_color
        cor_secundaria = self.branding.secondary_color
        cor_clara = self.branding.light_color
        fact_tables = self._tabelas_fato_mermaid()
        fact_tables_js = json.dumps(fact_tables, ensure_ascii=True)

        css_style = """
        html { -webkit-font-smoothing: antialiased; -moz-osx-font-smoothing: grayscale; }
        body { font-family: 'Inter', 'Segoe UI', system-ui, -apple-system, sans-serif; color: #2C3E50; line-height: 1.6; padding: 20px 40px; font-size: 10.5pt; }
        h1, h2, h3, h4, h5 { color: #1A5276; font-weight: 600; margin-top: 1.5em; margin-bottom: 0.5em; page-break-after: avoid; }
        h1 { border-bottom: 2px solid #2980B9; padding-bottom: 0.3em; font-size: 24pt; }
        h2 { border-bottom: 1px solid #EAECEE; padding-bottom: 0.3em; font-size: 18pt; }
        h3 { font-size: 14pt; color: #2980B9; }
        table { border-collapse: collapse; width: 100%; margin: 1.5em 0; font-size: 9.5pt; box-shadow: 0 1px 3px rgba(0,0,0,0.1); table-layout: auto; word-break: break-word; }
        th { background-color: #2980B9; color: white; padding: 10px 12px; text-align: left; font-weight: 600; }
        td { border-bottom: 1px solid #EAECEE; padding: 10px 12px; }
        tr:nth-child(even) { background-color: #F8F9F9; }
        tr { page-break-inside: avoid; }
        code { font-family: 'Consolas', 'Courier New', monospace; background-color: #F2F4F4; padding: 2px 5px; border-radius: 4px; font-size: 9pt; color: #C0392B; word-break: break-all; }
        pre code { display: block; padding: 15px; border-left: 4px solid #2980B9; overflow-x: auto; line-height: 1.4; white-space: pre-wrap; word-wrap: break-word; border-radius: 0 6px 6px 0; }
        blockquote { border-left: 4px solid #3498DB; background: #EBF5FB; padding: 12px 16px; margin: 1.5em 0; color: #2C3E50; border-radius: 0 4px 4px 0; }
        hr { border: 0; height: 1px; background: #EAECEE; margin: 2em 0; }
        pre code.hljs { border-left: 4px solid #2980B9; border-radius: 0 6px 6px 0; }
        :not(pre) > code { background-color: #F2F4F4 !important; color: #C0392B !important; padding: 2px 5px; border-radius: 4px; font-size: 9pt; }
        .language-mermaid { white-space: pre-wrap; word-wrap: break-word; font-family: 'Consolas', monospace; }
        .mermaid { display: flex; justify-content: center; margin: 2em 0; page-break-inside: avoid; }
        .mermaid .fact-table rect { fill: #2980B9; stroke: #1A5276; stroke-width: 1.5px; }
        .mermaid .fact-table text { fill: #FFFFFF; font-weight: 600; }
        .cover-page { display: flex; flex-direction: column; justify-content: center; align-items: center; min-height: 85vh; text-align: center; padding: 2em 0; }
        .cover-page h1.cover-title { font-size: 36pt; border-bottom: 3px solid #2980B9; padding-bottom: 0.4em; margin-bottom: 0.6em; color: #1A5276; }
        .cover-logo { margin-bottom: 1.5em; }
        .cover-project { font-size: 20pt; color: #2C3E50; font-weight: 600; margin: 0.3em 0; }
        .cover-date { font-size: 12pt; color: #7F8C8D; margin-top: 1em; }
        .page-break { height: 0; margin: 0; padding: 0; border: none; }
        @media print {
            @page { size: A4; margin: 2cm; }
            body { padding: 0; -webkit-print-color-adjust: exact; print-color-adjust: exact; }
            p, li { orphans: 3; widows: 3; }
            pre { page-break-inside: avoid; }
            a[href^="http"]:after { content: " (" attr(href) ")"; font-size: 80%; color: #555; }
            a { color: inherit; text-decoration: none; }
            .page-break { page-break-after: always; break-after: page; }
        }
        """
        css_style = (
            css_style
            .replace("#1A5276", cor_primaria)
            .replace("#2980B9", cor_secundaria)
            .replace("#3498DB", cor_secundaria)
            .replace("#EBF5FB", cor_clara)
        )

        print_script = ""

        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{html.escape(self.branding.document_title)} - {html.escape(str(self.nome_projeto))}</title>
            <link rel="preconnect" href="https://fonts.googleapis.com">
            <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
            <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600&display=swap" rel="stylesheet">
            <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/atom-one-light.min.css">
            <script src="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/highlight.min.js"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/languages/sql.min.js"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/languages/powershell.min.js"></script>
            <style>{css_style}</style>
            <script type="module">
                import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
                mermaid.initialize({{ startOnLoad: true, theme: 'default' }});
                const factTables = new Set({fact_tables_js});

                document.addEventListener("DOMContentLoaded", () => {{
                    document.querySelectorAll("code.language-mermaid").forEach(el => {{
                        const div = document.createElement("div");
                        div.className = "mermaid";
                        div.textContent = el.textContent;
                        if (el.parentNode.tagName === 'PRE') {{
                            el.parentNode.replaceWith(div);
                        }}
                    }});

                    document.querySelectorAll("code.language-dax").forEach(el => {{
                        el.classList.remove("language-dax");
                        el.classList.add("language-sql");
                    }});
                    document.querySelectorAll("code.language-powerquery").forEach(el => {{
                        el.classList.remove("language-powerquery");
                        el.classList.add("language-powershell");
                    }});

                    hljs.highlightAll();

                    if (factTables.size > 0) {{
                        const applyFactHighlight = () => {{
                            const entities = document.querySelectorAll(".mermaid svg g.entity");
                            if (!entities.length) {{
                                return false;
                            }}
                            entities.forEach(entity => {{
                                const label = entity.querySelector("text");
                                if (!label) {{
                                    return;
                                }}
                                const name = (label.textContent || "").trim();
                                if (factTables.has(name)) {{
                                    entity.classList.add("fact-table");
                                }}
                            }});
                            return true;
                        }};

                        const waitForMermaid = (attempt = 0) => {{
                            if (applyFactHighlight()) {{
                                return;
                            }}
                            if (attempt < 30) {{
                                setTimeout(() => waitForMermaid(attempt + 1), 150);
                            }}
                        }};

                        waitForMermaid();
                    }}
                }});

                {print_script}
            </script>
        </head>
        <body>
            {html_body}
        </body>
        </html>
        """

    def salvar_documentacao_html(self, caminho_saida: Optional[str] = None, auto_print: bool = False):
        """
        Salva a documentacao em HTML imprimivel para o usuario gerar PDF pelo navegador.
        """
        if caminho_saida is None:
            caminho_saida = self.caminho_projeto / f"{self.nome_projeto}_documentacao.html"

        html = self._gerar_html_documentacao(auto_print=auto_print)

        with open(caminho_saida, 'w', encoding='utf-8-sig') as f:
            f.write(html)

        print(f"[OK] Documentacao HTML salva em: {caminho_saida}")
        return caminho_saida

    def salvar_documentacao_docx(self, caminho_saida: Optional[str] = None):
        """
        Salva a documentação em formato Word (.docx) com formatação profissional

        Args:
            caminho_saida: Caminho para salvar. Se None, salva na pasta do projeto
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor, Emu
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.section import WD_ORIENT
            from docx.oxml.ns import qn, nsdecls
            from docx.oxml import parse_xml
        except ImportError:
            print("[ERRO] Pacote 'python-docx' não instalado. Execute: pip install python-docx")
            return None

        if caminho_saida is None:
            caminho_saida = self.caminho_projeto / f"{self.nome_projeto}_documentacao.docx"

        doc = Document()

        # ====================================================================
        # CORES E CONSTANTES
        # ====================================================================
        AZUL_PRI = RGBColor(0x00, 0x3D, 0x6B)    # Azul escuro primário
        AZUL_SEC = RGBColor(0x00, 0x6D, 0xAA)    # Azul médio
        AZUL_CLR = RGBColor(0xD6, 0xE8, 0xF5)    # Azul claro fundo
        CINZA_TX = RGBColor(0x44, 0x44, 0x44)     # Texto principal
        CINZA_LT = RGBColor(0x77, 0x77, 0x77)     # Texto secundário
        BRANCO   = RGBColor(0xFF, 0xFF, 0xFF)

        HEX_AZUL_PRI = "003D6B"
        HEX_AZUL_SEC = "006DAA"
        HEX_AZUL_CLR = "D6E8F5"
        HEX_CINZA_BG = "F7F8FA"
        HEX_CODE_BG  = "F0F2F5"
        HEX_CODE_BRD = "006DAA"
        HEX_TABLE_ALT = "FAFBFC"

        FONT_MAIN = "Aptos"
        FONT_CODE = "Consolas"
        FONT_BODY = Pt(10)
        FONT_TABLE = Pt(8.5)
        FONT_TABLE_COMPACT = Pt(8)
        FONT_SMALL = Pt(8)
        FONT_MICRO = Pt(7.5)
        FONT_CODE_SIZE = Pt(8)
        FONT_H1 = Pt(18)
        FONT_H2 = Pt(14)
        FONT_H3 = Pt(11.5)
        FONT_H4 = Pt(10.5)
        FONT_COVER_TITLE = Pt(24)
        FONT_COVER_PROJECT = Pt(30)
        FONT_DAX_TITLE = Pt(11)
        FONT_DAX_TEXT = Pt(8.5)

        def _rgb_from_hex(cor: str) -> RGBColor:
            cor = cor.lstrip("#")
            return RGBColor(int(cor[0:2], 16), int(cor[2:4], 16), int(cor[4:6], 16))

        AZUL_PRI = _rgb_from_hex(self.branding.primary_color)
        AZUL_SEC = _rgb_from_hex(self.branding.secondary_color)
        AZUL_CLR = _rgb_from_hex(self.branding.light_color)
        HEX_AZUL_PRI = self._hex_sem_hash(self.branding.primary_color)
        HEX_AZUL_SEC = self._hex_sem_hash(self.branding.secondary_color)
        HEX_AZUL_CLR = self._hex_sem_hash(self.branding.light_color)
        HEX_CODE_BRD = HEX_AZUL_SEC

        # ====================================================================
        # CONFIGURAÇÃO DA PÁGINA
        # ====================================================================
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

        # ====================================================================
        # ESTILOS GLOBAIS
        # ====================================================================
        def _set_style_font(style, name=FONT_MAIN, size=None, color=None, bold=None):
            style.font.name = name
            style._element.rPr.rFonts.set(qn('w:eastAsia'), name)
            if size is not None:
                style.font.size = size
            if color is not None:
                style.font.color.rgb = color
            if bold is not None:
                style.font.bold = bold

        def _set_run_font(run, name=FONT_MAIN, size=None, color=None, bold=None, italic=None):
            run.font.name = name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
            if size is not None:
                run.font.size = size
            if color is not None:
                run.font.color.rgb = color
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            return run

        style_normal = doc.styles['Normal']
        _set_style_font(style_normal, FONT_MAIN, FONT_BODY, CINZA_TX)
        style_normal.paragraph_format.space_after = Pt(5)
        style_normal.paragraph_format.line_spacing = 1.08

        # Estilos de heading
        for lvl in range(1, 5):
            style_h = doc.styles[f'Heading {lvl}']
            _set_style_font(style_h, FONT_MAIN, color=AZUL_PRI, bold=True)
            if lvl == 1:
                style_h.font.size = FONT_H1
                style_h.paragraph_format.space_before = Pt(16)
                style_h.paragraph_format.space_after = Pt(8)
            elif lvl == 2:
                style_h.font.size = FONT_H2
                style_h.paragraph_format.space_before = Pt(12)
                style_h.paragraph_format.space_after = Pt(6)
            elif lvl == 3:
                style_h.font.size = FONT_H3
                style_h.font.color.rgb = AZUL_SEC
                style_h.paragraph_format.space_before = Pt(10)
                style_h.paragraph_format.space_after = Pt(4)
            else:
                style_h.font.size = FONT_H4
                style_h.font.color.rgb = AZUL_SEC
                style_h.paragraph_format.space_before = Pt(8)
                style_h.paragraph_format.space_after = Pt(3)

        # ====================================================================
        # FUNÇÕES AUXILIARES
        # ====================================================================
        def _add_separator():
            """Adiciona linha separadora sutil"""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Borda inferior como separador
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{HEX_AZUL_CLR}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            return p

        def _add_table(headers, rows, compact=False):
            """Cria tabela profissional com estilo refinado"""
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Remove bordas padrão e aplica estilo customizado
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            # Bordas suaves da tabela
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E5EB"/>'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="E0E5EB"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(borders)
            cell_margins = parse_xml(
                f'<w:tblCellMar {nsdecls("w")}>'
                f'  <w:top w:w="{80 if compact else 110}" w:type="dxa"/>'
                f'  <w:left w:w="120" w:type="dxa"/>'
                f'  <w:bottom w:w="{80 if compact else 110}" w:type="dxa"/>'
                f'  <w:right w:w="120" w:type="dxa"/>'
                f'</w:tblCellMar>'
            )
            tblPr.append(cell_margins)

            font_size = FONT_TABLE_COMPACT if compact else FONT_TABLE

            # Cabeçalho
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(header)
                _set_run_font(run, FONT_MAIN, font_size, BRANCO, bold=True)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_AZUL_PRI}"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            # Dados
            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    cell = table.rows[r_idx + 1].cells[c_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(str(cell_text))
                    _set_run_font(run, FONT_MAIN, font_size, CINZA_TX)
                    # Fundo alternado
                    if r_idx % 2 == 0:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_TABLE_ALT}"/>')
                        cell._tc.get_or_add_tcPr().append(shading)

            doc.add_paragraph("").paragraph_format.space_after = Pt(4)
            return table

        CODE_TOKEN_RE = re.compile(
            r"(//.*$|--.*$|/\*.*?\*/|#\"(?:[^\"]|\"\")*\"|\"(?:[^\"]|\"\")*\"|'[^']*'|\[[^\]]+\]|\b\d+(?:[.,]\d+)?\b|[A-Za-z_#][A-Za-z0-9_#]*(?:\.[A-Za-z_#][A-Za-z0-9_#]*)*|[=+\-*/<>^&|(),{};:])"
        )
        POWER_QUERY_HIGHLIGHT_KEYWORDS = {
            "let", "in", "each", "if", "then", "else", "try", "otherwise",
            "and", "or", "not", "true", "false", "null", "as", "is", "meta",
        }
        DAX_HIGHLIGHT_KEYWORDS = {
            "var", "return", "if", "then", "else", "true", "false", "blank",
            "in", "not", "and", "or",
        }
        POWER_QUERY_HIGHLIGHT_FUNCTIONS = set(POWER_QUERY_FUNCTION_CATALOG.keys())
        DAX_HIGHLIGHT_FUNCTIONS = set(DAX_FUNCTION_CATALOG.keys())
        POWER_QUERY_GENERIC_FUNCTION_PREFIXES = (
            "Table.", "Text.", "Date.", "DateTime.", "List.", "Record.",
            "Number.", "Value.", "Binary.",
        )

        def _normalizar_linguagem_codigo(language=""):
            language_norm = str(language or "").strip().lower()
            if "dax" in language_norm:
                return "dax"
            if "power" in language_norm or language_norm in {"m", "powerquery"}:
                return "powerquery"
            return language_norm

        def _tokenizar_linha_codigo(line, language_key):
            tokens = []
            pos = 0
            for match in CODE_TOKEN_RE.finditer(str(line or "")):
                start, end = match.span()
                if start > pos:
                    tokens.append((line[pos:start], "text"))

                token = match.group(0)
                token_upper = token.upper()
                token_lower = token.lower()
                token_kind = "text"

                if token.startswith("//") or token.startswith("--") or token.startswith("/*"):
                    token_kind = "comment"
                elif token.startswith('"') or token.startswith("'"):
                    token_kind = "string"
                elif token.startswith('#"') or token.startswith("["):
                    token_kind = "identifier"
                elif re.fullmatch(r"\d+(?:[.,]\d+)?", token):
                    token_kind = "number"
                elif re.fullmatch(r"[=+\-*/<>^&|(),{};:]", token):
                    token_kind = "operator"
                elif language_key == "powerquery":
                    if token_lower in POWER_QUERY_HIGHLIGHT_KEYWORDS:
                        token_kind = "keyword"
                    elif token in POWER_QUERY_HIGHLIGHT_FUNCTIONS or any(token.startswith(prefix) for prefix in POWER_QUERY_GENERIC_FUNCTION_PREFIXES):
                        token_kind = "function"
                elif language_key == "dax":
                    if token_upper in DAX_HIGHLIGHT_FUNCTIONS:
                        token_kind = "function"
                    elif token_lower in DAX_HIGHLIGHT_KEYWORDS:
                        token_kind = "keyword"

                tokens.append((token, token_kind))
                pos = end

            if pos < len(line):
                tokens.append((line[pos:], "text"))

            return tokens or [("", "text")]

        def _estilizar_run_codigo(run, token_kind):
            colors = {
                "text": RGBColor(0x1E, 0x1E, 0x2E),
                "comment": RGBColor(0x5B, 0x6F, 0x59),
                "string": RGBColor(0xA1, 0x5C, 0x38),
                "number": RGBColor(0x7C, 0x3A, 0xA8),
                "keyword": AZUL_PRI,
                "function": AZUL_SEC,
                "identifier": RGBColor(0x47, 0x54, 0x67),
                "operator": RGBColor(0x66, 0x70, 0x85),
            }
            _set_run_font(
                run,
                FONT_CODE,
                FONT_CODE_SIZE,
                colors.get(token_kind, colors["text"]),
                bold=token_kind in {"keyword", "function"},
                italic=token_kind == "comment",
            )
            return run

        def _add_technical_label(text, space_before=8):
            p_label = doc.add_paragraph()
            p_label.paragraph_format.space_after = Pt(2)
            p_label.paragraph_format.space_before = Pt(space_before)
            p_label.paragraph_format.left_indent = Cm(0.15)
            run_label = p_label.add_run(text)
            _set_run_font(run_label, FONT_MAIN, FONT_MICRO, AZUL_SEC, bold=True)
            return p_label

        def _add_code_block(code, language="", label: Optional[str] = None, compact=False):
            """Adiciona bloco de código com syntax highlighting leve no DOCX."""
            language_key = _normalizar_linguagem_codigo(language)

            # Label da linguagem (se fornecida)
            if language or label:
                label_text = label or ("EXPRESSÃO DAX" if language_key == "dax" else language.upper())
                _add_technical_label(label_text, 8 if language_key == "dax" else 6)

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2) if language else Pt(6)
            p.paragraph_format.space_after = Pt(8 if compact else (10 if language_key == "dax" else 6))
            p.paragraph_format.left_indent = Cm(0.35)
            p.paragraph_format.right_indent = Cm(0.15)
            p.paragraph_format.line_spacing = Pt(11)

            # Fundo cinza + borda lateral azul
            pPr = p._p.get_or_add_pPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{HEX_CODE_BG}"/>')
            pPr.append(shading)
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="18" w:space="6" w:color="{HEX_CODE_BRD}"/>'
                f'  <w:top w:val="single" w:sz="2" w:space="2" w:color="D0D5DD"/>'
                f'  <w:bottom w:val="single" w:sz="2" w:space="2" w:color="D0D5DD"/>'
                f'  <w:right w:val="single" w:sz="2" w:space="2" w:color="D0D5DD"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

            code_lines = str(code or "").replace("\t", "    ").splitlines()
            if not code_lines:
                code_lines = [""]

            for line_idx, line in enumerate(code_lines):
                if language_key in {"powerquery", "dax"}:
                    for token_text, token_kind in _tokenizar_linha_codigo(line, language_key):
                        if token_text:
                            _estilizar_run_codigo(p.add_run(token_text), token_kind)
                else:
                    _estilizar_run_codigo(p.add_run(line), "text")

                if line_idx < len(code_lines) - 1:
                    p.add_run().add_break()

            return p

        def _add_info_box(text, tipo="info"):
            """Adiciona caixa informativa colorida"""
            cores = {
                "info":    (HEX_AZUL_CLR, HEX_AZUL_SEC, AZUL_PRI),
                "warning": ("FFF3CD",      "FFC107",     RGBColor(0x85, 0x6D, 0x00)),
                "success": ("D4EDDA",      "28A745",     RGBColor(0x15, 0x57, 0x24)),
            }
            bg, brd, txt_cor = cores.get(tipo, cores["info"])

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Cm(0.5)

            pPr = p._p.get_or_add_pPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg}"/>')
            pPr.append(shading)
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{brd}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

            run = p.add_run(text)
            _set_run_font(run, FONT_MAIN, FONT_TABLE, txt_cor)
            return p

        def _add_bullet_list(titulo, itens):
            if not itens:
                return
            h = doc.add_heading(titulo, level=4)
            h.paragraph_format.space_before = Pt(4)
            for item in itens:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(str(item))
                _set_run_font(run, FONT_MAIN, FONT_TABLE, CINZA_TX)

        def _add_regra_power_query_docx(regra: RegraPowerQuery):
            if not regra_power_query_tem_conteudo(regra):
                return

            h = doc.add_heading(_t("docx.section.business_rule_inferred"), level=3)
            h.paragraph_format.space_before = Pt(8)
            _add_info_box(_t("docx.info.pq_intro"), "info")

            if regra.observacoes:
                _add_info_box(_t("docx.info.pq_observations_intro"), "info")
                itens_doc = []
                rotulos = {
                    "geral": _t("pq.obs.general"),
                    "origem": _t("pq.obs.source"),
                    "regra": _t("pq.obs.rule"),
                    "observacao": _t("pq.obs.observation"),
                }
                for chave, valores in regra.observacoes.items():
                    for valor in valores:
                        itens_doc.append(f"{rotulos.get(chave, chave.title())}: {valor}")
                _add_bullet_list(_t("docx.pq.documented_observations"), itens_doc)

            if regra.linhas_regra:
                resumo = _resumir_etapas_power_query(regra.linhas_regra)
                if resumo:
                    p_resumo = doc.add_paragraph()
                    run_lbl = p_resumo.add_run(_t("docx.pq.summary_label") + ": ")
                    run_lbl.bold = True
                    p_resumo.add_run(f"{resumo}.")
                    p_resumo.paragraph_format.space_before = Pt(4)
                    p_resumo.paragraph_format.space_after = Pt(4)
                h_regras = doc.add_heading(_t("docx.section.pq_rules"), level=4)
                h_regras.paragraph_format.space_before = Pt(4)
                linhas_agrupadas = _agrupar_linhas_regra_consecutivas(regra.linhas_regra)
                rows = [
                    [
                        _localize_etapa_pq(linha.etapa).replace("`", ""),
                        linha.regra.replace("`", ""),
                        linha.descricao.replace("`", ""),
                    ]
                    for linha in linhas_agrupadas
                ]
                _add_table(
                    [
                        _t("docx.tcol.step"),
                        _t("docx.tcol.rule_or_filter"),
                        _t("docx.tcol.description"),
                    ],
                    rows,
                    compact=True,
                )

        def _add_leitura_dax_docx(leitura: LeituraDax):
            linhas = []
            if leitura.categorias:
                for categoria in DAX_CATEGORY_ORDER:
                    itens = leitura.categorias.get(categoria, [])
                    for item in itens:
                        linhas.append((DAX_CATEGORY_LABELS.get(categoria, categoria.title()), item.replace("`", "")))
            elif leitura.itens:
                linhas = [("Leitura técnica", item.replace("`", "")) for item in leitura.itens]

            if not linhas:
                return

            _add_technical_label("LEITURA TÉCNICA DAX", 8)

            for categoria, texto in linhas:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.35)
                p.paragraph_format.right_indent = Cm(0.2)
                pPr = p._p.get_or_add_pPr()
                pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F8FBFE"/>'))
                pPr.append(parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:left w:val="single" w:sz="10" w:space="5" w:color="{HEX_AZUL_CLR}"/>'
                    f'</w:pBdr>'
                ))

                run_cat = p.add_run(f"{categoria}: ")
                _set_run_font(run_cat, FONT_MAIN, FONT_DAX_TEXT, AZUL_PRI, bold=True)

                run_txt = p.add_run(texto)
                _set_run_font(run_txt, FONT_MAIN, FONT_DAX_TEXT, CINZA_TX)

        def _add_format_code_block(label, value):
            if not value:
                return
            _add_code_block(str(value), "", label=label.upper(), compact=True)

        def _add_dax_object_header(tipo, nome, detalhe_label=None, detalhe_valor=None, descricao=None, show_tipo=True):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(3)
            pPr = p._p.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="5" w:color="E0E7F0"/>'
                f'</w:pBdr>'
            ))

            if show_tipo:
                run_tipo = p.add_run(f"{tipo.upper()}  ")
                _set_run_font(run_tipo, FONT_MAIN, FONT_MICRO, AZUL_SEC, bold=True)

            run_nome = p.add_run(nome)
            _set_run_font(run_nome, FONT_MAIN, FONT_DAX_TITLE, AZUL_PRI, bold=True)

            if detalhe_valor:
                _add_format_code_block(detalhe_label or "Detalhe", detalhe_valor)

            if descricao:
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.space_before = Pt(0)
                p_desc.paragraph_format.space_after = Pt(5)
                p_desc.paragraph_format.left_indent = Cm(0.15)
                r_desc = p_desc.add_run(descricao)
                _set_run_font(r_desc, FONT_MAIN, FONT_DAX_TEXT, CINZA_LT, italic=True)

        def _add_stat_card(label, value, emoji=""):
            """Retorna lista [label, valor] para card de estatística"""
            return [f"{emoji} {label}" if emoji else label, str(value)]

        # ====================================================================
        # CONTADORES
        # ====================================================================
        total_colunas = sum(len(t.colunas) for t in self.tabelas)
        total_medidas = sum(len(t.medidas) for t in self.tabelas)
        total_calc = sum(len(t.colunas_calculadas) for t in self.tabelas)
        dicionario_dados = self.gerar_dicionario_dados()

        # ====================================================================
        # CAPA
        # ====================================================================
        # Espaçamento superior
        for _ in range(2):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)

        # Barra decorativa superior
        p_barra = doc.add_paragraph()
        p_barra.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p_barra._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="48" w:space="1" w:color="{HEX_AZUL_PRI}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        logo_path = self._obter_logo_path()
        if logo_path:
            try:
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_logo.paragraph_format.space_before = Pt(18)
                p_logo.paragraph_format.space_after = Pt(10)
                p_logo.add_run().add_picture(str(logo_path), width=Cm(2.2))
            except Exception as e:
                print(f"  [AVISO] Falha ao inserir logo no DOCX: {e}")

        # Título único
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_before = Pt(12 if logo_path else 24)
        p_titulo.paragraph_format.space_after = Pt(8)
        run = p_titulo.add_run(self.branding.document_title)
        _set_run_font(run, FONT_MAIN, FONT_COVER_TITLE, AZUL_SEC)

        # Nome do projeto
        p_proj = doc.add_paragraph()
        p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_proj.paragraph_format.space_after = Pt(12)
        run = p_proj.add_run(self.nome_projeto)
        _set_run_font(run, FONT_MAIN, FONT_COVER_PROJECT, AZUL_PRI, bold=True)

        # Barra decorativa inferior
        p_barra2 = doc.add_paragraph()
        p_barra2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr2 = p_barra2._p.get_or_add_pPr()
        pBdr2 = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="24" w:space="1" w:color="{HEX_AZUL_SEC}"/>'
            f'</w:pBdr>'
        )
        pPr2.append(pBdr2)

        # Espaçamento
        doc.add_paragraph("").paragraph_format.space_after = Pt(24)

        # Estatísticas resumidas na capa
        info_capa = [
            (_t("docx.stat.tables"), str(len(self.tabelas))),
            (_t("docx.stat.measures_dax"), str(total_medidas)),
            (_t("docx.stat.columns"), str(total_colunas)),
            (_t("docx.stat.relationships"), str(len(self._relacionamentos_usuario()))),
        ]

        tbl_info = doc.add_table(rows=len(info_capa), cols=2)
        tbl_info.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, (label, valor) in enumerate(info_capa):
            cell_l = tbl_info.rows[r_idx].cells[0]
            cell_l.text = ''
            p = cell_l.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(label + "  ")
            _set_run_font(run, FONT_MAIN, FONT_H4, CINZA_LT)

            cell_v = tbl_info.rows[r_idx].cells[1]
            cell_v.text = ''
            p = cell_v.paragraphs[0]
            run = p.add_run("  " + valor)
            _set_run_font(run, FONT_MAIN, FONT_H4, AZUL_PRI, bold=True)

        # Quebra de página após capa
        doc.add_page_break()

        # ====================================================================
        # SUMÁRIO (Table of Contents)
        # ====================================================================
        h_toc = doc.add_heading(_t("docx.toc.title"), level=1)

        # Insere campo TOC do Word (atualizado ao abrir o documento)
        p_toc = doc.add_paragraph()
        run_toc = p_toc.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_toc._r.append(fldChar1)
        run_toc2 = p_toc.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        run_toc2._r.append(instrText)
        run_toc3 = p_toc.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run_toc3._r.append(fldChar2)
        run_toc4 = p_toc.add_run("(Clique com botão direito e selecione \"Atualizar campo\" para gerar o sumário)")
        _set_run_font(run_toc4, FONT_MAIN, FONT_TABLE, CINZA_LT, italic=True)
        run_toc5 = p_toc.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_toc5._r.append(fldChar3)

        doc.add_page_break()

        # ====================================================================
        # CABEÇALHO E RODAPÉ
        # ====================================================================
        header = section.header
        header.is_linked_to_previous = False
        p_header = header.paragraphs[0]
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run_h = p_header.add_run(f"{self.nome_projeto}")
        _set_run_font(run_h, FONT_MAIN, FONT_SMALL, CINZA_LT)
        _set_run_font(p_header.add_run("  |  "), FONT_MAIN, FONT_SMALL, RGBColor(0xCC, 0xCC, 0xCC))
        run_h2 = p_header.add_run("Documentação Power BI")
        _set_run_font(run_h2, FONT_MAIN, FONT_SMALL, AZUL_SEC)

        # Borda inferior no cabeçalho
        pPr_h = p_header._p.get_or_add_pPr()
        pBdr_h = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="{HEX_AZUL_CLR}"/>'
            f'</w:pBdr>'
        )
        pPr_h.append(pBdr_h)

        # Rodapé com numeração
        footer = section.footer
        footer.is_linked_to_previous = False
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Borda superior no rodapé
        pPr_f = p_footer._p.get_or_add_pPr()
        pBdr_f = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="6" w:space="4" w:color="{HEX_AZUL_CLR}"/>'
            f'</w:pBdr>'
        )
        pPr_f.append(pBdr_f)

        run_f1 = p_footer.add_run("Página ")
        _set_run_font(run_f1, FONT_MAIN, FONT_SMALL, CINZA_LT)

        # Campo de número de página
        fldChar_pg1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_pg = p_footer.add_run()
        run_pg._r.append(fldChar_pg1)
        instrPg = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_pg2 = p_footer.add_run()
        run_pg2._r.append(instrPg)
        fldChar_pg2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_pg3 = p_footer.add_run()
        run_pg3._r.append(fldChar_pg2)

        run_f2 = p_footer.add_run(" de ")
        _set_run_font(run_f2, FONT_MAIN, FONT_SMALL, CINZA_LT)

        # Campo total de páginas
        fldChar_np1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_np = p_footer.add_run()
        run_np._r.append(fldChar_np1)
        instrNp = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        run_np2 = p_footer.add_run()
        run_np2._r.append(instrNp)
        fldChar_np2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_np3 = p_footer.add_run()
        run_np3._r.append(fldChar_np2)

        # ====================================================================
        # VISÃO GERAL
        # ====================================================================
        doc.add_heading(_t("docx.section.overview"), level=1)

        # Cards de estatísticas em tabela 2x3
        stats_table = doc.add_table(rows=2, cols=3)
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        stats_data = [
            (_t("docx.stat.tables"), len(self.tabelas)),
            (_t("docx.stat.measures"), total_medidas),
            (_t("docx.stat.columns"), total_colunas),
            (_t("docx.stat.calc_columns"), total_calc),
            (_t("docx.stat.relationships"), len(self._relacionamentos_usuario())),
            (_t("docx.stat.pages"), len(self.paginas)),
        ]

        for idx, (label, value) in enumerate(stats_data):
            row = idx // 3
            col = idx % 3
            cell = stats_table.rows[row].cells[col]
            cell.text = ''

            # Fundo azul claro
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_AZUL_CLR}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

            # Valor grande
            p_val = cell.paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_val.paragraph_format.space_before = Pt(8)
            p_val.paragraph_format.space_after = Pt(2)
            run_val = p_val.add_run(str(value))
            _set_run_font(run_val, FONT_MAIN, FONT_COVER_TITLE, AZUL_PRI, bold=True)

            # Label abaixo
            p_lbl = cell.add_paragraph()
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_lbl.paragraph_format.space_before = Pt(0)
            p_lbl.paragraph_format.space_after = Pt(8)
            run_lbl = p_lbl.add_run(label)
            _set_run_font(run_lbl, FONT_MAIN, FONT_TABLE, AZUL_SEC, bold=True)

        # Resumo do modelo dimensional (apenas quando ha convencao fat_/dim_)
        resumo_dim = self._resumo_modelo_dimensional()
        if resumo_dim:
            p_dim = doc.add_paragraph()
            p_dim.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_dim.paragraph_format.space_before = Pt(10)
            p_dim.paragraph_format.space_after = Pt(4)
            run_dim_lbl = p_dim.add_run(_t("docx.stat.dimensional_label") + ": ")
            _set_run_font(run_dim_lbl, FONT_MAIN, FONT_TABLE, AZUL_PRI, bold=True)
            # _resumo_modelo_dimensional retorna markdown com **; converte para texto puro
            texto_limpo = resumo_dim.replace("**", "")
            run_dim_val = p_dim.add_run(f"{texto_limpo}.")
            _set_run_font(run_dim_val, FONT_MAIN, FONT_TABLE, CINZA_LT)

        if dicionario_dados:
            _add_separator()
            doc.add_heading(_t("docx.section.dictionary"), level=1)
            _add_info_box(_t("docx.info.dict_intro"), "info")
            rows_dic = []
            for termo in dicionario_dados:
                rows_dic.append([
                    termo.termo,
                    str(termo.frequencia),
                    self._localizar_categoria_termo(termo.categoria),
                    ", ".join(termo.fontes[:5]) or "—",
                    "; ".join(termo.exemplos[:3]) or "—",
                ])
            _add_table(
                [
                    _t("docx.tcol.term"),
                    _t("docx.tcol.frequency"),
                    _t("docx.tcol.category"),
                    _t("docx.tcol.sources"),
                    _t("docx.tcol.examples"),
                ],
                rows_dic,
                compact=True,
            )

        _add_separator()

        # ====================================================================
        # PÁGINAS DO RELATÓRIO
        # ====================================================================
        doc.add_heading(_t("docx.section.pages"), level=1)

        if self.paginas:
            _add_info_box(_t("docx.info.pages_count", n=len(self.paginas)), "success")
            rows_pag = []
            for i, pagina in enumerate(self.paginas, 1):
                dimensoes = f"{pagina.largura} × {pagina.altura}"
                filtros_interpretados = self._interpretar_filtros_pagina(pagina.filtros) if pagina.filtros else []
                qtd = str(len(filtros_interpretados)) if filtros_interpretados else "—"
                rows_pag.append([str(i), pagina.nome_exibicao, pagina.tipo, dimensoes, qtd])
            _add_table(
                [
                    _t("docx.tcol.idx"),
                    _t("docx.tcol.page_name"),
                    _t("docx.tcol.page_type"),
                    _t("docx.tcol.dimensions"),
                    _t("docx.tcol.filters"),
                ],
                rows_pag,
            )

            # Detalhamento dos filtros por página
            paginas_com_filtros = [
                (p, self._interpretar_filtros_pagina(p.filtros))
                for p in self.paginas if p.filtros
            ]
            paginas_com_filtros = [(p, f) for p, f in paginas_com_filtros if f]

            if paginas_com_filtros:
                doc.add_heading(_t("docx.section.page_filters"), level=2)

                for pagina, filtros in paginas_com_filtros:
                    p_nome = doc.add_paragraph()
                    p_nome.paragraph_format.space_before = Pt(8)
                    p_nome.paragraph_format.space_after = Pt(4)
                    run = p_nome.add_run(f"▸ {pagina.nome_exibicao}")
                    _set_run_font(run, FONT_MAIN, FONT_BODY, AZUL_PRI, bold=True)

                    rows_filt = []
                    for f in filtros:
                        valores_str = ', '.join(f['valores']) if f['valores'] else '—'
                        rows_filt.append([f['tabela'], f['coluna'], f['tipo'], valores_str])
                    _add_table(
                        [
                            _t("docx.tcol.table"),
                            _t("docx.tcol.column"),
                            _t("docx.tcol.type"),
                            _t("docx.tcol.values"),
                        ],
                        rows_filt,
                        compact=True,
                    )
        else:
            _add_info_box(
                _t("docx.info.pages_empty_thin_report"),
                "warning"
            )

        _add_separator()

        # ====================================================================
        # MODELO DE DADOS
        # ====================================================================
        doc.add_heading(_t("docx.section.model"), level=1)

        # Lista de Relacionamentos (exclui tabelas técnicas de calendário, igual ao Markdown)
        rel_validos_docx = self._relacionamentos_usuario()
        if rel_validos_docx:
            doc.add_heading(_t("docx.section.relationships_list"), level=2)

            _add_info_box(
                _t("docx.info.rel_count", n=len(rel_validos_docx)),
                "info",
            )

            rows_rel = []
            for rel in rel_validos_docx:
                origem = f"{rel.tabela_origem}.{rel.coluna_origem}"
                destino = f"{rel.tabela_destino}.{rel.coluna_destino}"
                bidi = _t("docx.value.yes") if rel.filtro_bidirecional else _t("docx.value.no")
                ativo = _t("docx.value.yes") if rel.esta_ativo else _t("docx.value.no")
                rows_rel.append([origem, "→", destino, bidi, ativo])

            _add_table(
                [
                    _t("docx.rel.col.from"),
                    "",
                    _t("docx.rel.col.to"),
                    _t("docx.rel.col.bidirectional"),
                    _t("docx.rel.col.active"),
                ],
                rows_rel,
                compact=True,
            )

        # Grupos de Consulta
        if self.info_modelo.grupos_consulta:
            doc.add_heading(_t("docx.section.query_groups"), level=2)
            rows_gq = [[g['nome'], str(g['ordem'])] for g in sorted(self.info_modelo.grupos_consulta, key=lambda x: x['ordem'])]
            _add_table([_t("docx.tcol.gq_name"), _t("docx.tcol.gq_order")], rows_gq)

        # ====================================================================
        # RESUMO DAS TABELAS
        # ====================================================================
        doc.add_page_break()
        doc.add_heading(_t("docx.section.tables_summary"), level=1)

        p = doc.add_paragraph()
        run = p.add_run(f"O modelo contém {len(self.tabelas)} tabelas com {total_colunas} colunas, "
                        f"{total_medidas} medidas e {total_calc} colunas calculadas.")
        _set_run_font(run, FONT_MAIN, FONT_BODY, CINZA_LT, italic=True)

        rows_resumo = []
        for idx, tabela in enumerate(self.tabelas, 1):
            fonte = "Importação"
            if tabela.particao:
                if tabela.particao.grupo_consulta:
                    fonte = tabela.particao.grupo_consulta
                elif tabela.particao.codigo_fonte and _codigo_fonte_eh_dax(tabela.particao.codigo_fonte):
                    fonte = "DAX"
            rows_resumo.append([
                str(idx), tabela.nome,
                str(len(tabela.colunas)), str(len(tabela.medidas)),
                str(len(tabela.colunas_calculadas)), fonte
            ])

        _add_table(
            ["#", "Tabela", "Colunas", "Medidas", "Col. Calc.", "Fonte"],
            rows_resumo
        )

        _add_separator()

        # ====================================================================
        # DETALHAMENTO DAS TABELAS
        # ====================================================================
        doc.add_page_break()
        doc.add_heading(_t("docx.section.tables_detail"), level=1)

        for idx, tabela in enumerate(self.tabelas, 1):
            # Separador entre tabelas (exceto a primeira)
            if idx > 1:
                p_sep = doc.add_paragraph()
                p_sep.paragraph_format.space_before = Pt(12)
                p_sep.paragraph_format.space_after = Pt(12)
                run_sep = p_sep.add_run("─" * 40)
                _set_run_font(run_sep, FONT_MAIN, FONT_SMALL, RGBColor(0xE0, 0xE0, 0xE0))
                p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_heading(f"{idx}. {tabela.nome}", level=2)

            if hasattr(tabela, 'descricao') and tabela.descricao:
                p_desc = doc.add_paragraph()
                r_desc = p_desc.add_run(tabela.descricao)
                _set_run_font(r_desc, FONT_MAIN, FONT_BODY, CINZA_LT, italic=True)

            # Card de metadados
            status = _t("docx.tcard.status_visible") if not tabela.esta_oculta else _t("docx.tcard.status_hidden")
            refresh = _t("docx.value.yes") if not tabela.excluida_refresh else _t("docx.value.no")
            fonte_tipo = _t("docx.tcard.source_import")
            if tabela.particao:
                if tabela.particao.grupo_consulta:
                    fonte_tipo = tabela.particao.grupo_consulta
                elif tabela.particao.codigo_fonte and _codigo_fonte_eh_dax(tabela.particao.codigo_fonte):
                    fonte_tipo = _t("docx.tcard.source_dax")

            _add_table(
                [
                    _t("docx.tcard.status"),
                    _t("docx.tcard.refresh"),
                    _t("docx.tcard.columns"),
                    _t("docx.tcard.measures"),
                    _t("docx.tcard.source"),
                ],
                [[status, refresh, str(len(tabela.colunas)), str(len(tabela.medidas)), fonte_tipo]],
            )

            # Colunas
            if tabela.colunas:
                h = doc.add_heading(_t("docx.section.columns"), level=3)
                h.paragraph_format.space_before = Pt(8)
                rows_col = []
                for col in tabela.colunas:
                    oculta = _t("docx.value.yes") if col.esta_oculta else _t("docx.value.no")
                    rows_col.append([col.nome, col.tipo_dado, col.sumarizacao, oculta])
                _add_table(
                    [
                        _t("docx.tcol.col_name"),
                        _t("docx.tcol.col_type"),
                        _t("docx.tcol.summarization"),
                        _t("docx.tcol.hidden"),
                    ],
                    rows_col,
                    compact=True,
                )

            # Colunas Calculadas
            if tabela.colunas_calculadas:
                h = doc.add_heading(_t("docx.section.calc_columns"), level=3)
                h.paragraph_format.space_before = Pt(8)
                for coluna in tabela.colunas_calculadas:
                    _add_dax_object_header(
                        "Coluna calculada",
                        coluna.nome,
                        "Tipo de dado",
                        coluna.tipo_dado,
                    )
                    expr = coluna.expressao_dax.strip() if coluna.expressao_dax else "// Expressão não capturada"
                    _add_leitura_dax_docx(analisar_dax(expr))
                    _add_code_block(expr, "DAX")

            # Medidas
            if tabela.medidas:
                h = doc.add_heading(_t("docx.section.measures_dax"), level=3)
                h.paragraph_format.space_before = Pt(8)

                # Tabela resumo (Nome + Tipo inferido)
                rows_med = []
                for m in tabela.medidas:
                    tipo_med = inferir_tipo_medida(
                        m.nome,
                        m.expressao_dax or "",
                        analisar_dax(m.expressao_dax or ""),
                    )
                    rows_med.append([m.nome, tipo_med])
                _add_table(
                    [_t("docx.tcol.measure_name"), _t("docx.tcol.measure_type")],
                    rows_med,
                    compact=True,
                )

                # Código de cada medida
                h_cod = doc.add_heading(_t("docx.section.measures_code"), level=4)
                h_cod.paragraph_format.space_before = Pt(6)

                for medida in tabela.medidas:
                    _add_dax_object_header(
                        "Medida",
                        medida.nome,
                        None,
                        None,
                        medida.descricao if hasattr(medida, 'descricao') else None,
                        show_tipo=False,
                    )

                    expr = medida.expressao_dax.strip() if medida.expressao_dax else "// Expressão não capturada"
                    _add_leitura_dax_docx(analisar_dax(expr))
                    _add_code_block(expr, "DAX")

                    # Formato dinâmico (se existir)
                    if medida.formato_dinamico and str(medida.formato_dinamico).strip():
                        p_fd = doc.add_paragraph()
                        p_fd.paragraph_format.space_before = Pt(4)
                        p_fd.paragraph_format.space_after = Pt(2)
                        p_fd.paragraph_format.left_indent = Cm(0.15)
                        run_fd = p_fd.add_run("Formato Dinâmico:")
                        _set_run_font(run_fd, FONT_MAIN, FONT_DAX_TEXT, AZUL_SEC, bold=True)
                        _add_code_block(str(medida.formato_dinamico).strip(), "DAX")

            # Hierarquias
            if tabela.hierarquias:
                h = doc.add_heading(_t("docx.section.hierarchies"), level=3)
                h.paragraph_format.space_before = Pt(8)
                for hier in tabela.hierarquias:
                    niveis_str = " → ".join(hier.niveis)
                    p = doc.add_paragraph()
                    run = p.add_run(f"  {hier.nome}: ")
                    _set_run_font(run, FONT_MAIN, FONT_BODY, AZUL_PRI, bold=True)
                    _set_run_font(p.add_run(niveis_str), FONT_MAIN, FONT_BODY, CINZA_TX)

            # Fonte de Dados
            if tabela.particao and tabela.particao.codigo_fonte:
                codigo_formatado = self._formatar_codigo_fonte(tabela.particao.codigo_fonte)
                is_dax = _codigo_fonte_eh_dax(tabela.particao.codigo_fonte)
                if not is_dax:
                    _add_regra_power_query_docx(analisar_power_query_m(codigo_formatado))

                h = doc.add_heading(_t("docx.section.source_data"), level=3)
                h.paragraph_format.space_before = Pt(8)

                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r_m = p.add_run("Modo: ")
                _set_run_font(r_m, FONT_MAIN, FONT_TABLE, CINZA_LT, bold=True)
                r_v = p.add_run(tabela.particao.modo)
                _set_run_font(r_v, FONT_MAIN, FONT_TABLE, CINZA_TX)
                if tabela.particao.grupo_consulta:
                    _set_run_font(p.add_run("  │  "), FONT_MAIN, FONT_TABLE, RGBColor(0xCC, 0xCC, 0xCC))
                    r_g = p.add_run("Grupo: ")
                    _set_run_font(r_g, FONT_MAIN, FONT_TABLE, CINZA_LT, bold=True)
                    _set_run_font(p.add_run(tabela.particao.grupo_consulta), FONT_MAIN, FONT_TABLE, CINZA_TX)

                codigo_formatado = self._formatar_codigo_fonte(tabela.particao.codigo_fonte)
                lang = "DAX" if _codigo_fonte_eh_dax(codigo_formatado) else "Power Query M"
                _add_code_block(codigo_formatado, lang)


        # ====================================================================
        # VISUAIS PERSONALIZADOS
        # ====================================================================
        if self.visuais_personalizados:
            doc.add_heading(_t("docx.section.custom_visuals"), level=1)
            rows_vis = [[v] for v in self.visuais_personalizados]
            _add_table([_t("docx.tcol.visual_id")], rows_vis)

        # ====================================================================
        # RECURSOS DE IMAGEM
        # ====================================================================
        if self.recursos_imagem:
            doc.add_heading(_t("docx.section.image_resources"), level=1)
            rows_img = [[r['nome'], r['tipo']] for r in self.recursos_imagem]
            _add_table([_t("docx.tcol.img_name"), _t("docx.tcol.img_type")], rows_img)

        # ====================================================================
        # SALVAR
        # ====================================================================
        try:
            doc.save(str(caminho_saida))
            print(f"[OK] Documentação DOCX salva em: {caminho_saida}")
            return caminho_saida
        except PermissionError:
            msg_erro = f"[ERRO] Permissão negada ao salvar {caminho_saida}. O arquivo DOCX está aberto no Microsoft Word? Feche-o e tente novamente."
            print(msg_erro)
            raise PermissionError(msg_erro)


# ============================================================================
# EXEMPLO DE USO
# ============================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Uso: python documentador_pbip.py <caminho_para_projeto>")
        print("Exemplo: python documentador_pbip.py C:/Projetos/MeuDashboard")
        sys.exit(1)

    caminho_projeto = sys.argv[1]

    try:
        # Criar documentador
        print("=" * 60)
        print("DOCUMENTADOR DE PROJETOS POWER BI (.pbip)")
        print("=" * 60)

        doc = DocumentadorPBIP(caminho_projeto)

        # Extrair informações
        doc.extrair_informacoes()

        # Gerar e salvar documentação
        print("Gerando documentação...")
        caminho_doc = doc.salvar_documentacao()

        print("\n" + "=" * 60)
        print("[OK] CONCLUIDO!")
        print("=" * 60)
        print(f"\nResumo:")
        print(f"  - Tabelas: {len(doc.tabelas)}")
        print(f"  - Relacionamentos: {len(doc.relacionamentos)}")
        print(f"  - Paginas: {len(doc.paginas)}")
        print(f"  - Documentacao: {caminho_doc}")

    except Exception as e:
        print(f"\n[ERRO] {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
